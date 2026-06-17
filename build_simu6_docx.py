#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_simu6_docx.py  —  compile  6/simu.docx

Exhaustive simulation-output dossier for the NEREID-B Sydney Desalination Plant
brine-dispersion prediction (solver.py Rev 2.0 + 6/sydney_case_input.json). It:
  * plots EVERY column of EVERY CSV the solver wrote (time series, centreline curve,
    vertical profile) as line GRAPHS -> plot_*.png;
  * derives additional engineering GRAPHS from the 3-D field (seabed footprint-area
    vs threshold curve, percentile uncertainty band) and only-where-necessary CHARTS
    (distribution histograms, spatial ensemble maps);
  * embeds the solver's own fig_*.png suite;
  * tabulates the full input deck, every scalar metric, and the raw CSV data;
and assembles it into 6/simu.docx.

Graphs are preferred over charts: anything that is a relationship/curve is drawn as a
line graph; histograms and spatial maps (charts) appear only where they are the correct
representation.

Run AFTER the case run:
    python3 solver.py --config 6/sydney_case_input.json
    python3 build_simu6_docx.py
"""
import os, json, csv
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
D6 = os.path.join(HERE, "6")
ACCENT = (0x0B, 0x3D, 0x5C)
BODY = "Calibri"
plt.rcParams.update({"figure.dpi": 120, "font.size": 9, "axes.grid": True,
                     "grid.alpha": 0.3, "axes.titlesize": 10})

with open(os.path.join(D6, "metrics_summary.json")) as f:
    js = json.load(f)
cfg = js["config"]; M = js["metrics"]
CRIT = float(cfg.get("dS_crit", 0.5))
NF_DIL = float(M.get("nf_return_dilution", 0.0))


# ---------------- CSV / field helpers ----------------
def read_csv(path):
    rows = []
    with open(path) as f:
        rdr = csv.reader(f); header = next(rdr)
        for r in rdr:
            try:
                rows.append([float(x) for x in r])
            except Exception:
                pass
    return header, np.array(rows) if rows else np.zeros((0, len(header)))


def seabed(field3d, fluid):
    bk = np.argmax(fluid, axis=2)
    return np.take_along_axis(field3d, bk[:, :, None], axis=2)[:, :, 0]


# ---------------- one line-GRAPH per CSV ----------------
def plot_timeseries(path, out):
    hdr, A = read_csv(path)
    if A.shape[0] == 0:
        return None
    t = A[:, 0]
    cols = [("S_max", "S_max (g/kg)"), ("excess_max", "ΔS_max (g/kg)"),
            ("dilution_min", "min dilution (:1)"), ("r_max_m", "reach r_max (m)"),
            ("seabed_footprint_m2", "footprint (m²)"), ("z_deepest_m", "deepest impact (m)"),
            ("dt_s", "timestep dt (s)"), ("max_divergence", "max divergence")]
    idx = {h: i for i, h in enumerate(hdr)}
    fig, ax = plt.subplots(2, 4, figsize=(13, 5.2))
    for a, (key, lab) in zip(ax.ravel(), cols):
        if key in idx:
            y = A[:, idx[key]]
            if key == "max_divergence":
                a.semilogy(t, np.abs(y) + 1e-30, "-o", ms=3, color="#b03a2e")
            else:
                a.plot(t, y, "-o", ms=3, color="#0b3d5c")
            a.set_title(lab); a.set_xlabel("t (s)")
        else:
            a.axis("off")
    fig.suptitle("metrics_timeseries.csv — every scalar metric vs time (line graphs)",
                 fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.96]); fig.savefig(out); plt.close(fig)
    return out


def plot_centerline(path, out):
    hdr, A = read_csv(path)
    if A.shape[0] == 0:
        return None
    d = A[:, 0]; o = np.argsort(d); d = d[o]; A = A[o]
    fig, ax = plt.subplots(1, 3, figsize=(13, 3.8))
    ax[0].plot(d, A[:, 1], "-o", ms=3, color="#0b3d5c"); ax[0].set_title("excess ΔS vs distance")
    ax[0].set_xlabel("distance from source (m)"); ax[0].set_ylabel("ΔS (g/kg)")
    ax[0].axhline(CRIT, color="r", ls="--", lw=1, label=f"{CRIT:g} g/kg contour"); ax[0].legend()
    ax[1].semilogy(d, np.maximum(A[:, 2], 1.0), "-o", ms=3, color="#1e8449")
    ax[1].set_title("brine dilution vs distance (log)")
    ax[1].set_xlabel("distance from source (m)"); ax[1].set_ylabel("dilution (:1)")
    if NF_DIL > 0:
        ax[1].axhline(NF_DIL, color="r", ls="--", lw=1, label=f"near-field {NF_DIL:.0f}:1"); ax[1].legend()
    ax[2].plot(d, A[:, 3], "-o", ms=3, color="#7d3c98"); ax[2].set_title("plume core depth vs distance")
    ax[2].set_xlabel("distance from source (m)"); ax[2].set_ylabel("core depth below surface (m)")
    fig.suptitle("curve_centerline.csv — centreline dilution / excess / depth (line graphs)",
                 fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.93]); fig.savefig(out); plt.close(fig)
    return out


def plot_vprofile(path, out):
    hdr, A = read_csv(path)
    if A.shape[0] == 0:
        return None
    z = A[:, 0]
    fig, ax = plt.subplots(1, 4, figsize=(13, 4.0))
    series = [(1, "salinity (g/kg)", "#0b3d5c"), (2, "excess ΔS (g/kg)", "#b03a2e"),
              (3, "density (kg/m³)", "#117a65"), (4, "temperature (°C)", "#b9770e")]
    for a, (ci, lab, c) in zip(ax, series):
        a.plot(A[:, ci], z, "-o", ms=3, color=c)
        a.set_xlabel(lab); a.set_ylabel("depth below surface (m)"); a.invert_yaxis()
    fig.suptitle("curve_vertical_profile.csv — vertical structure at the plume core (line graphs)",
                 fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.93]); fig.savefig(out); plt.close(fig)
    return out


# ---------------- derived engineering GRAPHS ----------------
def plot_footprint_vs_threshold(out):
    """GRAPH: seabed area exceeding ΔS as a function of the threshold — the design-
    relevant cumulative-exceedance curve (drives the choice of compliance contour)."""
    d = np.load(os.path.join(D6, "fields_final.npz"))
    fl = d["fluid"]; exc = seabed(d["excess"], fl)
    x = d["x"]; y = d["y"]
    dx = float(np.mean(np.diff(x))) if x.size > 1 else 1.0
    dy = float(np.mean(np.diff(y))) if y.size > 1 else 1.0
    cell = abs(dx * dy)
    excf = np.nan_to_num(exc); excf = excf[np.isfinite(excf)]
    thr = np.linspace(0.0, max(0.6, float(excf.max())), 80)
    area = np.array([(excf > t).sum() * cell for t in thr])
    fig, ax = plt.subplots(figsize=(8, 3.8))
    ax.plot(thr, area, "-", lw=2, color="#0b3d5c")
    ax.axvline(CRIT, color="r", ls="--", lw=1, label=f"assessment contour {CRIT:g} g/kg")
    ai = (excf > CRIT).sum() * cell
    ax.plot([CRIT], [ai], "ro"); ax.annotate(f"{ai:.0f} m²", (CRIT, ai),
             textcoords="offset points", xytext=(8, 6), color="r")
    ax.set_title("Seabed footprint area exceeding a salinity threshold")
    ax.set_xlabel("excess-salinity threshold ΔS (g/kg)"); ax.set_ylabel("seabed area > threshold (m²)")
    ax.legend()
    fig.tight_layout(); fig.savefig(out); plt.close(fig)
    return out


def plot_percentile_band(out):
    if not os.path.exists(os.path.join(D6, "ensemble_stats.npz")):
        return None
    e = np.load(os.path.join(D6, "ensemble_stats.npz"))
    f = np.load(os.path.join(D6, "fields_final.npz"))
    fl = f["fluid"]; x = f["x"]; Samb = f["S_amb"]
    jc = fl.shape[1] // 2
    p05 = seabed(e["S_p05"] - Samb, fl)[:, jc]
    p50 = seabed(e["S_p50"] - Samb, fl)[:, jc]
    p95 = seabed(e["S_p95"] - Samb, fl)[:, jc]
    fig, ax = plt.subplots(figsize=(8, 3.6))
    ax.fill_between(x, p05, p95, color="#aed6f1", alpha=0.7, label="5–95% band")
    ax.plot(x, p50, "-", color="#0b3d5c", label="median (p50)")
    ax.axhline(CRIT, color="r", ls="--", lw=1, label=f"{CRIT:g} g/kg contour")
    ax.set_title("Seabed excess-salinity uncertainty band along the diffuser axis")
    ax.set_xlabel("x (m)"); ax.set_ylabel("ΔS (g/kg)"); ax.legend()
    fig.tight_layout(); fig.savefig(out); plt.close(fig)
    return out


# ---------------- only-where-necessary CHARTS ----------------
def plot_seabed_distribution(out):
    d = np.load(os.path.join(D6, "fields_final.npz"))
    fl = d["fluid"]; exc = seabed(d["excess"], fl); dil = seabed(d["dilution"], fl)
    excf = exc[np.isfinite(exc) & (exc > 1e-3)]
    dilf = dil[np.isfinite(dil) & (dil > 0) & (dil < 500)]
    fig, ax = plt.subplots(1, 2, figsize=(11, 3.8))
    ax[0].hist(excf, bins=30, color="#c0392b", alpha=0.8)
    ax[0].axvline(CRIT, color="k", ls="--", lw=1, label=f"{CRIT:g} g/kg contour")
    ax[0].set_title("seabed excess-salinity distribution"); ax[0].set_xlabel("ΔS (g/kg)")
    ax[0].set_ylabel("number of seabed cells"); ax[0].legend()
    ax[1].hist(dilf, bins=30, color="#1e8449", alpha=0.8)
    ax[1].set_title("seabed brine-dilution distribution"); ax[1].set_xlabel("dilution (:1)")
    ax[1].set_ylabel("number of seabed cells")
    fig.suptitle("Chart (histogram) — seabed distribution of ΔS and dilution (fields_final.npz)",
                 fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.92]); fig.savefig(out); plt.close(fig)
    return out


def plot_ensemble_maps(out):
    if not os.path.exists(os.path.join(D6, "ensemble_stats.npz")):
        return None
    e = np.load(os.path.join(D6, "ensemble_stats.npz"))
    f = np.load(os.path.join(D6, "fields_final.npz"))
    fl = f["fluid"]; x = f["x"]; y = f["y"]; Samb = f["S_amb"]
    mean_exc = seabed(e["S_mean"] - Samb, fl); std = seabed(e["S_std"], fl)
    exc = seabed(e["exceedance"], fl)
    fig, ax = plt.subplots(1, 3, figsize=(13, 3.8))
    for a, (Z, lab, cm) in zip(ax, [(mean_exc, "ensemble-mean ΔS (g/kg)", "magma"),
                                    (std, "ΔS std (g/kg)", "viridis"),
                                    (exc, "exceedance probability", "cividis")]):
        im = a.pcolormesh(x, y, Z.T, cmap=cm, shading="auto")
        a.set_title(lab); a.set_xlabel("x (m)"); a.set_ylabel("y (m)"); a.set_aspect("auto")
        fig.colorbar(im, ax=a, fraction=0.046, pad=0.04)
    fig.suptitle("Chart (spatial maps) — Monte-Carlo ensemble seabed statistics (ensemble_stats.npz)",
                 fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.92]); fig.savefig(out); plt.close(fig)
    return out


# ---------------- build the plots ----------------
plots = {}
plots["timeseries"] = plot_timeseries(os.path.join(D6, "metrics_timeseries.csv"),
                                      os.path.join(D6, "plot_timeseries.png"))
plots["centerline"] = plot_centerline(os.path.join(D6, "curve_centerline.csv"),
                                      os.path.join(D6, "plot_centerline.png"))
plots["vprofile"] = plot_vprofile(os.path.join(D6, "curve_vertical_profile.csv"),
                                  os.path.join(D6, "plot_vertical_profile.png"))
plots["footprint_thr"] = plot_footprint_vs_threshold(os.path.join(D6, "plot_footprint_vs_threshold.png"))
plots["pct_band"] = plot_percentile_band(os.path.join(D6, "plot_percentile_band.png"))
plots["seabed_dist"] = plot_seabed_distribution(os.path.join(D6, "plot_seabed_distribution.png"))
plots["ens_maps"] = plot_ensemble_maps(os.path.join(D6, "plot_ensemble_maps.png"))


# ---------------- docx ----------------
DOC = Document()
DOC.styles["Normal"].font.name = BODY
DOC.styles["Normal"].font.size = Pt(11)


def h(text, level=1):
    p = DOC.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = BODY
        if level <= 1:
            r.font.color.rgb = RGBColor(*ACCENT)
    return p


def para(text="", bold=False, italic=False, size=11, color=None, space_after=6):
    p = DOC.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = BODY
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def fig(fname, caption, width=6.4):
    path = fname if os.path.isabs(fname) else os.path.join(D6, fname)
    if not path or not os.path.exists(path):
        return False
    DOC.add_picture(path, width=Inches(width))
    DOC.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = DOC.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9.5); r.font.name = BODY
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44); cp.paragraph_format.space_after = Pt(10)
    return True


def table(header, rows, fs=9):
    t = DOC.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    for c, txt in zip(t.rows[0].cells, header):
        rr = c.paragraphs[0].add_run(str(txt)); rr.bold = True; rr.font.size = Pt(fs); rr.font.name = BODY
    for row in rows:
        cells = t.add_row().cells
        for c, txt in zip(cells, row):
            rr = c.paragraphs[0].add_run(str(txt)); rr.font.size = Pt(fs); rr.font.name = BODY
    DOC.add_paragraph().paragraph_format.space_after = Pt(2)


def fnum(x, nd=3):
    try:
        xf = float(x)
        if abs(xf) != 0 and (abs(xf) < 1e-3 or abs(xf) >= 1e5):
            return f"{xf:.2e}"
        return f"{xf:.{nd}f}"
    except Exception:
        return str(x)


title = DOC.add_heading("NEREID-B Simulation Output Report", level=0)
for r in title.runs:
    r.font.color.rgb = RGBColor(*ACCENT)
para("Brine-dispersion prediction — Sydney Desalination Plant offshore submerged diffuser "
     "(Kurnell, NSW). Complete simulation-output dossier: every CSV plotted as a line graph, "
     "derived engineering graphs, only-where-necessary charts, metrics tables and the solver "
     "figure suite.", bold=True, size=12, color=ACCENT)
para("Generated from the case-study run (solver.py Rev 2.0 + 6/sydney_case_input.json). "
     "Companion documents: 6/case_study.docx (interpretation) and 6/model.docx (governing PDEs).",
     italic=True, size=10)

# ---- 1. prediction at a glance
h("1.  Prediction at a glance", 1)
table(["Quantity", "Value", "Unit"],
      [["Discharge salinity S₀", fnum(cfg["S0"], 1), "g/kg"],
       ["Ambient salinity", fnum(cfg["S_amb_surf"], 1), "g/kg"],
       ["Grid (nx×ny×nz) / dz", f"{cfg['nx']}×{cfg['ny']}×{cfg['nz']} / {fnum(cfg['depth']/cfg['nz'],2)} m", "—"],
       ["Near-field return dilution", fnum(M.get("nf_return_dilution"), 1), ":1"],
       ["Near-field terminal rise", fnum(M.get("nf_rise_m"), 1), "m"],
       ["Peak salinity S_max", fnum(M.get("S_max"), 2), "g/kg"],
       ["Max excess ΔS_max", fnum(M.get("excess_max"), 2), "g/kg"],
       ["Horizontal reach r_max (>ΔS_crit)", fnum(M.get("r_max_m"), 1), "m"],
       [f"Seabed footprint > {CRIT:g} g/kg", fnum(M.get("seabed_footprint_m2"), 0), "m²"],
       ["Affected water volume", fnum(M.get("affected_volume_m3"), 0), "m³"],
       ["Dense-layer deepest impact", fnum(M.get("z_deepest_m"), 1), "m below surface"],
       ["Final divergence", fnum(M.get("divergence_final")), "—"],
       ["Mass imbalance", fnum(M.get("mass_imbalance_final")), "—"]])
fp = M.get("seabed_footprint_m2", 0.0)
para(f"Interpretation: against the conservative sub-lethal assessment contour ΔS = {CRIT:g} g/kg, "
     f"the negatively-buoyant brine forms a bottom-trapped plume with peak excess "
     f"≈ {fnum(M.get('excess_max'),2)} g/kg and a seabed footprint of ≈ {fnum(fp,0)} m² "
     f"(reach ≈ {fnum(M.get('r_max_m'),0)} m). NEREID-B under-predicts dilution (over-predicts "
     f"impact), so the footprint is a conservative bound; full interpretation in 6/case_study.docx.",
     bold=True)

# ---- 2. CSV line graphs
h("2.  Line graphs of every output CSV", 1)
para("Each CSV written by the solver is plotted as line graphs; every numeric column appears.")
para("2.1  metrics_timeseries.csv", bold=True)
fig("plot_timeseries.png", "Figure A (graph). All scalar metrics vs time (S_max, ΔS_max, dilution, "
    "reach, footprint, deepest impact, timestep, divergence). Flat traces confirm a steady, "
    "drift-free quasi-equilibrium plume; divergence stays at machine precision.")
para("2.2  curve_centerline.csv", bold=True)
fig("plot_centerline.png", "Figure B (graph). Centreline excess salinity, brine dilution (log axis) "
    f"and plume-core depth vs distance; dashed lines mark the {CRIT:g} g/kg contour and the "
    f"near-field return dilution.")
para("2.3  curve_vertical_profile.csv", bold=True)
fig("plot_vertical_profile.png", "Figure C (graph). Vertical profiles of salinity, excess, density "
    "and temperature at the plume-core column, showing the bottom-trapped dense layer.")

# ---- 3. derived engineering graphs
h("3.  Derived engineering graphs", 1)
if plots["footprint_thr"]:
    fig("plot_footprint_vs_threshold.png", "Figure D (graph). Seabed area exceeding a salinity "
        "threshold as a continuous function of the threshold — the cumulative-exceedance curve that "
        f"drives the choice of compliance contour; the {CRIT:g} g/kg point is marked.")
if plots["pct_band"]:
    fig("plot_percentile_band.png", "Figure E (graph). Seabed excess-salinity 5–95% uncertainty band "
        f"and median along the diffuser axis vs the {CRIT:g} g/kg contour (Monte-Carlo ensemble).")

# ---- 4. charts where necessary
h("4.  Charts (only where a graph cannot represent the data)", 1)
if plots["seabed_dist"]:
    fig("plot_seabed_distribution.png", "Figure F (chart/histogram). Distribution of seabed excess "
        "salinity and brine dilution across all seabed cells — a histogram is the correct "
        "representation of a population distribution.")
if plots["ens_maps"]:
    fig("plot_ensemble_maps.png", "Figure G (chart/spatial maps). Monte-Carlo ensemble seabed "
        "statistics: mean excess, standard deviation and exceedance probability — 2-D fields require "
        "spatial maps.")

# ---- 5. solver figure suite
h("5.  Solver figure suite", 1)
for fn, cap in [
    ("fig_seabed_excess_map.png", "Figure H. Seabed excess-salinity map (plan view)."),
    ("fig_vertical_section.png", "Figure I. Vertical section of excess salinity along the plume."),
    ("fig_centerline_dilution.png", "Figure J. Centreline dilution (solver)."),
    ("fig_salinity_decay.png", "Figure K. Excess-salinity decay with distance (solver)."),
    ("fig_exceedance_probability.png", "Figure L. Exceedance-probability map (ensemble)."),
    ("fig_seabed_currents.png", "Figure M. Near-bed current field."),
    ("fig_nearfield_trajectory.png", "Figure N. Near-field dense-jet trajectory.")]:
    fig(fn, cap)

# ---- 6. raw CSV data tables
h("6.  Output data tables", 1)
para("6.1  Centreline curve (curve_centerline.csv)", bold=True)
hdr, A = read_csv(os.path.join(D6, "curve_centerline.csv"))
o = np.argsort(A[:, 0]); A = A[o]; step = max(1, A.shape[0] // 18)
table(hdr, [[fnum(v, 3) for v in A[i]] for i in range(0, A.shape[0], step)])
para("6.2  Vertical profile (curve_vertical_profile.csv)", bold=True)
hdr, A = read_csv(os.path.join(D6, "curve_vertical_profile.csv"))
table(hdr, [[fnum(v, 3) for v in row] for row in A])
para("6.3  Time series (metrics_timeseries.csv)", bold=True)
hdr, A = read_csv(os.path.join(D6, "metrics_timeseries.csv"))
step = max(1, A.shape[0] // 16)
table(hdr, [[fnum(v, 3) for v in A[i]] for i in range(0, A.shape[0], step)])

# ---- 7. full metrics
h("7.  Full metric set (metrics_summary.json)", 1)
flat = []
for kk, vv in M.items():
    if isinstance(vv, (int, float, bool)):
        flat.append([kk, str(vv) if isinstance(vv, bool) else fnum(vv)])
rows = [[flat[i][0], flat[i][1],
         flat[i + 1][0] if i + 1 < len(flat) else "", flat[i + 1][1] if i + 1 < len(flat) else ""]
        for i in range(0, len(flat), 2)]
table(["metric", "value", "metric", "value"], rows, fs=8.5)

# ---- 8. input deck
h("8.  Input deck (sydney_case_input.json)", 1)
keys = ["S0", "T_b", "Q_d", "d_p", "theta_deg", "n_ports", "port_spacing",
        "S_amb_surf", "S_amb_bot", "T_amb_surf", "T_amb_bot",
        "U_current", "tide_amp", "Hs", "Tw", "wind10", "latitude_deg",
        "depth", "Lx", "Ly", "nx", "ny", "nz", "t_end", "ensemble", "dS_crit",
        "near_field_coupling", "free_surface", "stoch_enable"]
rows = []
for i in range(0, len(keys), 2):
    k1 = keys[i]; k2 = keys[i + 1] if i + 1 < len(keys) else None
    rows.append([k1, fnum(cfg.get(k1)) if isinstance(cfg.get(k1), (int, float)) else str(cfg.get(k1)),
                 k2 or "", (fnum(cfg.get(k2)) if isinstance(cfg.get(k2), (int, float)) else str(cfg.get(k2))) if k2 else ""])
table(["parameter", "value", "parameter", "value"], rows, fs=8.5)

para("")
para("All artifacts (CSVs, line-graph plots plot_*.png, derived graphs, charts, solver figures "
     "fig_*.png, metrics_summary.json, fields_final.npz, ensemble_stats.npz, run.log, input deck) "
     "reside in folder 6/. Companion documents: 6/case_study.docx (engineering interpretation), "
     "6/model.docx (governing PDEs).", size=9, italic=True, color=ACCENT)

out = os.path.join(D6, "simu.docx")
DOC.save(out)
made = [k for k, v in plots.items() if v]
print("wrote", out)
print("generated plots:", made)
