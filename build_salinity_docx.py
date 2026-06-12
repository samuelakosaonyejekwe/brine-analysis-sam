# -*- coding: utf-8 -*-
"""
Generator for salinity.docx — the NEREID-B coupled stochastic PDE model
for salinity distribution around a negatively-buoyant brine outfall.

Produces a native Microsoft Word document using python-docx.
Equations are rendered in Unicode mathematical notation (Cambria Math),
each numbered and displayed on its own centered line.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = Document()

# ----------------------------------------------------------------------
# Styling helpers
# ----------------------------------------------------------------------
EQ_FONT = "Cambria Math"
BODY_FONT = "Calibri"

# base style
normal = DOC.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)

_eq_counter = [0]


def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def h(text, level=1):
    p = DOC.add_heading(text, level=level)
    return p


def para(text="", bold=False, italic=False, size=11, align=None, color=None,
         space_after=6):
    p = DOC.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = BODY_FONT
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def rich(parts, align=None, space_after=6):
    """parts: list of (text, dict-of-attrs)."""
    p = DOC.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for text, attrs in parts:
        r = p.add_run(text)
        r.bold = attrs.get("bold", False)
        r.italic = attrs.get("italic", False)
        r.font.size = Pt(attrs.get("size", 11))
        r.font.name = attrs.get("font", BODY_FONT)
        if attrs.get("color"):
            r.font.color.rgb = RGBColor(*attrs["color"])
    return p


def eq(expr, label=True, indent=0.3):
    """Display equation, centered, numbered."""
    p = DOC.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    # equation text, centered via a tab structure: we just center it
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(expr)
    r.font.name = EQ_FONT
    r.font.size = Pt(12)
    # set east-asian / math font too
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), EQ_FONT)
    rFonts.set(qn("w:hAnsi"), EQ_FONT)
    rFonts.set(qn("w:cs"), EQ_FONT)
    if label:
        _eq_counter[0] += 1
        tab = p.add_run("\t\t(" + str(_eq_counter[0]) + ")")
        tab.font.name = BODY_FONT
        tab.font.size = Pt(11)
    return p


def bullet(text, sub=False):
    style = "List Bullet 2" if sub else "List Bullet"
    p = DOC.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def numbered(text):
    p = DOC.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def hrule():
    para("", space_after=2)


def make_table(header, rows, col_widths=None, font_size=9):
    t = DOC.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = BODY_FONT
        _set_cell_bg(hdr[i], "1F4E79")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = EQ_FONT if i == 0 else BODY_FONT
    if col_widths:
        for i, w in enumerate(col_widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)
    return t


# ======================================================================
#  TITLE PAGE
# ======================================================================
title = DOC.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("NEREID-B")
tr.bold = True
tr.font.size = Pt(30)
tr.font.name = BODY_FONT
tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = DOC.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Nonlinear Eulerian Reactive-osmotic Effluent "
                 "Integro-Dispersion model for Brine outfalls")
sr.italic = True
sr.font.size = Pt(15)
sr.font.name = BODY_FONT

sub2 = DOC.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr2 = sub2.add_run("A Novel Unsteady, Stochastic, Coupled "
                   "Hydro–Haline–Thermal Partial-Differential-Equation System "
                   "for Predicting the Three-Dimensional, Time-Evolving "
                   "Salinity Distribution of a Negatively-Buoyant Brine Plume "
                   "Discharged from a Submarine Pipe into a Moving Sea")
sr2.font.size = Pt(12)
sr2.font.name = BODY_FONT

para("", space_after=10)
meta = DOC.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Mathematical Engineering Specification  ·  Rev. 1.0\n"
                  "Prepared: 11 June 2026")
mr.font.size = Pt(11)
mr.font.name = BODY_FONT
mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

DOC.add_page_break()

# ======================================================================
#  ABSTRACT
# ======================================================================
h("Abstract", 1)
para(
    "NEREID-B is a fully-coupled, unsteady, non-Boussinesq, stochastic "
    "partial-differential-equation (PDE) system that predicts the three-"
    "dimensional space-time evolution of salinity produced when a "
    "concentrated brine effluent (e.g. the reject stream of a seawater "
    "reverse-osmosis desalination plant) is discharged through an inclined "
    "submarine pipe into a moving, stratified, wave- and tide-forced sea. "
    "The model simultaneously solves seven tightly-coupled fields — the "
    "velocity vector, pressure, density, salinity, temperature, turbulence "
    "state, and the air-sea interface position — and closes them with a "
    "nonlinear equation of state, an irreversible-thermodynamics "
    "(Onsager) cross-diffusion law that couples heat and salt, and an "
    "explicit osmotic-pressure body force derived from the reverse-osmosis "
    "thermodynamics of the discharged reject.",
)
para(
    "Three features make the formulation novel and patentable: (i) the "
    "osmotic-pressure gradient of the hyper-saline reject is promoted to a "
    "first-class momentum and mass-flux term rather than being neglected as "
    "in all conventional plume models; (ii) the unpredictability of the sea, "
    "wind and turbulence is represented intrinsically as space-time colored-"
    "noise stochastic forcing, so that the model returns a full probability "
    "density of the salinity field — a predicted mean plume together with a "
    "quantified confidence envelope of how far and how deep the brine "
    "reaches — rather than a single deterministic answer; and (iii) salinity "
    "and temperature are coupled through a complete anisotropic, state-"
    "dependent dispersion tensor that folds together molecular diffusion, "
    "shear dispersion, wave stirring, and bathymetric steering. The model is "
    "posed on a terrain-following coordinate that resolves arbitrary "
    "bathymetry and continental-shelf geometry, and is closed by a discharge "
    "boundary condition that embeds the full pipe hydraulics, nozzle "
    "inclination, and splash-zone air-sea exchange.",
)

# ======================================================================
#  1. INTRODUCTION & NOVELTY
# ======================================================================
h("1. Problem Statement, Scope and Novelty", 1)

h("1.1 Physical problem", 2)
para(
    "A submarine outfall pipe of length L_p, internal diameter d_p and "
    "terminal inclination θ discharges a brine effluent of salinity S₀ "
    "(typically 55–80 g kg⁻¹ for SWRO reject, against an ambient of "
    "≈35–39 g kg⁻¹) and momentum into a receiving sea. Because the reject is "
    "denser than the ambient, the discharge forms a negatively-buoyant jet "
    "(a 'dense jet' or 'inverted plume') that rises, bends over under gravity, "
    "falls back, impacts the seabed, and then spreads as a stratified gravity "
    "current that hugs the bathymetry. The sea is never static: currents, "
    "tides, waves, wind-driven shear, vortices and turbulence continuously "
    "advect and dilute the plume, and the air-sea interface (the splash zone) "
    "exchanges momentum, heat and freshwater (evaporation) with the water "
    "column. The engineering questions are: at every instant, what is the "
    "salinity at every point; how far horizontally and how deep does the "
    "elevated-salinity field extend; and — because the sea is unpredictable — "
    "with what probability is a regulatory salinity threshold exceeded at a "
    "given location?",
)

h("1.2 Why existing models are insufficient", 2)
bullet("Integral/entrainment jet models (e.g. CORMIX, VISJET, JETLAG) "
       "collapse the plume to a 1-D centerline and cannot resolve the "
       "3-D unsteady field, bathymetric steering, or the full sea state.")
bullet("Boussinesq Reynolds-Averaged Navier-Stokes (RANS) models assume "
       "small density differences and therefore lose accuracy for hyper-"
       "saline reject, where (ρ−ρ₀)/ρ₀ can exceed 0.04.")
bullet("All conventional models neglect osmotic pressure and the "
       "irreversible-thermodynamics coupling of heat and salt, and treat the "
       "sea deterministically, giving no measure of predictive uncertainty.")

h("1.3 Statement of novelty (patentable claims)", 2)
para("NEREID-B advances the state of the art through the following coupled "
     "innovations, each of which is, to the authors' knowledge, absent from "
     "prior salinity-dispersion models:", bold=False)
numbered("Osmotic-hydrodynamic coupling: the gradient of the osmotic pressure "
         "Π(S,T) of the hyper-saline field enters the momentum equation as a "
         "thermodynamic body force and the salt equation as an osmotic flux, "
         "capturing reverse-osmosis-driven micro-transport at the brine front.")
numbered("Intrinsic stochastic forcing: ocean currents, tides, wind and "
         "sub-grid turbulence are represented by Ornstein–Uhlenbeck colored-"
         "noise PDE forcing, so the model is a genuine stochastic PDE (SPDE) "
         "whose solution is the probability density functional of the salinity "
         "field, P[S(x,t)].")
numbered("Non-Boussinesq variable-density formulation valid for arbitrarily "
         "large brine/ambient density contrast.")
numbered("Onsager cross-diffusion: Soret (thermo-diffusion) and Dufour "
         "(diffusion-thermo) terms couple the salt and heat equations beyond "
         "their density coupling.")
numbered("Fully anisotropic, state-dependent dispersion tensor D_S that "
         "fuses molecular diffusion, turbulent diffusion, shear (Taylor) "
         "dispersion, wave-orbital stirring and bathymetric anisotropy.")
numbered("Embedded pipe-hydraulics and splash-zone boundary operators that "
         "couple the internal pipe flow, the inclined nozzle jet, and the "
         "two-phase air-sea interface directly to the field equations.")

# ======================================================================
#  2. DOMAIN, COORDINATES, NOTATION
# ======================================================================
h("2. Computational Domain, Coordinates and Bathymetry", 1)
para(
    "Let x = (x, y, z) be a right-handed Cartesian frame with z positive "
    "upward and z = 0 at the still-water level. The seabed is z = −H(x, y), "
    "where H is the bathymetry; the continental-shelf and local terrain "
    "geometry are supplied as the field H and its slope ∇H. The instantaneous "
    "free surface is z = η(x, y, t). To resolve sloping bathymetry and the "
    "moving surface exactly, the vertical coordinate is mapped to a terrain-"
    "following sigma coordinate σ ∈ [−1, 0]:",
)
eq("σ = (z − η) / (H + η),   so that  σ = 0 at the surface and σ = −1 at the bed.")
para(
    "All field equations below are written in physical (x, y, z) form for "
    "clarity; the model is integrated in the (x, y, σ, t) image where the "
    "domain is a fixed prism. The total water depth is D(x,y,t) = H + η. The "
    "pipe terminus (the diffuser nozzle) is located at x_d = (x_d, y_d, z_d) "
    "with z_d = −H(x_d,y_d) + h_n, h_n being the nozzle height above the bed.",
)

# ======================================================================
#  3. GOVERNING COUPLED PDE SYSTEM
# ======================================================================
h("3. The Coupled Governing PDE System", 1)
para("The state vector solved by NEREID-B is",)
eq("q(x,t) = ( ρ,  u = (u,v,w),  p,  S,  T,  k,  ε,  η,  α )ᵀ ,")
para("comprising density ρ, velocity u, pressure p, absolute salinity S, "
     "conservative temperature T, turbulent kinetic energy k, its dissipation "
     "ε, the free-surface elevation η, and the air-water volume fraction α. "
     "The governing balance laws follow.")

# 3.1 Mass
h("3.1 Conservation of mass (non-Boussinesq continuity)", 2)
para("Because the brine/ambient density contrast is not small, mass is "
     "conserved in fully compressible-in-density (non-Boussinesq) form:")
eq("∂ρ/∂t + ∇·(ρ u) = 0 .")
para("The incompressible Boussinesq limit ∇·u = 0 is recovered automatically "
     "where |ρ − ρ₀|/ρ₀ ≪ 1, but is never assumed.")

# 3.2 Momentum
h("3.2 Momentum balance (Reynolds-averaged Navier–Stokes with buoyancy, "
  "rotation, waves, osmosis and stochastic forcing)", 2)
eq("∂(ρu)/∂t + ∇·(ρ u⊗u) = −∇p + ∇·τ + ρg − 2ρ Ω×u + F_wave + F_osm + F_stoch .")
para("Term by term:")
bullet("Deviatoric + turbulent (Reynolds) stress, with eddy viscosity μ_t:")
eq("τ = (μ + μ_t)(∇u + ∇uᵀ − (2/3)(∇·u) I) − (2/3) ρ k I .")
bullet("Buoyancy: ρg with g = −g ẑ; the active driver of the dense plume "
       "because ρ depends on S and T through the equation of state (§4.1).")
bullet("Coriolis (Earth rotation): −2ρ Ω×u, Ω = Ω(0, cosφ, sinφ), φ = latitude.")
bullet("Wave forcing as the divergence of the radiation-stress tensor S_rad "
       "(depth-dependent Longuet-Higgins / vortex-force form):")
eq("F_wave = −∇·S_rad + ρ (u_St × ω) ,   u_St = Stokes drift,  ω = ∇×u .")
bullet("Osmotic body force (NOVEL) — the gradient of osmotic pressure Π of "
       "the saline field acts on the fluid through the partial molar volume "
       "of water v̄_w:")
eq("F_osm = − (φ_v / v̄_w) ∇Π(S,T) ,   φ_v = local solute volume fraction .")
bullet("Stochastic forcing F_stoch (NOVEL) — see §3.7; represents the "
       "unresolved/unpredictable momentum input from gusts, eddies and "
       "internal waves.")

# 3.3 Salinity
h("3.3 Absolute-salinity transport (advection – anisotropic dispersion – "
  "osmotic & Soret flux – reaction – noise)", 2)
para("Salinity is transported by advection and by the divergence of a "
     "generalized flux J_S built from Onsager irreversible thermodynamics:")
eq("∂S/∂t + u·∇S = −(1/ρ) ∇·J_S + R_S + ξ_S ,")
eq("J_S = −ρ D_S·∇S  −  ρ D_ST ∇T  −  ρ (D_S L_p / R_g T) ∇Π .")
para("The three flux contributions are, respectively: anisotropic dispersion "
     "(D_S, a full tensor, §4.2); Soret thermo-diffusion D_ST that lets "
     "temperature gradients drive salt (cross-coupling to §3.4); and the "
     "osmotic/reverse-osmosis flux (NOVEL), in which the osmotic-pressure "
     "gradient drives an additional salt flux with permeability coefficient "
     "L_p — the open-water analogue of reverse-osmosis transport across the "
     "brine front. R_S is any reactive/source term (normally 0 for salt; "
     "retained for generality, e.g. anti-scalant tracers), and ξ_S is "
     "scalar stochastic forcing.")

# 3.4 Temperature
h("3.4 Conservative-temperature (heat) transport with Dufour coupling", 2)
eq("∂T/∂t + u·∇T = −(1/ρ c_p) ∇·J_Q + Q_rad/(ρ c_p) + ξ_T ,")
eq("J_Q = −ρ c_p D_T·∇T  −  ρ c_p D_TS ∇S .")
para("Here D_T is the anisotropic thermal dispersion tensor, D_TS is the "
     "Dufour (diffusion-thermo) coefficient that closes the reciprocal "
     "Onsager pair with D_ST, Q_rad is solar/longwave radiative heating, and "
     "ξ_T is thermal stochastic forcing. Equations (3.3) and (3.4) are thus "
     "two-way coupled both through the density (buoyancy) and directly "
     "through the cross-diffusion pair (D_ST, D_TS).")

# 3.5 EOS
h("3.5 Nonlinear equation of state (haline–thermal–baric closure)", 2)
para("Density is closed by the full nonlinear TEOS-10 / Gibbs-function "
     "equation of state — not a linearized approximation — because the dense "
     "plume's dynamics are exquisitely sensitive to small density errors:")
eq("ρ = ρ_EOS(S, T, p) = ρ₀ [ 1 + β_S(S−S₀) − α_T(T−T₀) + γ_p(p−p₀) + Φ_nl ] ,")
para("where α_T is the thermal-expansion, β_S the haline-contraction and "
     "γ_p the compressibility coefficient, all themselves functions of "
     "(S,T,p), and Φ_nl collects the cabbeling/thermobaric nonlinear terms of "
     "TEOS-10. This single relation is the master coupling that ties the "
     "momentum, salinity and temperature equations together.")

# 3.6 Turbulence
h("3.6 Turbulence closure (buoyancy-modified, stratification-damped k–ε)", 2)
para("Turbulent mixing — which controls dilution — is closed by a two-"
     "equation k–ε model modified for buoyancy so that the stable "
     "stratification of the dense plume correctly suppresses vertical mixing:")
eq("∂(ρk)/∂t + ∇·(ρu k) = ∇·[(μ + μ_t/σ_k)∇k] + P_k + G_b − ρε ,")
eq("∂(ρε)/∂t + ∇·(ρu ε) = ∇·[(μ + μ_t/σ_ε)∇ε] + (ε/k)[C₁(P_k + C₃ G_b) − C₂ ρε] ,")
eq("μ_t = ρ C_μ k²/ε ,   P_k = τ : ∇u ,   G_b = (μ_t/ρ Sc_t) g·∇ρ .")
para("The buoyancy-production term G_b is negative in stable stratification, "
     "damping turbulence exactly where the brine layer is densest — the "
     "physical mechanism that makes brine pool on the seabed. Standard "
     "constants (C_μ=0.09, C₁=1.44, C₂=1.92, σ_k=1.0, σ_ε=1.3, Sc_t=0.7) are "
     "defaults; C₃ controls stratified mixing.")

# 3.7 Stochastic
h("3.7 Stochastic forcing of an unpredictable sea (the SPDE layer)", 2)
para("The sea, wind and turbulence are not deterministic. NEREID-B promotes "
     "the system to a stochastic PDE by driving the momentum and scalar "
     "equations with space-time colored noise generated by Ornstein–Uhlenbeck "
     "(OU) processes, one per forcing channel m (currents, tide phase, wind, "
     "sub-grid turbulence):")
eq("∂ζ_m/∂t = −ζ_m/τ_m + σ_m 𝓛_m^{1/2} Ẇ_m(x,t) ,")
para("where τ_m is the de-correlation time, σ_m the intensity, 𝓛_m a spatial "
     "smoothing (covariance) operator giving the noise a realistic spatial "
     "correlation length ℓ_m, and Ẇ_m space-time white noise. The forcings in "
     "§3.2–3.4 are then")
eq("F_stoch = ρ Σ_m a_m ζ_m ,   ξ_S = Σ_m b_m ζ_m·∇S ,   ξ_T = Σ_m c_m ζ_m·∇T .")
para("Solving the ensemble (Monte-Carlo over realizations of Ẇ_m, or the "
     "associated Fokker–Planck moment equations) yields the predictive "
     "probability density of the salinity field:")
eq("P[S(x,t)] ,  with mean ⟨S⟩, variance σ_S², and exceedance probability "
   "ℙ(S(x,t) > S_crit) .")
para("This is the formal mechanism by which NEREID-B 'captures "
     "unpredictability': it predicts not a single number but a calibrated "
     "distribution of how far and how deep the brine plume may reach.")

# ======================================================================
#  4. CONSTITUTIVE / CLOSURE RELATIONS
# ======================================================================
h("4. Constitutive and Closure Relations", 1)

h("4.1 Osmotic pressure (Pitzer-corrected van 't Hoff)", 2)
para("The osmotic pressure that drives F_osm and the osmotic salt flux is "
     "computed from a van 't Hoff law with an activity (osmotic) coefficient "
     "φ_os(S,T) that corrects for the strong non-ideality of concentrated "
     "brine (Pitzer model):")
eq("Π(S,T) = φ_os(S,T) · ν · (ρ_w S / M_s) · R_g T ,")
para("with ν the number of dissolved ions per formula unit of salt, M_s the "
     "mean molar mass of the dissolved salts, R_g the gas constant and ρ_w "
     "the water density. For seawater Π ≈ 2.7 MPa at S = 35 g kg⁻¹ and rises "
     "steeply for reject brine, so ∇Π across the brine front is dynamically "
     "non-negligible — the basis of the osmotic coupling.")

h("4.2 The anisotropic, state-dependent dispersion tensor (NOVEL closure)", 2)
para("The salt dispersion tensor fuses every relevant mixing mechanism into a "
     "single symmetric positive-definite tensor aligned with the local flow, "
     "shear, wave and bathymetric directions:")
eq("D_S = D_mol I + D_turb(k,ε) I + D_shear(|∇u|) ê_u⊗ê_u "
   "+ D_wave(H_s,T_w) ê_w⊗ê_w + D_bath(∇H) (I − n̂⊗n̂) .")
para("The five contributions are molecular diffusion (D_mol), isotropic "
     "turbulent diffusion (D_turb = μ_t/ρ Sc_t), longitudinal shear/Taylor "
     "dispersion along the flow unit-vector ê_u, wave-orbital stirring along "
     "the wave direction ê_w (a function of significant wave height H_s and "
     "period T_w), and an along-slope enhancement tangent to the bed (n̂ = "
     "bed normal) that steers the dense gravity current along the bathymetry. "
     "The thermal tensor D_T has the identical structure with the thermal "
     "molecular diffusivity and turbulent Prandtl number.")

# ======================================================================
#  5. BOUNDARY & INITIAL CONDITIONS
# ======================================================================
h("5. Boundary and Initial Conditions", 1)

h("5.1 Discharge (diffuser-nozzle) condition with embedded pipe hydraulics", 2)
para("At the nozzle Γ_d the jet velocity, direction and salinity are imposed. "
     "The exit speed U_d and pressure are obtained from the internal pipe "
     "hydraulics (Darcy–Weisbach) so that pipe length, diameter, roughness "
     "and inclination feed the field directly:")
eq("U_d = Q_d /(π d_p²/4) ,   p_d = p_man − ρ_b g Δz_p "
   "− (f L_p/d_p + Σ K_minor)(ρ_b U_d²/2) ,")
eq("u|_Γd = U_d (cosθ cosψ, cosθ sinψ, sinθ) ,   S|_Γd = S₀ ,   T|_Γd = T_b ,")
para("where Q_d is the discharge flow rate, f the Darcy friction factor "
     "(roughness-/Reynolds-dependent), Δz_p the pipe rise, K_minor the minor-"
     "loss coefficients, p_man the manifold pressure, ρ_b the brine density, "
     "θ the nozzle elevation (inclination) angle and ψ its azimuth. The jet "
     "momentum flux M₀ = ρ_b U_d² A_d and buoyancy flux B₀ = g(ρ_b−ρ_a)/ρ_a · "
     "Q_d set the densimetric Froude number that governs the trajectory.")

h("5.2 Free surface and splash zone (two-phase air–sea coupling)", 2)
para("The air-water interface is tracked by a volume-of-fluid fraction α "
     "(α=1 water, 0 air), advected with the flow, which resolves the splash "
     "zone and surface deformation:")
eq("∂α/∂t + ∇·(α u) = 0 ,   ρ = α ρ_water + (1−α) ρ_air .")
para("Across the interface the dynamic and kinematic conditions, wind stress, "
     "and the evaporative salinity flux apply:")
eq("∂η/∂t + u_h·∇η = w ,   (kinematic)")
eq("τ_surf = ρ_a C_d |U₁₀ − u_h|(U₁₀ − u_h) ,   (wind stress, U₁₀ = 10-m wind)")
eq("(ρ D_S ∇S)·n̂ = S (E − P)/ρ_w ,   (evaporation E minus precipitation P "
   "concentrates salt) .")
para("The air velocity / wind field U₁₀(x,t) is itself one of the stochastic "
     "OU channels of §3.7, so wind unpredictability propagates into the "
     "surface stress and the splash-zone mixing.")

h("5.3 Seabed and shoreline", 2)
eq("u·n̂ = 0 (no-normal-flow) , wall-function τ_bed = ρ C_D^bed |u_b| u_b , "
   "(D_S ∇S)·n̂ = 0 (no salt flux through bed) .")
para("The bed-normal n̂ follows the bathymetry/terrain ∇H, so the dense "
     "gravity current is steered along the real continental-shelf slope.")

h("5.4 Open-ocean lateral boundaries (currents, tides)", 2)
para("Far-field boundaries impose the ambient stratification and the "
     "current/tidal forcing, again with a stochastic component:")
eq("u|_∞ = U_current(z,t) + U_tide(t) + ζ_current ,   S|_∞ = S_amb(z) ,   "
   "T|_∞ = T_amb(z) ,")
eq("U_tide(t) = Σ_j A_j cos(ω_j t + φ_j)   (harmonic tidal constituents) .")

h("5.5 Initial condition", 2)
eq("q(x,0) = q_amb(x): quiescent or measured ambient profiles of "
   "S, T, ρ, u; plume fields seeded at the nozzle.")

# ======================================================================
#  6. COUPLING MAP
# ======================================================================
h("6. Coupling Structure", 1)
para("The seven balance laws are not solved in isolation; they form a single "
     "coupled operator. The principal coupling pathways are:")
bullet("ρ ↔ (S,T,p): the nonlinear EOS (3.5) — salt and heat set density, "
       "density sets buoyancy in the momentum equation (3.2).")
bullet("u ↔ (S,T): velocity advects salt and heat (3.3, 3.4); their gradients "
       "set buoyancy that drives velocity (3.2).")
bullet("S ↔ T: direct Onsager cross-diffusion (Soret D_ST, Dufour D_TS) "
       "beyond the density route.")
bullet("(k,ε) ↔ all: turbulence sets μ_t, D_turb in every equation; mean "
       "shear and buoyancy gradients set k and ε (3.6).")
bullet("Π ↔ (u,S): osmotic pressure of the saline field forces momentum and "
       "salt flux (F_osm, J_S) — the novel osmotic loop.")
bullet("ζ_m ↔ all: stochastic channels perturb momentum, scalars and "
       "boundaries, turning the deterministic PDE into an SPDE.")

# ======================================================================
#  7. OUTPUT FUNCTIONALS
# ======================================================================
h("7. Predicted Outputs (Diagnostic Functionals)", 1)
para("From the solved fields, NEREID-B reports the engineering quantities of "
     "interest, each as a time series and as a probability distribution:")
eq("Dilution:  𝒟(x,t) = (S₀ − S_amb)/(S(x,t) − S_amb) ,")
eq("Excess-salinity footprint:  𝒜(t) = ∫∫ 𝟙[ S(x,y,z*,t) − S_amb > ΔS_crit ] dx dy ,")
eq("Maximum reach:  r_max(t) = max{ |x_h − x_d| : S − S_amb > ΔS_crit } ,")
eq("Maximum depth of impact:  z_max(t) = min{ z : S − S_amb > ΔS_crit } ,")
eq("Seabed impact (return) point and the gravity-current run-out distance ,")
eq("Exceedance risk:  ℝ(x,t) = ℙ[ S(x,t) − S_amb > ΔS_crit ] from the SPDE ensemble .")
para("ΔS_crit is the regulatory salinity-increment limit (commonly +2 to +5 % "
     "of ambient, or a fixed +2 g kg⁻¹ above background). The pair "
     "(r_max, z_max) with its confidence envelope is the direct answer to "
     "'how far and how deep does the brine go'.")

# ======================================================================
#  8. DIMENSIONLESS NUMBERS
# ======================================================================
h("8. Governing Dimensionless Groups", 1)
para("Non-dimensionalizing the system reveals the controlling parameters, "
     "useful for scaling, validation and reduced-order use:")
make_table(
    ["Group", "Definition", "Physical role"],
    [
        ["Fr_d", "U_d /√(g' d_p)", "Densimetric Froude — jet momentum vs buoyancy; sets rise height"],
        ["g'", "g (ρ_b − ρ_a)/ρ_a", "Reduced gravity of the reject"],
        ["Re", "ρ U_d d_p /μ", "Reynolds — turbulence onset"],
        ["Ri", "g'/(∂u/∂z)² · ∂ρ/∂z", "Richardson — stratified mixing suppression"],
        ["Pe", "U_d d_p /D_mol", "Péclet — advection vs diffusion of salt"],
        ["Sc", "ν/D_mol", "Schmidt — momentum vs salt diffusivity"],
        ["Ra", "g' L³/(ν D_S)", "Rayleigh — buoyant convection strength"],
        ["Ro", "U/(f L)", "Rossby — rotation (Coriolis) importance"],
        ["Π*", "Π/(ρ U_d²)", "Osmotic-to-inertial pressure ratio (novel group)"],
        ["Kᵤ", "σ_m τ_m /U", "Stochastic-forcing intensity (novel group)"],
    ],
    col_widths=[0.7, 1.9, 3.6],
)

# ======================================================================
#  9. PARAMETER INVENTORY
# ======================================================================
DOC.add_page_break()
h("9. Complete Parameter Inventory (Model Inputs)", 1)
para("Every quantity the model ingests, grouped by source. Symbols match the "
     "equations above. These are the inputs that must be supplied or measured "
     "to drive NEREID-B.")

h("9.1 Effluent / brine discharge", 2)
make_table(
    ["Symbol", "Parameter", "Units"],
    [
        ["S₀", "Brine (effluent) absolute salinity", "g kg⁻¹"],
        ["Q_d", "Volumetric discharge rate", "m³ s⁻¹"],
        ["U_d", "Nozzle exit velocity", "m s⁻¹"],
        ["p_d", "Discharge pressure at nozzle", "Pa"],
        ["T_b", "Brine temperature", "°C"],
        ["ρ_b", "Brine density", "kg m⁻³"],
        ["θ, ψ", "Nozzle inclination & azimuth angles", "deg"],
        ["S_anti", "Anti-scalant / additive tracer conc.", "mg L⁻¹"],
    ],
    col_widths=[0.8, 3.6, 1.4],
)

h("9.2 Pipe / diffuser geometry & hydraulics", 2)
make_table(
    ["Symbol", "Parameter", "Units"],
    [
        ["L_p", "Pipe length", "m"],
        ["d_p", "Pipe / port internal diameter", "m"],
        ["ε_r", "Pipe wall roughness", "m"],
        ["f", "Darcy friction factor", "–"],
        ["N_port", "Number of diffuser ports", "–"],
        ["s_port", "Port spacing", "m"],
        ["h_n", "Nozzle height above seabed", "m"],
        ["Δz_p", "Pipe vertical rise", "m"],
        ["K_minor", "Minor-loss coefficients", "–"],
    ],
    col_widths=[0.8, 3.6, 1.4],
)

h("9.3 Ambient sea — physical state", 2)
make_table(
    ["Symbol", "Parameter", "Units"],
    [
        ["S_amb(z)", "Ambient salinity profile", "g kg⁻¹"],
        ["T_amb(z)", "Ambient temperature profile", "°C"],
        ["ρ_amb(z)", "Ambient density (stratification)", "kg m⁻³"],
        ["p(z)", "Hydrostatic + dynamic pressure", "Pa"],
        ["N²", "Brunt–Väisälä (stratification) frequency", "s⁻²"],
        ["c_p", "Specific heat capacity of seawater", "J kg⁻¹ K⁻¹"],
        ["μ", "Dynamic viscosity", "Pa s"],
        ["D_mol", "Molecular salt diffusivity", "m² s⁻¹"],
    ],
    col_widths=[0.9, 3.5, 1.4],
)

h("9.4 Ambient sea — dynamics (the moving sea)", 2)
make_table(
    ["Symbol", "Parameter", "Units"],
    [
        ["U_current(z,t)", "Ocean current profile (magnitude & direction)", "m s⁻¹"],
        ["A_j, ω_j, φ_j", "Tidal constituents (amp., freq., phase)", "m, s⁻¹, rad"],
        ["H_s, T_w, ê_w", "Significant wave height, period, direction", "m, s, –"],
        ["u_St", "Stokes drift (wave-induced)", "m s⁻¹"],
        ["ω = ∇×u", "Vorticity / vortex field", "s⁻¹"],
        ["φ_lat", "Latitude (Coriolis parameter f)", "deg"],
        ["τ_m, σ_m, ℓ_m", "Stochastic de-correlation time, intensity, length", "s, –, m"],
    ],
    col_widths=[1.2, 3.3, 1.3],
)

h("9.5 Atmosphere / air–sea interface", 2)
make_table(
    ["Symbol", "Parameter", "Units"],
    [
        ["U₁₀", "Wind velocity at 10 m (air velocity)", "m s⁻¹"],
        ["C_d", "Air–sea drag coefficient", "–"],
        ["ρ_a", "Air density", "kg m⁻³"],
        ["E, P", "Evaporation, precipitation rates", "m s⁻¹"],
        ["Q_rad", "Net radiative heat flux", "W m⁻²"],
        ["T_air", "Air temperature", "°C"],
    ],
    col_widths=[0.9, 3.5, 1.4],
)

h("9.6 Bathymetry, terrain & domain", 2)
make_table(
    ["Symbol", "Parameter", "Units"],
    [
        ["H(x,y)", "Bathymetry / seabed depth", "m"],
        ["∇H", "Bed slope (continental-shelf geometry)", "–"],
        ["C_D^bed", "Bed friction coefficient", "–"],
        ["x_d", "Discharge location", "m"],
        ["L_x,L_y,D", "Domain extent & total depth", "m"],
    ],
    col_widths=[0.9, 3.5, 1.4],
)

h("9.7 Thermodynamic / coupling coefficients", 2)
make_table(
    ["Symbol", "Parameter", "Units"],
    [
        ["α_T, β_S, γ_p", "Thermal-expansion, haline-contraction, compressibility", "K⁻¹, kg g⁻¹, Pa⁻¹"],
        ["φ_os", "Osmotic (activity) coefficient", "–"],
        ["L_p", "Osmotic permeability coefficient", "s"],
        ["D_ST, D_TS", "Soret & Dufour cross-diffusion coefficients", "m² s⁻¹ K⁻¹ (eq.)"],
        ["Sc_t, Pr_t", "Turbulent Schmidt & Prandtl numbers", "–"],
        ["v̄_w, M_s, ν", "Partial molar volume, salt molar mass, ion number", "m³ mol⁻¹, kg mol⁻¹, –"],
    ],
    col_widths=[1.1, 3.4, 1.4],
)

# ======================================================================
#  10. NUMERICAL STRATEGY
# ======================================================================
h("10. Recommended Numerical Solution Strategy", 1)
numbered("Spatial discretization on a terrain-following σ-grid by a finite-"
         "volume method, refined near the diffuser and the bed; unstructured "
         "horizontal mesh to capture the coastline and bathymetry.")
numbered("Pressure–velocity coupling for variable density by a PISO/SIMPLE "
         "projection with the non-Boussinesq continuity constraint.")
numbered("Scalar transport (S, T) with a high-resolution, monotone (TVD) "
         "advection scheme to avoid spurious salinity over/undershoots.")
numbered("Semi-implicit time integration; CFL- and buoyancy-limited adaptive "
         "time step for the unsteady evolution from discharge to far field.")
numbered("Stochastic layer by Monte-Carlo ensemble (or polynomial-chaos / "
         "stochastic-Galerkin surrogate) over OU realizations to build "
         "P[S(x,t)], the mean, variance and exceedance maps.")
numbered("Two-phase free surface by VOF (α) with the wind-stress and "
         "evaporation surface fluxes.")
numbered("Coupling solved by an outer Picard/Newton iteration each step until "
         "ρ, u, S, T, k, ε are mutually converged.")

# ======================================================================
#  11. VALIDATION & WELL-POSEDNESS
# ======================================================================
h("11. Validation Plan and Well-Posedness", 1)
bullet("Analytical limits: recover the classical 30°/45°/60° inclined dense-"
       "jet trajectory and dilution scaling (Roberts et al.) when stochastic, "
       "osmotic and rotation terms are switched off.")
bullet("Laboratory benchmarks: compare rise height, return-point dilution and "
       "spreading-layer thickness against published dense-jet experiments.")
bullet("Field validation: CTD / ADCP transects of an operating SWRO outfall; "
       "verify the predicted P[S] envelope contains the observed salinities at "
       "the stated confidence.")
bullet("Conservation checks: global salt and mass budgets closed to machine "
       "precision; entropy-production positivity of the Onsager fluxes "
       "guarantees thermodynamic consistency and dissipativity (a sufficient "
       "condition for well-posedness of the coupled parabolic-hyperbolic "
       "system).")

# ======================================================================
#  12. ASSUMPTIONS & LIMITATIONS
# ======================================================================
h("12. Assumptions and Limitations", 1)
bullet("Seawater is a Newtonian continuum; the reject carries no significant "
       "suspended-solid rheology (extendable via a particle-laden closure).")
bullet("The OU stochastic forcing is Gaussian and stationary over the "
       "forecast window; non-stationary storm forcing requires re-estimated "
       "(τ_m, σ_m, ℓ_m).")
bullet("Chemical speciation is lumped into a single absolute-salinity scalar; "
       "multi-ion reverse-osmosis selectivity is a possible extension via "
       "Maxwell–Stefan multicomponent fluxes.")
bullet("Soret/Dufour and osmotic coefficients must be supplied for the "
       "specific brine chemistry; defaults are provided for NaCl-dominated "
       "SWRO reject.")

# ======================================================================
#  13. NOMENCLATURE
# ======================================================================
DOC.add_page_break()
h("13. Nomenclature (Principal Symbols)", 1)
make_table(
    ["Symbol", "Meaning", "Units"],
    [
        ["ρ, ρ₀", "Density, reference density", "kg m⁻³"],
        ["u=(u,v,w)", "Velocity vector", "m s⁻¹"],
        ["p", "Pressure", "Pa"],
        ["S", "Absolute salinity", "g kg⁻¹"],
        ["T", "Conservative temperature", "°C / K"],
        ["k, ε", "Turbulent kinetic energy, dissipation", "m² s⁻², m² s⁻³"],
        ["μ, μ_t", "Molecular, turbulent (eddy) viscosity", "Pa s"],
        ["D_S, D_T", "Salt, heat dispersion tensors", "m² s⁻¹"],
        ["Π", "Osmotic pressure", "Pa"],
        ["g, g'", "Gravity, reduced gravity", "m s⁻²"],
        ["Ω, f", "Rotation vector, Coriolis parameter", "s⁻¹"],
        ["η, α", "Free-surface elevation, water fraction", "m, –"],
        ["S_rad", "Wave radiation-stress tensor", "N m⁻¹"],
        ["ζ_m, ξ_S, ξ_T", "Stochastic forcing channels & scalar noise", "various"],
        ["F_osm, F_wave", "Osmotic, wave body forces", "N m⁻³"],
        ["𝒟, 𝒜, r_max, z_max", "Dilution, footprint, reach, depth of impact", "–, m², m, m"],
    ],
    col_widths=[1.3, 3.3, 1.2],
)

# ======================================================================
#  14. CLOSING
# ======================================================================
h("14. Summary", 1)
para(
    "NEREID-B unifies, for the first time in a single self-consistent PDE "
    "system, the non-Boussinesq hydrodynamics of a dense brine jet, the "
    "nonlinear haline-thermal-baric equation of state, irreversible-"
    "thermodynamics (Soret/Dufour) cross-diffusion, an explicit osmotic / "
    "reverse-osmosis coupling, a state-dependent anisotropic dispersion "
    "tensor, terrain-following bathymetry, two-phase splash-zone air-sea "
    "exchange, and an intrinsic stochastic-forcing layer that converts the "
    "model into a predictor of the full probability density of salinity. The "
    "result is a model that predicts, at every instant and with quantified "
    "uncertainty, the three-dimensional salinity field around a submarine "
    "outfall and precisely how far and how deep the discharged brine "
    "penetrates a moving, unpredictable sea — accurately, sustainably, and in "
    "a form suitable for patent protection.",
)

# ----------------------------------------------------------------------
DOC.save("/home/akosa/salinity_prediction/salinity.docx")
print("Saved salinity.docx")
print("Equations numbered:", _eq_counter[0])
