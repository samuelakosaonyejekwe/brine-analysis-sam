# -*- coding: utf-8 -*-
"""
Generator for input.docx — the complete input-data requirements / data
acquisition & measurement plan needed to drive the NEREID-B coupled
stochastic PDE brine-salinity model.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = Document()
BODY_FONT = "Calibri"
EQ_FONT = "Cambria Math"

normal = DOC.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)


def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def h(text, level=1):
    return DOC.add_heading(text, level=level)


def para(text="", bold=False, italic=False, size=11, align=None, color=None,
         space_after=6):
    p = DOC.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = BODY_FONT
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def bullet(text, sub=False):
    style = "List Bullet 2" if sub else "List Bullet"
    p = DOC.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def numbered(text):
    p = DOC.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def make_table(header, rows, col_widths=None, font_size=8.5, mono_first=False):
    t = DOC.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = BODY_FONT
        _set_cell_bg(hdr[i], "1F4E79")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = EQ_FONT if (mono_first and i == 0) else BODY_FONT
    if col_widths:
        for i, w in enumerate(col_widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)
    return t


# ======================================================================
#  TITLE
# ======================================================================
title = DOC.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("NEREID-B — Input Data Requirements")
tr.bold = True
tr.font.size = Pt(23)
tr.font.name = BODY_FONT
tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = DOC.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Data Acquisition, Measurement and Pre-Processing Plan "
                 "Required to Drive the Coupled Stochastic PDE "
                 "Brine-Salinity Model")
sr.italic = True
sr.font.size = Pt(13)
sr.font.name = BODY_FONT

meta = DOC.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Companion to salinity.docx & output.docx  ·  "
                  "Rev. 1.0  ·  11 June 2026")
mr.font.size = Pt(11)
mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
mr.font.name = BODY_FONT

para("", space_after=8)
para(
    "This document specifies every input the NEREID-B model requires, how "
    "each input is obtained (measurement instrument, survey, model or "
    "literature), the spatial and temporal resolution and accuracy needed, "
    "and the update frequency. Inputs are grouped by physical category. Each "
    "table lists the symbol, the quantity, how it is acquired, the required "
    "resolution / accuracy, and its update cadence. A priority tier (Required "
    "/ Recommended / Optional) is given so a project can scope a minimum-"
    "viable dataset and grow it. The document closes with the field-survey / "
    "monitoring programme, data-format and QA-QC requirements, and a "
    "readiness checklist.",
)

# Legend
para("Priority key:  R = Required (model will not run without it)   ·   "
     "S = Strongly recommended (materially affects accuracy)   ·   "
     "O = Optional (refinement / specific scenarios).", italic=True, size=10)

# ======================================================================
#  1. INPUT DATA OVERVIEW
# ======================================================================
h("1. Input Data Classes", 1)
para("NEREID-B ingests seven classes of input:")
bullet("A. Effluent / brine discharge characteristics (the source term).")
bullet("B. Pipe & diffuser geometry and hydraulics (the boundary device).")
bullet("C. Ambient sea state — physical (salinity, temperature, density, pressure).")
bullet("D. Ambient sea dynamics — currents, tides, waves, turbulence (the moving sea).")
bullet("E. Atmospheric / air–sea forcing (wind, heat, evaporation).")
bullet("F. Bathymetry, terrain & domain geometry.")
bullet("G. Thermodynamic, coupling & stochastic coefficients (closure data).")
para("Classes A–B define the source and how it is injected; C–F define the "
     "receiving environment and its forcing; G supplies the physical "
     "constants and the statistics of the unpredictable forcing.")

# ======================================================================
#  CLASS A — DISCHARGE
# ======================================================================
h("2. Class A — Effluent / Brine Discharge (Source Term)", 1)
make_table(
    ["Sym.", "Quantity", "How acquired", "Resolution / accuracy", "Update", "Pri."],
    [
        ["S₀", "Brine salinity / TDS", "In-line conductivity & density meter on reject line; lab TDS", "±0.1 g kg⁻¹; continuous", "1 min–1 h", "R"],
        ["Q_d", "Discharge flow rate", "Magnetic / ultrasonic flowmeter on outfall", "±2 %", "1 min", "R"],
        ["U_d", "Nozzle exit velocity", "Derived: Q_d / port area", "from Q_d & geometry", "1 min", "R"],
        ["p_d", "Discharge pressure", "Pressure transducer at manifold/port", "±1 %", "1 min", "S"],
        ["T_b", "Brine temperature", "In-line RTD/thermistor", "±0.1 °C", "1 min", "R"],
        ["ρ_b", "Brine density", "Density meter or from S₀,T_b via EOS", "±0.5 kg m⁻³", "1 min", "R"],
        ["composition", "Ion composition (Na,Cl,Ca,Mg,SO₄)", "Periodic lab ion chromatography", "—", "weekly", "S"],
        ["S_anti", "Anti-scalant/additive load", "Plant dosing records", "as dosed", "per batch", "O"],
        ["pH, DO", "Reject pH, dissolved oxygen", "In-line probes", "±0.1 pH", "1 h", "O"],
    ],
    col_widths=[0.5, 1.3, 1.9, 1.5, 0.7, 0.4], mono_first=True,
)

# ======================================================================
#  CLASS B — PIPE GEOMETRY
# ======================================================================
h("3. Class B — Pipe & Diffuser Geometry and Hydraulics", 1)
make_table(
    ["Sym.", "Quantity", "How acquired", "Resolution / accuracy", "Update", "Pri."],
    [
        ["L_p", "Pipe length", "Engineering drawings / as-built survey", "±0.1 m", "static", "R"],
        ["d_p", "Pipe & port diameter", "As-built / design spec", "±1 mm", "static", "R"],
        ["ε_r", "Wall roughness", "Material spec; ROV inspection for ageing", "—", "annual", "S"],
        ["N_port, s_port", "Port count & spacing", "Diffuser design", "exact", "static", "R"],
        ["θ, ψ", "Nozzle elevation & azimuth angle", "Design + as-built ROV survey", "±1°", "static", "R"],
        ["h_n", "Nozzle height above bed", "ROV / multibeam survey", "±0.1 m", "annual", "S"],
        ["Δz_p", "Pipe vertical rise", "As-built profile", "±0.1 m", "static", "S"],
        ["K_minor, f", "Minor-loss & friction factors", "Hydraulic design / handbook + calibration", "—", "static", "S"],
        ["x_d", "Diffuser location (geo-position)", "DGPS / survey", "±0.5 m", "static", "R"],
    ],
    col_widths=[0.6, 1.4, 1.8, 1.5, 0.6, 0.4], mono_first=True,
)

# ======================================================================
#  CLASS C — AMBIENT PHYSICAL
# ======================================================================
DOC.add_page_break()
h("4. Class C — Ambient Sea: Physical State", 1)
make_table(
    ["Sym.", "Quantity", "How acquired", "Resolution / accuracy", "Update", "Pri."],
    [
        ["S_amb(z)", "Ambient salinity profile", "CTD casts; moored conductivity sensors", "±0.01 g kg⁻¹; ΔZ≈0.5 m", "seasonal + real-time", "R"],
        ["T_amb(z)", "Ambient temperature profile", "CTD; thermistor chain", "±0.01 °C; ΔZ≈0.5 m", "seasonal + real-time", "R"],
        ["ρ_amb(z)", "Density / stratification", "Derived from S,T,p via TEOS-10", "from CTD", "seasonal", "R"],
        ["p(z)", "Pressure profile", "CTD pressure sensor / hydrostatic", "±0.1 dbar", "static law", "R"],
        ["N²(z)", "Stratification frequency", "Computed from ρ(z)", "from CTD", "seasonal", "S"],
        ["c_p, μ", "Specific heat, viscosity", "Seawater property libs (from S,T,p)", "literature", "static", "S"],
        ["D_mol", "Molecular salt diffusivity", "Literature value (≈1.5×10⁻⁹ m² s⁻¹)", "literature", "static", "S"],
        ["turbidity/SS", "Suspended-solid background", "Optical turbidity sensor", "±5 %", "monthly", "O"],
    ],
    col_widths=[0.7, 1.4, 1.8, 1.4, 0.7, 0.4], mono_first=True,
)
para("Profiles should span at least one full seasonal cycle to capture "
     "summer/winter stratification extremes, plus event-driven casts. A "
     "vertical resolution of ≈0.5 m near the pycnocline is recommended.",
     italic=True, size=10)

# ======================================================================
#  CLASS D — AMBIENT DYNAMICS
# ======================================================================
h("5. Class D — Ambient Sea Dynamics (the Moving Sea)", 1)
make_table(
    ["Sym.", "Quantity", "How acquired", "Resolution / accuracy", "Update", "Pri."],
    [
        ["U_current(z,t)", "Current speed & direction profile", "ADCP (bottom-mounted/moored); HF radar; ocean model", "±1 cm s⁻¹; ΔZ≈1 m", "10 min, ≥30 d", "R"],
        ["A_j,ω_j,φ_j", "Tidal harmonic constituents", "Harmonic analysis of tide-gauge / ADCP record", "≥35 d record", "per analysis", "R"],
        ["H_s,T_w,ê_w", "Wave height, period, direction", "Wave buoy / AWAC; wave model (SWAN)", "±10 %", "30 min", "S"],
        ["u_St", "Stokes drift", "Derived from wave spectrum", "from waves", "30 min", "O"],
        ["spectra", "Directional wave spectra", "Directional wave buoy", "—", "1 h", "O"],
        ["ω", "Background vorticity/eddies", "ADCP array / model", "—", "10 min", "O"],
        ["φ_lat", "Latitude (Coriolis f)", "Site coordinates", "exact", "static", "R"],
        ["τ_m,σ_m,ℓ_m", "Stochastic forcing statistics", "Estimated from ADCP/wind time-series (autocorrelation, variance, length-scale)", "from records", "per calibration", "S"],
    ],
    col_widths=[0.9, 1.4, 1.8, 1.2, 0.7, 0.4], mono_first=True,
)
para("The stochastic-layer parameters (τ_m, σ_m, ℓ_m) are estimated by "
     "fitting Ornstein–Uhlenbeck statistics — de-correlation time, variance "
     "and spatial correlation length — to the observed current, wave and wind "
     "fluctuation records. These convert the measured variability of the sea "
     "into the model's uncertainty forcing.", italic=True, size=10)

# ======================================================================
#  CLASS E — ATMOSPHERE
# ======================================================================
h("6. Class E — Atmospheric / Air–Sea Forcing", 1)
make_table(
    ["Sym.", "Quantity", "How acquired", "Resolution / accuracy", "Update", "Pri."],
    [
        ["U₁₀", "Wind speed & direction at 10 m", "Met station / anemometer; reanalysis (ERA5)", "±0.5 m s⁻¹", "10 min–1 h", "S"],
        ["σ_wind", "Wind gust statistics", "From anemometer record", "—", "per calibration", "O"],
        ["C_d", "Air–sea drag coefficient", "Parameterisation (wind-speed dependent)", "literature", "static", "S"],
        ["ρ_a, T_air", "Air density & temperature", "Met station / reanalysis", "±0.3 °C", "1 h", "S"],
        ["E, P", "Evaporation & precipitation", "Met station; flux parameterisation", "±10 %", "daily", "O"],
        ["Q_rad", "Net radiative heat flux", "Pyranometer / reanalysis", "±10 W m⁻²", "1 h", "O"],
        ["RH", "Relative humidity", "Met station", "±3 %", "1 h", "O"],
    ],
    col_widths=[0.7, 1.4, 1.8, 1.4, 0.7, 0.4], mono_first=True,
)

# ======================================================================
#  CLASS F — BATHYMETRY
# ======================================================================
DOC.add_page_break()
h("7. Class F — Bathymetry, Terrain & Domain", 1)
make_table(
    ["Sym.", "Quantity", "How acquired", "Resolution / accuracy", "Update", "Pri."],
    [
        ["H(x,y)", "Bathymetry / seabed depth", "Multibeam echosounder survey; LiDAR nearshore; ENC charts", "≤5 m grid; ±0.1 m", "project + post-storm", "R"],
        ["∇H", "Bed slope / shelf geometry", "Derived from H gradient", "from DEM", "static", "R"],
        ["seabed type", "Sediment / substrate class", "Grab samples; side-scan sonar", "class map", "project", "S"],
        ["C_D^bed", "Bed friction coefficient", "From substrate class / calibration", "literature", "static", "S"],
        ["coastline", "Shoreline & structures", "Survey / satellite / charts", "≤5 m", "annual", "R"],
        ["L_x,L_y,D", "Domain extent & depth", "Modeller choice from site", "—", "static", "R"],
        ["receptors", "Sensitive-habitat locations", "Ecological survey (seagrass, reef)", "GPS points", "project", "S"],
    ],
    col_widths=[0.8, 1.4, 1.8, 1.3, 0.7, 0.4], mono_first=True,
)

# ======================================================================
#  CLASS G — COEFFICIENTS
# ======================================================================
h("8. Class G — Thermodynamic, Coupling & Closure Coefficients", 1)
make_table(
    ["Sym.", "Quantity", "How acquired", "Resolution / accuracy", "Update", "Pri."],
    [
        ["α_T,β_S,γ_p", "Expansion/contraction/compress. coeffs", "TEOS-10 equation of state (from S,T,p)", "standard", "computed", "R"],
        ["φ_os", "Osmotic (activity) coefficient", "Pitzer model / lab osmometry for the brine", "±2 %", "per brine", "S"],
        ["L_p", "Osmotic permeability coefficient", "Calibration / irreversible-thermo estimate", "calibrated", "per calibration", "O"],
        ["D_ST,D_TS", "Soret & Dufour coefficients", "Literature / lab; small for seawater", "literature", "static", "O"],
        ["Sc_t,Pr_t", "Turbulent Schmidt & Prandtl numbers", "Calibration (≈0.7)", "calibrated", "per calibration", "S"],
        ["C_μ,C₁,C₂,σ", "k–ε turbulence constants", "Standard values; calibration", "standard", "static", "S"],
        ["v̄_w,M_s,ν", "Molar volume, salt molar mass, ion number", "Chemistry of the brine", "literature", "static", "S"],
    ],
    col_widths=[0.9, 1.5, 1.7, 1.2, 0.7, 0.4], mono_first=True,
)

# ======================================================================
#  9. FIELD SURVEY / MONITORING PROGRAMME
# ======================================================================
h("9. Field-Survey and Monitoring Programme", 1)
para("A staged data-acquisition campaign supplies the time-varying inputs and "
     "the calibration/validation data:")
numbered("Baseline (pre-operation) survey: seasonal CTD grid, multibeam "
         "bathymetry, sediment sampling, habitat mapping — establishes ambient "
         "S,T,ρ climatology and the domain.")
numbered("Mooring deployment: bottom-mounted ADCP + thermistor/conductivity "
         "chain + wave gauge for ≥30–45 days per season — supplies currents, "
         "tides, waves, stratification and the stochastic statistics.")
numbered("Meteorological feed: on-site met station and/or ERA5 reanalysis — "
         "wind, heat, evaporation.")
numbered("Source monitoring: continuous in-line instrumentation on the reject "
         "line (flow, conductivity, density, temperature, pressure).")
numbered("Validation survey: CTD transects and towed/AUV salinity mapping "
         "around the operating outfall — the data the model is tested against "
         "(see output.docx, Tier 8b).")
numbered("Event & post-storm surveys: re-survey bathymetry and re-sample "
         "currents after major storms to update H and forcing statistics.")

# ======================================================================
#  10. DATA FORMAT, PRE-PROCESSING & QA/QC
# ======================================================================
h("10. Data Format, Pre-Processing and QA/QC", 1)
bullet("Coordinate reference: all spatial data in a single projected CRS "
       "(e.g. UTM) with a defined vertical datum (chart datum / MSL).")
bullet("Time base: all time series in UTC with consistent sampling; gaps "
       "flagged, not zero-filled.")
bullet("Formats: gridded fields as NetCDF/GeoTIFF; time series as CSV/Parquet; "
       "bathymetry as XYZ/DEM; metadata per CF conventions.")
bullet("Pre-processing: de-spiking, drift correction and calibration of "
       "sensors; harmonic analysis to separate tidal/residual currents; "
       "interpolation of CTD casts to model levels.")
bullet("QA/QC: range and spike checks, sensor cross-comparison, mass-balance "
       "sanity checks, and documented uncertainty for every input (feeds the "
       "stochastic layer and the output confidence envelopes).")
bullet("Gap-filling: short gaps by interpolation; long gaps by validated "
       "ocean/atmospheric model reanalysis, with the substitution flagged.")

# ======================================================================
#  11. MINIMUM-VIABLE vs FULL DATASET
# ======================================================================
DOC.add_page_break()
h("11. Minimum-Viable vs Full Dataset", 1)
para("The model can run at increasing fidelity as data is added:")
make_table(
    ["Tier", "Dataset", "Enables"],
    [
        ["Minimum (R only)", "Discharge S₀/Q_d/T_b, pipe geometry & angle, "
         "ambient S/T profile, bathymetry, currents+tides, latitude",
         "Deterministic unsteady plume; r_max, z_max, dilution, footprint"],
        ["Recommended (R+S)", "Adds waves, wind, stratification time series, "
         "osmotic & turbulence coefficients, stochastic statistics",
         "Full coupled accuracy + uncertainty envelopes & exceedance maps"],
        ["Full (R+S+O)", "Adds ion chemistry, heat fluxes, Soret/Dufour, "
         "directional spectra, event surveys",
         "Maximum-fidelity research-grade prediction & all novel couplings"],
    ],
    col_widths=[1.3, 3.0, 2.4], font_size=9,
)

# ======================================================================
#  12. READINESS CHECKLIST
# ======================================================================
h("12. Input-Readiness Checklist", 1)
para("Before a production run, confirm:")
bullet("☐ Continuous source monitoring (flow, salinity, density, T, p) live.")
bullet("☐ As-built diffuser geometry, port angles and position verified by ROV.")
bullet("☐ Seasonal CTD profiles of ambient S, T, ρ available.")
bullet("☐ ≥30-day ADCP current record + tidal harmonics derived.")
bullet("☐ Wave and wind time series (measured or reanalysis) secured.")
bullet("☐ Multibeam bathymetry / DEM on a consistent CRS & datum.")
bullet("☐ Osmotic, turbulence and stochastic coefficients estimated/calibrated.")
bullet("☐ Validation dataset (outfall CTD survey) reserved for model testing.")
bullet("☐ All inputs QA/QC'd with documented uncertainties.")

para("", space_after=6)
para(
    "Note on inputs easily overlooked but required here: the stochastic-"
    "forcing statistics (τ_m, σ_m, ℓ_m) that quantify the sea's "
    "unpredictability, the osmotic/activity coefficient of the specific "
    "brine, sensor uncertainties (which propagate into the output confidence "
    "envelopes), the vertical datum reconciliation between bathymetry and "
    "tide data, and post-storm bathymetry refresh — each is included above "
    "because NEREID-B's coupled, stochastic formulation depends on them.",
    italic=True,
)

DOC.save("/home/akosa/salinity_prediction/input.docx")
print("Saved input.docx")
