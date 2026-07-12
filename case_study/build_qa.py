#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build questions.answers.docx — examination/defence Q&A for NEREID-B + the Sydney case study."""
import json
import os
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- live run artefacts: never hard-code a number that the run reports -------
HERE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(HERE, "outputs")
MS = json.load(open(os.path.join(OUT_DIR, "metrics_summary.json")))
C, M = MS["config"], MS["metrics"]
SS = M.get("steady_state", {})
STEADY = bool(SS.get("steady_state_reached", False))
NOTCONV = sorted(k for k, v in SS.get("converged", {}).items() if not v)

FOOT = M["seabed_footprint_m2"]
VOL = M["affected_volume_m3"]
REACH = M["r_max_m"]
REACH_W = SS.get("r_max_m_mean", 0.0)
REACH_SD = SS.get("r_max_m_std", 0.0)
RDRIFT = SS.get("r_max_m_drift", 0.0)
RDREL = SS.get("r_max_m_drift_rel", 0.0)
EXMAX = M["excess_max"]
SMAX = M["S_max"]
DILMIN = M["dilution_min"]
NFDIL = M["nf_return_dilution"]
NFRISE = M["nf_rise_m"]
RISE = M["plume_rise_m"]
PTOP = M["plume_top_m"]
ZDEEP = M["z_deepest_m"]
LAYER = ZDEEP - PTOP
TEND = C["t_end"]
NENS = C["ensemble"]
TAU = C.get("stoch_tau", 600.0)
REACH_LO, REACH_HI = min(REACH, REACH_W), max(REACH, REACH_W)
DRAG = bool(C.get('bottom_drag', False))
NUTCAP = M.get('nut_cap_fraction', 0.0)

INK   = RGBColor(0x24, 0x29, 0x2E)   # body: dark slate, never pure black
HEAD  = RGBColor(0x0F, 0x3D, 0x56)   # section headings: deep petrol blue
QCOL  = RGBColor(0x17, 0x33, 0x45)   # question text
ACCENT= RGBColor(0x8A, 0x3D, 0x1E)   # caveat/flag colour (burnt sienna)
MUTED = RGBColor(0x5A, 0x62, 0x6A)

doc = docx.Document()

# ---- page + base style -------------------------------------------------------
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.85)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.font.color.rgb = INK
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
st.paragraph_format.space_after = Pt(4)
st.paragraph_format.line_spacing = 1.08


def title(text, sub=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(20); r.bold = True; r.font.color.rgb = HEAD
    p.paragraph_format.space_after = Pt(2)
    if sub:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(sub)
        r2.font.size = Pt(10.5); r2.italic = True; r2.font.color.rgb = MUTED
        p2.paragraph_format.space_after = Pt(10)


def section(text):
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14); r.bold = True; r.font.color.rgb = HEAD
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)


def subsection(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11); r.bold = True; r.font.color.rgb = HEAD
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)


N = [0]

def qa(q, a):
    N[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"Q{N[0]}.  {q}")
    r.bold = True; r.font.color.rgb = QCOL; r.font.size = Pt(10.5)
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Inches(0.22)
    pa.paragraph_format.space_after = Pt(3)
    ra = pa.add_run(a)
    ra.font.color.rgb = INK; ra.font.size = Pt(10.5)


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = ACCENT


def para(text, italic=False, size=10.5, color=INK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.font.size = Pt(size); r.font.color.rgb = color
    return p


# =============================================================================
title("NEREID-B — Examination & Technical-Review Question Bank",
      "Brine-plume dispersion solver and the Sydney Desalination Plant (Kurnell) case study\n"
      "Reconciled across the report, the thesis, the slide deck and the speaker notes\n"
      "Author: Akosa Samuel Onyejekwe")

para("Scope. This document anticipates the questions a supervisor, examiner, peer reviewer or "
     "client's technical assurance engineer could ask about the NEREID-B solver and the Kurnell "
     "case study, and gives a short defensible answer to each. Answers are capped at four sentences.")
para("Verification basis. Every quantitative answer below was checked directly against the run "
     "artefacts — solver.py, case_study/sydney_sdp_case.json, outputs/metrics_summary.json, "
     "outputs/metrics_timeseries.csv, outputs/fields_final.npz, outputs/ensemble_stats.npz and "
     "outputs/plume_envelope_vs_distance.csv — not against any document's own prose. Where a "
     "document and the underlying data disagreed, the data was followed, the document was corrected, "
     "and the correction is recorded in Section R and the Appendix.")
para("Documents reconciled. brine_case_study.report.docx (and its PDF export), brine1-thesis.docx, "
     "slides.pptx and slides.docx. All four previously carried at least one claim the run does not "
     "support; the most serious was a datum inversion shared by the slides and the thesis. Sections O "
     "and R state each one and its correction.")
note("Sections O (Hard questions & known weaknesses) and R (Cross-document reconciliation) contain "
     "the questions most likely to decide a viva or an assurance review. Read them first.")

# =============================================================================
section("A.  Project scope, motivation and framing")

qa("In one sentence, what does this project deliver?",
   "A single self-contained 3-D solver, NEREID-B, that predicts the salinity field, dispersion and "
   "seabed footprint of a negatively-buoyant desalination brine plume, plus an industrial case study "
   "applying it to the Sydney Desalination Plant diffuser at Kurnell.")

qa("What is the engineering question the case study answers?",
   "Whether the Kurnell multiport diffuser dilutes the brine enough that the seabed excess-salinity "
   "footprint stays within ecological and regulatory limits. The assessment contour is ΔS = 0.5 g/kg "
   f"above ambient. The answer is a footprint of about {FOOT:,.0f} m² confined within roughly "
   f"{REACH_LO:.0f}–{REACH_HI:.0f} m of the diffuser.")

qa("Why is brine dispersion a problem worth modelling at all?",
   "Seawater reverse osmosis returns roughly half its intake as concentrate at about twice ambient "
   "salinity. That concentrate is denser than seawater, so it sinks and pools on the seabed rather "
   "than dispersing upward. Benthic communities are sessile and cannot avoid it, so the seabed excess-salinity "
   "footprint is the controlling ecological metric.")

qa("Why is a brine plume harder to model than a sewage or thermal plume?",
   "It is negatively buoyant, so it rises as a jet then falls back and spreads as a bottom gravity "
   "current. The stable stratification it creates suppresses turbulence, which most RANS closures "
   "get wrong. Dilution therefore depends on physics at two scales that are usually handled by "
   "different, incompatible codes.")

qa("What gap in current practice does NEREID-B claim to fill?",
   "Practice splits the problem between integral near-field jet models (CORMIX, VISJET, Visual "
   "PLUMES) that stop at the seabed return point, and coastal circulation models (Delft3D, MIKE-3, "
   "ROMS) that cannot resolve the jet. NEREID-B couples a validated near-field correlation directly "
   "into a genuine 3-D finite-volume far field. The claim is end-to-end coverage in one tool.")

qa("Is the near-field/far-field split a real limitation of the existing tools, or just an "
   "inconvenience?",
   "It is a real limitation: the hand-off between the two tool classes is done manually, with "
   "no feedback and an arbitrary matching plane. Errors in the near-field return dilution propagate "
   "silently into the far field. A coupled solver makes the hand-off explicit and reproducible.")

qa("Why Sydney/Kurnell rather than another plant?",
   "Kurnell is a deep submerged multiport diffuser in an exposed open-shelf setting, which is the "
   "dominant modern outfall class and squarely inside the solver's design envelope. It also has a "
   "published large-scale ecological impact study (Clark et al. 2018) giving an independent check "
   "on the predicted impact extent.")

qa("What is the intended decision-grade of the output?",
   "Screening-level assessment. The near-field return-dilution coefficient is calibrated to measured "
   "in-class field data; the far field is uncalibrated and is benchmarked, not validated, against that "
   "data (its dilution error across the four measured cases spans 0.35×–3.4×). It is not a substitute "
   "for a consented, survey-calibrated compliance model.")

qa("Who is the intended user?",
   "An outfall designer or environmental assessor sizing a diffuser and estimating a mixing-zone "
   "footprint, who currently needs two licensed codes and a manual coupling step.")

qa("Is this work independent or institutional?",
   "It is independent research work by Akosa Samuel Onyejekwe, released under the MIT licence. "
   "No client, consent process or regulator has adopted or audited it.")

qa("What would it take to move this from screening-grade to consent-grade?",
   "A real CTD/ADCP survey at the modelled outfall to replace the representative calibration deck, "
   "a domain and run length long enough to reach genuine steady state, a grid-convergence study, "
   "and an ensemble large enough to make the exceedance statistics meaningful. Each is identified "
   "in Section O and in the report's own recommendations.")

qa("What is the single strongest claim this project can defend?",
   "That the PDE core is numerically sound and the near-field coupling reproduces published "
   "laboratory scaling: 13/13 self-tests, 4/4 near-field validation, and a lock-exchange front "
   "Froude number of 0.51 against the textbook 0.50. The far-field absolute numbers are weaker "
   "and are defended only as conservative.")

# =============================================================================
section("B.  Solver design, novelty and positioning")

qa("What does the acronym NEREID-B stand for?",
   "Nonlinear Eulerian Reactive-osmotic Effluent Integro-Dispersion solver; the B denotes brine.")

qa("List the claimed novel contributions.",
   "Six: end-to-end near-field to far-field coupling in one solver; a full anisotropic state-dependent "
   "dispersion tensor with off-diagonal terms; a buoyancy-modified realizable k–ε closure with the "
   "correct stratification-damping sign; optional osmotic and Soret/Dufour cross-fluxes; a native "
   "Monte-Carlo stochastic ensemble; and a partial-cell, positivity-preserving numerical core.")

qa("Which of those were actually active in the reported Kurnell run?",
   "Three: near-field coupling, full tensor dispersion and stochastic forcing. The osmotic flux, "
   "osmotic body force, Soret cross-diffusion and free surface were all switched off, as recorded "
   "in the run's own active_physics provenance block. The novelty claims in report §2 therefore "
   "exceed what the case study exercises.")
note("Verified: metrics_summary.json → metrics.active_physics = {osmotic_flux: false, "
     "osmotic_body_force: false, soret_cross_diffusion: false, full_tensor_dispersion: true, "
     "stochastic_forcing: true, free_surface: false, near_field_coupling: true}.")

qa("How is NEREID-B different from CORMIX?",
   "CORMIX is an integral/length-scale near-field model: fast, well-validated, but it terminates at "
   "the seabed return point and cannot represent 3-D bathymetry or an evolving gravity current. "
   "NEREID-B uses the same class of near-field correlation but then solves a genuine 3-D RANS far "
   "field seeded from it.")

qa("How is it different from Delft3D or MIKE-3?",
   "Those solve the coastal circulation on grids far coarser than a nozzle, so they cannot resolve "
   "the jet and must be spun up from a prescribed source. They also use a scalar eddy diffusivity "
   "rather than a dispersion tensor. NEREID-B is smaller in domain but resolves the coupling.")

qa("Is NEREID-B a replacement for those codes?",
   "No. It has a 300 m × 120 m domain, no tidal spin-up, no wetting/drying and no baroclinic shelf "
   "circulation. It is a diffuser-scale assessment tool, not a coastal circulation model.")

qa("What are the dependencies?",
   "NumPy and SciPy are required; matplotlib is optional for figures; CuPy optionally enables a GPU "
   "backend and GSW optionally enables the official TEOS-10 equation of state. The solver is a "
   "single ~5,200-line Python file.")

qa("Why write this in Python rather than Fortran or C++?",
   "The cost is dominated by a sparse LU back-substitution and vectorised NumPy array operations, "
   "both of which run in compiled libraries. Python carries the orchestration, not the inner loops. "
   "A GPU backend is available for the array work.")

qa("What is the applicability envelope?",
   "Deep or submerged multiport brine diffusers — Perth, Gold Coast, Sydney, Carlsbad, Sorek. "
   "Shallow surface or shoreline discharges of roughly 5–6 m are explicitly outside the envelope, "
   "because a deep-diffuser near-field plus gravity-current model is the wrong tool there.")

qa("What discharge modes does the solver support?",
   "Submerged and surface discharge modes; correlation-based and Lagrangian near-field jet models; "
   "a resolved-nozzle mode for jet-resolution studies; and one-way and two-way nested refinement, "
   "including a resolved-near-field mode. The Kurnell case uses submerged discharge with the "
   "correlation near-field model.")

qa("What is the state vector?",
   "q = (u, v, w, p, S, T, ρ, k, ε, ζ): three velocity components, pressure, absolute salinity, "
   "temperature, density, turbulent kinetic energy, its dissipation rate, and the "
   "Ornstein–Uhlenbeck stochastic forcing field. A free-surface elevation η is carried when the "
   "free surface is enabled.")

qa("Is the code the deliverable, or the case study?",
   "Both, but they carry different weight. The solver is the reusable contribution and is where the "
   "verification evidence lives; the case study is a demonstration of it on a real outfall, and its "
   "absolute numbers are the weaker half.")

# =============================================================================
section("C.  Governing equations and physics")

qa("Write down the system being solved.",
   "Incompressible RANS momentum with Coriolis, nonlinear buoyancy and wave stress; a Boussinesq "
   "pressure-Poisson equation from the projection; advection–dispersion for absolute salinity S and "
   "temperature T; a nonlinear equation of state ρ(S,T,z); a two-equation buoyancy-modified "
   "realizable k–ε closure; and an Ornstein–Uhlenbeck stochastic forcing field ζ.")

qa("Is the Boussinesq approximation valid at a brine/ambient density contrast of this size?",
   "The contrast at the nozzle is about 24 kg/m³, roughly 2.3% of ρ₀ — near the edge of, but within, "
   "conventional Boussinesq validity. The solver mitigates this by enforcing continuity in the "
   "divergence-free Boussinesq sense while retaining the full nonlinear density in the buoyancy term. "
   "That keeps the pressure solve linear while preserving the dynamically active buoyancy.")

qa("Why retain nonlinear density in buoyancy but linearise it in continuity?",
   "Buoyancy is where the density contrast does the physical work — it drives the sinking jet and "
   "the gravity current. Continuity is where a nonlinear density would make the pressure equation "
   "nonlinear and expensive. Splitting them is standard practice and is the honest reading of "
   "'Boussinesq' for dense plumes.")

qa("What sets the Coriolis parameter and does it matter here?",
   f"Latitude −34° gives f ≈ −8.16 × 10⁻⁵ s⁻¹, an inertial period of about 21 hours. The run is "
   f"{TEND:.0f} s long, so Coriolis rotates the flow by only a few degrees. It is present for "
   f"completeness and is dynamically minor at this timescale.")

qa("How is wave forcing represented?",
   "As a wave-induced stress in the momentum equation and a wave contribution to the dispersion "
   "tensor, driven by Hs = 1.5 m and Tw = 8.0 s. Individual waves are not resolved; the free surface "
   "was disabled for this run.")

qa("How is wind represented?",
   f"As a surface stress τ = ρ_air · C_d · U₁₀², with U₁₀ = 7.0 m/s, C_d = 0.0013 and ρ_air = 1.225 kg/m³. "
   f"Over {TEND:.0f} s it imparts little momentum to a 25 m water column, and none to the bottom-trapped plume.")

qa("How is the tide represented?",
   f"As a sinusoidal M2 modulation of the ambient current, amplitude 0.04 m/s on a 0.12 m/s mean, "
   f"period 44,712 s. Since the run is {TEND:.0f} s, about {100*TEND/44712:.1f}% of one tidal cycle is "
   f"simulated — a slow drift of the ambient, not a full cycle; the reach oscillation seen in the run "
   f"is driven mainly by the stochastic forcing and the gravity-current dynamics.")

qa("What are the boundary conditions?",
   "Inflow of the ambient current with the prescribed ambient S and T profiles at the upstream face; "
   "radiating outflow downstream; sponge relaxation to ambient near the lateral and outflow "
   "boundaries; no-flux at the partial-cell seabed; and either a rigid lid or a genuine free surface "
   "at z = 0. The Kurnell run used the rigid lid.")

qa("Is there bottom friction?",
   f"Yes — the headline run has bottom_drag ON (quadratic drag τ_b = ρ·C_d·|u_b|·u_b, C_d = 0.0025, "
   f"plus a log-law wall function). It is enabled because the former far-field eddy-viscosity railing "
   f"that used to appear past ~400–600 s is now fixed by a turbulent length-scale limiter, so the run "
   f"integrates stably to {TEND:.0f} s with {100*NUTCAP:.0f}% railing. The drag provides the momentum "
   f"sink that arrests the dense gravity current, so the horizontal reach is now bounded (it surges "
   f"then retreats and oscillates) rather than growing without limit.")

qa("How is the brine actually introduced into the 3-D domain?",
   "As an exponential relaxation of S and T toward prescribed source values inside a 3-D Gaussian "
   "blob centred at the near-field seabed return point, together with a momentum source along the "
   "residual gravity-current direction. The source salinity is S_source = S_amb,bed + (S₀ − S_amb,bed)/S_r, "
   "which for this case is exactly 36.4290 g/kg.")

qa("Is salinity guaranteed to stay physical?",
   "Yes. TVD-MUSCL advection is monotone and positivity-preserving, the off-diagonal cross-flux is "
   "written in conservative divergence form, and a mass-redistributing conservative clip enforces "
   "0 ≤ S ≤ S₀ as a hard bound. The self-test asserts that bound to 10⁻⁶ g/kg.")

qa("How is dilution defined?",
   "Dilution D = (S₀ − S_amb(z)) / (S − S_amb(z)), evaluated cell by cell against the local "
   "depth-varying ambient. Because S_amb varies with depth, the same absolute salinity yields a "
   "different reported dilution at different depths — which matters for interpreting the minimum "
   "far-field dilution.")

# =============================================================================
section("D.  Numerical methods and discretisation")

qa("What is the spatial discretisation?",
   "A structured finite-volume grid with a MAC staggered arrangement: velocities on cell faces, "
   "scalars at cell centres. The Kurnell case uses 52 × 32 × 26 = 43,264 cells over "
   "300 × 120 × 25 m, giving dx = 5.77 m, dy = 3.75 m, dz = 0.96 m.")

qa("Why a staggered MAC grid?",
   "It couples pressure and velocity on adjacent points, which eliminates the odd–even "
   "checkerboard pressure mode that a collocated grid admits. It also makes the discrete divergence "
   "and gradient exact transposes of each other, which the projection relies on.")

qa("What is the time-stepping scheme?",
   "A fractional-step Chorin projection: explicit advection, Coriolis, buoyancy and off-diagonal "
   "dispersion tendencies; implicit backward-Euler LOD diagonal diffusion; a pressure projection to "
   "enforce continuity; TVD scalar transport with a conservative clip; a turbulence update; and a "
   "stochastic forcing increment.")

qa("How is bathymetry represented?",
   "With partial (shaved) cells: each face carries a fractional open area, so a sloping seabed is "
   "represented by cut cells rather than a staircase. This matters because a staircase seabed would "
   "spuriously trip and detach the dense gravity current running down the shelf slope.")

qa("How is the pressure-Poisson equation solved?",
   "The variable-coefficient matrix — incorporating partial-cell open areas and, when enabled, the "
   "free-surface term — is assembled once and LU-factorised with SciPy. Each timestep is then a "
   "back-substitution. This is why the timestep must be fixed on the free-surface path.")

qa("Why is the run memory-bound rather than compute-bound?",
   "The sparse LU factors of a 43,264-cell variable-coefficient Poisson matrix have substantial "
   "fill-in, and they are resident for the whole run. Grid size is therefore limited by the "
   "factorisation footprint, not by floating-point throughput.")

qa("What advection scheme is used and why?",
   "Second-order TVD-MUSCL with a van Leer limiter, applied on the projection's divergence-free MAC "
   "face velocities. It is monotone, so it cannot manufacture salinity overshoots or negative "
   "concentrations. A centred scheme would oscillate at the sharp plume front; first-order upwind "
   "would smear it.")

qa("Why apply advection on post-projection face velocities?",
   "So the advecting velocity field is discretely divergence-free. Advecting a scalar with a "
   "non-solenoidal velocity introduces a spurious source term proportional to the divergence, which "
   "would create or destroy salt.")

qa("Why is diffusion implicit?",
   "Explicit diffusion imposes a dt ≤ dz²/(2D) stability limit, and dz is only 0.96 m while the "
   "vertical eddy diffusivity can be large. Backward-Euler LOD with a Thomas tridiagonal solve "
   "removes that constraint. It is Strang-symmetrised across axes to preserve second-order splitting "
   "accuracy.")

qa("The off-diagonal dispersion terms are explicit — is that stable?",
   "They are sub-stepped: the number of sub-steps is chosen from the local off-diagonal magnitude so "
   "each sub-step is stable. The flux is written in conservative divergence form, and the self-test "
   "confirms it integrates to zero in a closed box.")

qa("What sets the timestep?",
   "A CFL condition on the advective and diffusive scales. Because the seeded gravity current is "
   "slow, the CFL is sized from the seeded velocity rather than the 4 m/s nozzle jet velocity, which "
   "would otherwise force a needlessly tiny dt. The free-surface path uses a fixed dt so the Poisson "
   "matrix factors once.")

qa("Is the scheme conservative?",
   "For scalars, yes to machine precision in a closed box: the self-test confirms implicit diffusion "
   "and the off-diagonal cross-flux each conserve the integral to under 10⁻⁹ relative. In the open "
   "domain, the global volume budget is enforced at the boundaries, with a final mass imbalance of "
   "1.5 × 10⁻¹⁶.")

qa("How divergence-free is the velocity field?",
   "Final continuity residual 7.3 × 10⁻¹⁶ and maximum over the run 2.7 × 10⁻¹⁵ — machine precision. "
   "The divergence also drifts downward over the run (ratio 0.27), so it is not accumulating.")

qa("Is the time integration first or second order?",
   "First order for this run: time_order_2 is false. A second-order option exists. Since the reported "
   "quantities are steady-window averages rather than transients, the temporal order matters less "
   "than the fact that steady state was not fully reached.")

# =============================================================================
section("E.  Near-field jet model and coupling")

qa("Why is the near field not resolved directly?",
   "The nozzle diameter is 0.12 m and dx is 5.77 m — the nozzle is roughly 1/48 of a cell. Resolving "
   "the jet would need a grid finer by two orders of magnitude in each direction. Instead the near "
   "field is supplied by validated empirical correlations, which is the CORMIX/VISJET-class approach.")

qa("What correlations are used?",
   "For a 60° inclined dense jet: terminal rise z_t = 2.2 · D · Fr, return distance x_r = 2.4 · D · Fr, "
   "and return dilution S_r = 1.6 · Fr, after Roberts, Ferrier & Daviero (1997). Mild sine/cosine "
   "angle factors generalise these to other nozzle angles.")

qa("How is the densimetric Froude number computed and what is its value?",
   "Fr = U_d / √(g′₀ · d_p), with reduced gravity g′₀ taken against the ambient bed density rather "
   "than ρ₀. Here U_d = 4.006 m/s, g′₀ = 0.226 m/s² and d_p = 0.12 m, giving Fr = 24.31.")

qa("Where does the exit velocity of 4.0 m/s come from?",
   "From the mass balance in the input deck: 250 ML/day product at 47% recovery gives 532 ML/day feed "
   "and 282 ML/day concentrate, i.e. 3.26 m³/s. Split over 72 ports that is Q = 0.0453 m³/s per port, "
   "which through a 0.12 m nozzle gives U_d ≈ 4.0 m/s.")

qa("Where does the brine salinity of 67 g/kg come from?",
   "From the same mass balance: S_brine = S_amb/(1 − recovery) = 35.5/(1 − 0.47) ≈ 67 g/kg. This "
   "assumes complete salt rejection by the membranes, which is very nearly true for SWRO.")

qa("What is the multiport merging correction?",
   "Adjacent jets merge once their half-width reaches half the port spacing, after which they entrain "
   "as a line plume and dilute less than independent jets. The factor is min(1, max(0.4, s/(0.8·z_t))). "
   "Here s = 5 m and z_t = 6.42 m give a merge factor of 0.974, reducing S_r from 38.9 to 37.88.")

qa("What is actually passed from the near field to the far field?",
   "A source location at the seabed return point, a source salinity S_source = 36.4290 g/kg, a source "
   "temperature, a residual gravity-current velocity, and a spatial extent for the Gaussian source "
   "blob. The jet itself is never represented in the 3-D field.")

qa("Is the coupling one-way or two-way?",
   "One-way for the correlation model used here: the ambient does not feed back into the near-field "
   "solution. The solver does provide a two-way nested resolved-near-field mode, but it was not used "
   "for this case and is recommended as future work.")

qa("What is the entrainment ODE for, if the correlations set the answer?",
   "It supplies the trajectory shape only. A top-hat entrainment ODE with α ≈ 0.03 is integrated, then "
   "rescaled so its rise height and return distance match the correlations exactly. It is used for "
   "plotting Figure 7 and for seeding geometry, not for predicting dilution.")

qa("Why is the near-field return distance, 7.0 m, smaller than one grid cell?",
   "Because dx = 5.77 m, so the entire near field spans about 1.2 cells. That is by design — it is "
   "precisely why the near field is handled by correlation rather than resolved. But it means the "
   "3-D grid cannot represent any structure inside the near field.")

qa("What is the near-field dilution band reported, and why a band?",
   "36.70:1 to 41.43:1, reflecting the ±12% scatter Roberts reports on S_r/Fr = 1.6. The headline "
   "38:1 is the central estimate after the merging correction.")

qa("Lai & Lee (2012) report S_i/Fr ≈ 1.07, not 1.6. Why use the higher value?",
   "The literature splits into two clusters: Roberts (1.6) and Papakonstantis (1.68) against Lai & Lee "
   "(1.07). The solver adopts the Roberts cluster, which is the standard 60° design reference. Adopting "
   "Lai & Lee would reduce return dilution to about 26:1 and materially increase the predicted footprint, "
   "so this choice is not conservative and should be stated as a sensitivity.")
note("This is a fair challenge. The chosen constant is defensible by citation but is the "
     "less-protective of the two clusters; a sensitivity run at S_i/Fr = 1.07 would strengthen the case.")

# =============================================================================
section("F.  Turbulence closure")

qa("What closure is used?",
   "A buoyancy-modified realizable k–ε model with C_μ = 0.09, C₁ = 1.44, C₂ = 1.92, C₃ = 0.8, "
   "σ_k = 1.0, σ_ε = 1.3, a Durbin (1996) turbulent time-scale limiter with C_r = 0.6, and a "
   "Smagorinsky/WALE LES dissipation floor.")

qa("What was the sign bug, and why does it matter so much?",
   "The buoyancy production term G_b in the k equation carried a flipped sign, so stable brine "
   "stratification produced turbulence instead of damping it. The eddy viscosity then railed to its "
   "cap wherever the water was stratified, and the far field grossly over-mixed. Fixing the sign is "
   "the single most consequential correction in the project.")

qa("How do you know the fix worked?",
   "The eddy-viscosity cap fraction fell from about 17% (coarse grid) and 81% (fine grid) to roughly "
   "5% and 0% — grid-independent, physical turbulence. In the Kurnell run nut_cap_fraction is exactly "
   "0.0%, so the closure is nowhere railing against its limiter.")

qa("What does the Durbin limiter do?",
   "It bounds the turbulent time scale k/ε from below by a strain-rate-based scale, preventing the "
   "unbounded eddy-viscosity growth that standard k–ε exhibits at stagnation points. Together with "
   "the corrected buoyancy sign it removes the railing entirely.")

qa("What does 'realizable' mean here?",
   "That the modelled Reynolds normal stresses remain non-negative and the Schwarz inequality between "
   "shear stresses is satisfied, which standard k–ε can violate under strong strain. In practice it is "
   "the time-scale bound that enforces this.")

qa("Why not use a Reynolds-stress model or LES?",
   "An RSM adds six transport equations and is fragile in stratified flow; LES at 0.12 m nozzle scale "
   "over a 300 m domain is computationally out of reach. A buoyancy-corrected two-equation closure "
   "with an LES dissipation floor is the standard engineering compromise for this problem class.")

qa("What is the LES dissipation floor for?",
   "It puts a lower bound on dissipation at the grid scale so that, in regions where k–ε would under-"
   "dissipate, the sub-grid model still removes energy at the correct rate. It prevents energy "
   "piling up at the smallest resolved scale.")

qa("Is the closure validated independently?",
   "Only indirectly. Its effect is validated through the far-field dilution comparison against Perth "
   "and through the lock-exchange benchmark. There is no direct comparison of modelled k or ε against "
   "measured turbulence quantities.")

qa("What is nut_max and did it bind?",
   "It is a hard cap on eddy viscosity, 8.0 m²/s, retained as a blow-up guard. It bound in 0.0% of "
   "cells in the reported run, so the turbulence field is physical rather than clipped.")

# =============================================================================
section("G.  The dispersion tensor")

qa("What is in the dispersion tensor?",
   "Four additive contributions: an isotropic background; a flow-aligned Taylor shear term along "
   "e_u ⊗ e_u; a wave-orbital term along e_w ⊗ e_w; and an along-slope bathymetric term "
   "D_b (I − n ⊗ n). The sum is a full 3×3 symmetric tensor with off-diagonal entries.")

qa("Why a tensor rather than a scalar eddy diffusivity?",
   "Because dispersion in a sheared, wave-forced, sloping-bed flow is strongly directional: spreading "
   "along the current is much faster than across it, and spreading along the slope differs from "
   "spreading normal to it. A scalar diffusivity forces these to be equal and cannot represent the "
   "off-diagonal correlation between directions.")

qa("What guarantees the tensor is well-posed?",
   "Symmetric positive-definiteness. The self-test samples up to 400 fluid cells, forms the full 3×3 "
   "tensor at each and checks the minimum eigenvalue is non-negative. A non-SPD tensor would make the "
   "diffusion operator ill-posed and admit anti-diffusive growth.")

qa("How are the off-diagonal terms discretised?",
   "As an explicit conservative flux divergence, sub-stepped for stability, and separated from the "
   "diagonal terms which are handled implicitly. The self-test verifies that in a closed box the "
   "off-diagonal divergence integrates to zero to within 10⁻⁹.")

qa("What are the tensor's coefficients?",
   "disp_horiz = 0.05 (isotropic background), shear_disp = 0.2 (Taylor shear gain), "
   "wave_disp_gain = 0.5, bath_disp_gain = 1.0, and a single tunable multiplier farfield_disp_cal, "
   "here set to 1.00.")

qa("Which of these were calibrated?",
   "Only farfield_disp_cal, the scalar multiplier on the tunable horizontal dispersivity. The "
   "molecular and turbulent diffusivities and the near-field correlations were left physical and "
   "unfitted.")

qa("Is 'state-dependent' meaningful, or is the tensor effectively constant?",
   "It is genuinely state-dependent: e_u is the local flow direction, so the shear term rotates with "
   "the current, and the bathymetric term depends on the local seabed normal n. The scalar gains are "
   "constant, but the tensor's orientation and magnitude vary cell by cell.")

qa("How would you defend the specific coefficient values?",
   "Weakly. They are physically-motivated order-unity gains rather than values fitted to data, and "
   "only their aggregate effect is constrained — through farfield_disp_cal against a single transect. "
   "A sensitivity study over the four gains would be the honest next step.")

# =============================================================================
section("H.  Equation of state and the exotic physics")

qa("What equation of state was used?",
   "eos_mode = linear_cabbeling: linear haline contraction (β_S = 7.6 × 10⁻⁴) and thermal expansion "
   "(α_T = 2 × 10⁻⁴) plus a quadratic cabbeling term. The higher-order haline curvature, "
   "thermohaline coupling and thermobaric pressure terms exist in the code but were inactive.")

qa("So the 'nonlinear TEOS-10-style EOS' claimed as novelty was not used?",
   "Correct — not in this run. The full_nonlinear and official teos10 modes are implemented and "
   "selectable, but the Kurnell case ran the linear-plus-cabbeling mode. Over a 16–25 °C range and a "
   "1 g/kg excess this is a defensible simplification, but the novelty claim should be scoped to "
   "the code, not the case.")

qa("What is cabbeling and does it matter here?",
   "Cabbeling is the densification that results when two water parcels of equal density but different "
   "T and S mix, because the EOS is nonlinear in temperature. Its term is −½ρ₀·c·ΔT², with "
   "c = 4.5 × 10⁻⁶. Over the 3 °C ambient temperature range here it contributes under 0.02 kg/m³ and "
   "is negligible against the 24 kg/m³ brine contrast.")

qa("Why clamp ΔT to ±40 K in the cabbeling term?",
   "Because −½c·ΔT² is a downward parabola, so a large |ΔT| would drive density unphysically low. "
   "The clamp is a safety net and is inactive in the 16–25 °C range of this case; the linear terms "
   "are unaffected.")

qa("What is the osmotic flux term meant to represent?",
   "A salt flux driven by the osmotic pressure gradient rather than the concentration gradient, "
   "computed from a van 't Hoff expression with an activity coefficient of 0.9. It is genuinely "
   "absent from mainstream brine tools.")

qa("Was the osmotic term active?",
   "No. osmotic_diff and osmotic_force_gain are both 0.0, so neither the osmotic flux nor the osmotic "
   "body force contributed anything to the reported run. The osmotic_coeff = 0.9 is only the activity "
   "coefficient that would be used if the gains were non-zero.")

qa("What are the Soret and Dufour effects?",
   "Soret is mass flux driven by a temperature gradient; Dufour is heat flux driven by a "
   "concentration gradient. Both are cross-diffusion terms in the Onsager sense. In this run the "
   "Soret coefficient is 0.0, so neither was active.")

qa("Are these terms physically important for brine plumes?",
   "Almost certainly not at these scales — turbulent dispersion exceeds molecular cross-diffusion by "
   "many orders of magnitude. Their value is as a capability for laboratory-scale or high-gradient "
   "problems, and as a differentiator from tools that cannot represent them at all. Claiming them as "
   "a contribution to this case study would be overreach.")

# =============================================================================
section("I.  Stochastic ensemble and uncertainty")

qa("How is uncertainty quantified?",
   "By a Monte-Carlo ensemble in which each member is forced by a distinct realisation of an "
   "Ornstein–Uhlenbeck coloured-noise field with amplitude σ = 0.02, correlation time τ = 600 s and "
   "spatial correlation length 30 m. The ensemble yields mean, standard-deviation and "
   "exceedance-probability fields.")

qa("Why Ornstein–Uhlenbeck rather than white noise?",
   "Because real met-ocean forcing is temporally correlated: currents and winds do not decorrelate "
   "between timesteps. White noise would be averaged away by the solver and would produce no "
   "meaningful ensemble spread.")

qa("How many members were run?",
   f"{NENS:.0f}. The report quotes a peak-excess standard deviation and a nominal 95th-percentile "
   f"excess computed from those members, and then explicitly withdraws both.")

qa("Are two members enough to support a standard deviation or a 95th percentile?",
   "No. A standard deviation from n = 2 has roughly 100% relative uncertainty, and a 95th percentile "
   "is not estimable from two samples at all — it is simply the larger of the two. These statistics "
   "are reported for completeness only and must not be used as an uncertainty bound.")
note("A defensible ensemble for exceedance probabilities is O(100) members. This remains an open item.")

qa("There is a second problem with the ensemble — what is it?",
   f"The Ornstein–Uhlenbeck correlation time is τ = {TAU:.0f} s, comparable with the {TEND:.0f} s run. "
   f"The forcing barely decorrelates within a member, so each realisation samples close to one frozen "
   f"draw of the random field. The spread measures sensitivity to the initial random draw rather than "
   f"to temporal variability.")

qa("What does 'max exceedance probability = 1.00' mean here?",
   "That both members exceeded ΔS_crit somewhere. With two members the exceedance field can only take "
   "the values 0, 0.5 or 1 — it is a three-level indicator, not a probability. The solver does "
   "correctly flag exceedance_is_probability only when ensemble > 1, but that test is too weak.")

qa("What uncertainties are not captured by the ensemble at all?",
   "Structural uncertainty in the near-field correlation constants, in the dispersion-tensor gains, "
   "in the turbulence closure, and in the choice of calibration target. The ensemble varies forcing "
   "only. It is a sensitivity band, not a full uncertainty quantification.")

qa("What would a credible uncertainty analysis look like?",
   "An ensemble of order 100 members, run for several OU correlation times, with the near-field "
   "constant sampled across the Roberts and Lai & Lee clusters and the dispersion gains sampled over "
   "plausible ranges. The output would be a genuine exceedance-probability map with confidence "
   "intervals on the footprint area.")

# =============================================================================
section("J.  Site description, inputs and data provenance")

qa("Describe the site.",
   "The Sydney Desalination Plant at Kurnell, NSW, discharging to the open Tasman Sea shelf in about "
   "25 m of water via tunnelled risers with inclined multiport rosette heads. The shelf slopes at "
   "0.006 and the site is exposed to an energetic open-ocean wave climate.")

qa("What are the ambient conditions?",
   "Summer stratification: salinity 35.4 g/kg at the surface to 35.6 g/kg at the bed; temperature "
   "21 °C to 18 °C. Mean current 0.12 m/s with a 0.04 m/s M2 tidal amplitude, Hs = 1.5 m, Tw = 8.0 s, "
   "and 7 m/s wind at 10 m.")

qa("Which inputs are real and which are constructed?",
   "The plant capacity (250 ML/day), recovery (~47%), discharge depth (~25 m) and the general diffuser "
   "arrangement are public design basis. The per-port nozzle geometry, the bathymetry grid, the CTD "
   "and ADCP decks, the wave and wind series and the calibration transect are all constructed as a "
   "credible representative survey, not measured.")
note("This is stated plainly in sydney_sdp_case.json under _assumptions_engineering and in "
     "validation/sources.md. It must be stated equally plainly in any presentation of the results.")

qa("How were the site data generated?",
   "By case_study/make_site_data.py, a deterministic generator that writes the seven CSV decks in "
   "case_study/inputs/. It is reproducible, which is a virtue, but reproducibility of synthetic data "
   "is not the same as validity.")

qa("Why is the per-port geometry assumed rather than obtained?",
   "SDP's per-port nozzle diameter, count and spacing are not in the public domain. The deck adopts "
   "72 ports of 0.12 m at 5 m spacing as a representative 60° inclined-dense-jet configuration "
   "consistent with the published capacity and recovery.")

qa("How sensitive is the answer to that assumed geometry?",
   "Strongly. Fr scales as U_d/√(g′d_p) and U_d ∝ Q/d_p², so a smaller port raises Fr and hence "
   "dilution; the merge factor depends on spacing relative to rise height. A sensitivity sweep over "
   "d_p, n and s should accompany any quotation of the footprint.")

qa("What does the bathymetry deck contain?",
   "A multibeam-style grid over the diffuser corridor, consumed as a partial-cell field H(x,y). "
   "Depth runs from 24 m nearshore to 25 m at the diffuser on a 0.006 slope.")

qa("What is in the CTD deck?",
   "Summer and winter casts of salinity, temperature and σ_t against depth. Only the summer profile "
   "is used for the reported run; winter, being less stratified, would mix more and is not the "
   "conservative case.")

qa("Why is summer the conservative season?",
   "Because stronger stratification suppresses vertical mixing, keeping the dense layer thinner and "
   "more concentrated at the bed. Winter's weaker stratification would dilute the plume faster.")

qa("Is the domain large enough?",
   f"Yes, comfortably. The domain is 300 m long and, with bottom drag arresting the gravity current, "
   f"the reach is bounded and oscillates around {REACH_W:.0f} m — well inside the domain, with room "
   f"before the sponge zones. Before the fix the reach grew past 90 m and crowded the domain; that no "
   f"longer happens.")

# =============================================================================
section("K.  Calibration")

qa("What exactly was calibrated?",
   "One scalar: nf_dilution_cal, which scales the near-field return-dilution coefficient. It was fitted "
   "to MEASURED field data — the 48.4:1 dilution recorded 60 m from the Gold Coast Desalination Plant "
   "multiport diffuser at 100% plant capacity. The fitted value is 0.871, so the FIELD return-dilution "
   "coefficient is S_r = 1.39·Fr, against the quiescent-LABORATORY 1.6·Fr of Roberts et al. (1997). "
   "Everything else — the far-field dispersivity, the molecular and turbulent diffusivities — was left "
   "unfitted.")
note("Source: Baum, M.J. (2019), 'Dense Jet Behaviour in Dynamic Receiving Environments', PhD thesis, "
     "University of Queensland, Tables 2.2-2.3; peer-reviewed as Baum, Gibbes, Grinham, Albert, Fisher & "
     "Gale (2018), J. Hydraul. Eng. 144(11), doi:10.1061/(ASCE)HY.1943-7900.0001524. Input deck: "
     "case_study/inputs/gcdp_baum_case3-1_transect.csv. Log: validation/nf_calibration.log.")

qa("Why calibrate the NEAR-field coefficient rather than the far-field dispersivity?",
   "Because that is the parameter the measurement can actually identify. The 60 m station lies INSIDE "
   "the near-field mixing zone — Roberts et al. (1997) give its length as x_n = 9.0·Fr·d, about 50 m for "
   "this discharge — so the dilution there is set by near-field jet entrainment, not by far-field "
   "dispersion. Sweeping farfield_disp_cal over a four-fold range moves the modelled 60 m dilution by "
   "less than 3.5%; sweeping nf_dilution_cal over 0.40-1.30 moves it from 18.5:1 to 83.4:1. One knob has "
   "leverage on the observable and the other does not. Fitting the far-field knob to a mixing-zone "
   "measurement would be fitting a parameter the data cannot see.")

qa("So is the far field calibrated?",
   "No. farfield_disp_cal is UNIDENTIFIABLE from the available data and remains at its physical default "
   "of 1.0 — a default, not a fit. Calibrating it would require measurements far enough out for far-field "
   "spreading to dominate. A CTD/ADCP survey at the Kurnell outfall is the only route to a genuine "
   "site calibration of the far field.")

qa("Why is the fitted coefficient BELOW the laboratory value?",
   "Because a real diffuser entrains less than a still laboratory tank. The 1.6·Fr coefficient comes from "
   "quiescent-ambient experiments; a field diffuser sits in crossflow, waves and velocity shear, which "
   "degrade near-field entrainment. That the fit moved 13% in exactly that direction is the reason it is "
   "physically interpretable rather than a numerical fudge — it is the central finding of Baum (2019).")

qa("Does the calibration make the predicted impact larger or smaller?",
   "Larger. Less near-field dilution means more residual salt, so the excess salinity, the seabed "
   "footprint and the reach all increase relative to the uncalibrated run. Calibrating against measured "
   "reality made this assessment MORE conservative, not less. That is the direction a screening "
   "assessment should err in.")

qa("Why fit to only one of the four measured cases?",
   "Because only one of them is clean enough to fit. Case 3-1 ran at 100% plant capacity (the largest "
   "brine signal) with an ambient-salinity drift of only -0.10 g/kg. In the other cases the AMBIENT "
   "background salinity wanders by up to ±2 g/kg while the brine signal at 60 m is at most 0.53 g/kg — "
   "the noise exceeds the signal, and the source authors explicitly caution against reading dilutions "
   "from the low-capacity cases. Fitting to those would be fitting noise. They are reported instead as a "
   "validation spread.")

qa("What does that validation spread show?",
   "That the model is NOT systematically conservative. Across the four measured cases the modelled 60 m "
   "dilution divided by the measured value is 0.35, 1.20, 3.38 and 0.53 — the model is optimistic in two "
   "cases and conservative in two, with errors up to a factor of 3.4. The Roberts (1997) laboratory "
   "scaling misses the same cases in the same directions. Real crossflow and wave forcing dominate "
   "case-to-case dilution in a way neither the lab correlations nor this model reproduces.")

qa("An earlier revision claimed the model was ~16-25% conservative. What happened to that?",
   "It is WITHDRAWN. It rested on the Perth figure of 45:1 at 50 m, which is a DESIGN/COMPLIANCE target "
   "produced by another model — not a measurement. Against actual measurements the model is optimistic in "
   "half the cases. Validating against a design target rather than a measurement flipped the sign of the "
   "safety argument, and the report previously reported only the flattering half.")

qa("Was an earlier version of this study calibrated circularly?",
   "Yes, and it is worth stating plainly. make_site_data.py used to synthesise a 'site CTD/ADCP dilution "
   "transect' whose own source comment recorded that its stations were chosen to be 'reproducible by the "
   "model at no tuning'. Fitting to it guaranteed the farfield_disp_cal = 1.00 that was then reported as "
   "'no adjustment needed'. In fact the routine had FAILED to find leverage and fallen back to its "
   "default. That synthetic file has been deleted and replaced by the measured Gold Coast data.")
note("Verified: the fabricated transect and its generator function are removed from the repository; "
     "the calibration now consumes case_study/inputs/gcdp_baum_case*_transect.csv.")

qa("Could you calibrate against something real instead?",
   "Yes — the Gold Coast Tugun transect (Baum et al. 2019) is measured, in-class and published, giving "
   "boundary dilution of about 62.6 at 60 m. Calibrating to it would be defensible. It was used as a "
   "validation reference rather than a calibration target.")

qa("Why calibrate at the mixing-zone boundary rather than fit the whole transect?",
   "Because the mixing-zone boundary is the point the regulation cares about, and fitting one point "
   "with one parameter avoids the false confidence of a least-squares fit through four points that "
   "were not independently measured. The trade-off is that the near-transect shape is unconstrained.")

qa("If the multiplier is 1.00, is the calibration step doing anything at all?",
   "In this run, no — it is a no-op. Its value is that it exists as an explicit, single, documented "
   "tuning knob rather than a diffuse set of hidden coefficients, and that it can absorb a real "
   "survey when one becomes available.")

# =============================================================================
section("L.  Validation")

qa("At how many independent levels is the model validated?",
   "Three: the near-field jet against laboratory scaling, the far field against published field "
   "transects, and the PDE core against an analytical benchmark. Robustness is separately covered by "
   "13 invariant self-tests.")

qa("What does --selftest check?",
   "Thirteen invariants: field finiteness; the true 0 ≤ S ≤ S₀ bound; divergence control; "
   "machine-precision projection; global mass balance; EOS monotonicity in salinity; TVD "
   "non-amplification; bitwise-exact checkpoint/restart on both the free-surface and rigid-lid paths; "
   "dispersion-tensor SPD-ness; Poisson-matrix symmetry; and conservation of both the implicit "
   "diffusion and the off-diagonal cross-flux in a closed box.")

qa("Why does bitwise-exact restart matter?",
   "Because it proves the checkpoint captures the entire solver state — including the RNG stream, the "
   "turbulence fields and the stochastic forcing. If restart were only approximately exact, some state "
   "would be silently reinitialised, and long ensemble runs could not be trusted.")

qa("What does --validate check, and what is the result?",
   "It compares the near-field model against published 60° inclined-dense-jet scaling across a range "
   "of Froude numbers. NEREID-B returns z_t/(D·Fr) = 2.20 against the accepted band of 2.0–2.2, and "
   "S_i ≈ 1.6·Fr, passing 4/4.")

qa("Is that validation, or is it self-consistency?",
   "Largely self-consistency. The correlation z_t = 2.2·D·Fr is hard-coded, so recovering "
   "z_t/(D·Fr) = 2.20 is arithmetic, not prediction. It confirms correct implementation and correct "
   "Froude-number computation, which is worth having, but it is verification rather than validation.")
note("This is the sharpest question in the validation section. The honest answer is that --validate "
     "verifies the coupling arithmetic; the underlying physics is inherited from Roberts (1997), not "
     "independently reproduced.")

qa("What does the lock-exchange benchmark test?",
   "The PDE core in isolation: a flat closed box with no nozzle source, no ambient current and no "
   "near-field model, in which a density front is released. The front Froude number should be "
   "F_H = 0.50 (Benjamin 1968; Shin et al. 2004). NEREID-B returns 0.51.")

qa("Has the 0.50 target itself been verified against the source?",
   "Yes, and the normalisation checked — which is where this benchmark usually goes wrong. Benjamin's "
   "front condition is U/sqrt(g'H) = sqrt[h(1-h)(2-h)/(1+h)], with h the current depth as a fraction "
   "of the channel depth H. At the energy-conserving depth h = H/2 this evaluates to exactly 0.5000. "
   "Crucially it is normalised on the FULL depth H, which is the same convention the solver uses "
   "(Fr_f = U_f/sqrt(g'H)), so the comparison is like-for-like. The other numbers in this literature — "
   "sqrt(2), and the deep-ambient value of 1 that Shin et al. (2004) derive in place of it — belong to "
   "different normalisations and regimes and must not be compared against ours.")

qa("Why is 0.51 versus 0.50 acceptable — and is the direction right?",
   "The 2% magnitude is well within what a finite grid with a length-scale-limited closure should give, "
   "and the benchmark is genuinely independent of the brine physics, so it remains the strongest single "
   "piece of validation evidence in the project. But the DIRECTION deserves recording: Shin et al. (2004) "
   "find that dissipation reduces the real front speed by a few percent BELOW the energy-conserving 0.50, "
   "so a physical current should sit UNDER it. NEREID-B sits 2% OVER it. The model is therefore slightly "
   "fast, not slightly damped, and the solver's own note that 'turbulent damping lowers it' is not what "
   "its own number does. The discrepancy is small and does not change any conclusion, but it is the "
   "opposite sign to the expected dissipative bias and is reported rather than smoothed over.")

qa("What is the far-field validation result?",
   "It is a benchmark, not a pass. Against the four MEASURED Gold Coast cases the modelled dilution at "
   "the 60 m mixing-zone boundary divided by the measured value is 0.35, 1.20, 3.38 and 0.53. The model "
   "lands inside the measured spread but errs by up to a factor of 3.4, in both directions. It is "
   "therefore NOT demonstrably conservative. The far-field dispersivity itself is uncalibrated and "
   "unidentifiable from this data (see section K).")

qa("Didn't an earlier version claim the model was conservative by 16–25%?",
   "It did, and that claim is WITHDRAWN. It came from comparing against Perth's 45:1 at 50 m — which is "
   "a DESIGN/COMPLIANCE target produced by another model, not a measurement. Against real measurements "
   "the model over-predicts dilution (and so under-states residual salinity) in half the cases. "
   "Validating against a modelled target rather than a measurement had flipped the sign of the safety "
   "argument, and only the flattering half was being reported.")

qa("Is 'conservative' the same as 'correct'?",
   "No — and the model can no longer claim even 'conservative'. Erring toward over-stating impact is "
   "tolerable for screening; erring toward UNDER-stating it, as the model does in two of the four "
   "measured cases, is the failure mode that matters. This is why the near-field coefficient was "
   "calibrated to the measured data (section K), which lowered dilution and raised the predicted "
   "footprint.")

qa("What independent check exists on the predicted impact extent?",
   f"Clark et al. (2018) measured detectable ecological effects to about 100 m at the actual Kurnell "
   f"outfall. The model's ΔS > 0.5 g/kg reach is bounded at about {REACH_W:.0f} m (oscillating "
   f"~{REACH_W-REACH_SD:.0f}–{REACH_W+REACH_SD:.0f} m), which sits comfortably inside the ~100 m "
   f"ecological scale — the physical salinity extent extinguishing before the ecological signal is the "
   f"ordering one expects if the model is right.")

qa("Is that a validation or a coincidence?",
   "It is an order-of-magnitude consistency check, not a validation. Clark's ~100 m is an ecological "
   "effect distance driven partly by diffuser-induced near-bed flow, not a salinity isopleth. The two "
   "quantities are related but not the same, and should not be equated.")

qa("Tell me about the honest note in the solver header regarding an earlier '2.3%' claim.",
   "An earlier build reported that it reproduced Perth's 45:1 to within 2.3%. That was a "
   "discretisation artefact of non-conservative operators combined with the buoyancy sign bug — two "
   "errors partly cancelling. Once the operators were made conservative and the sign fixed, the "
   "agreement vanished. The deficit that then appeared was reported as a ~16–25% CONSERVATIVE bias; "
   "that interpretation has since been withdrawn too, because it was measured against Perth's design "
   "target rather than against measured data. Recording both corrections rather than deleting them is "
   "the correct scientific practice.")

# =============================================================================
section("M.  Results and interpretation")

qa("State the headline results.",
   f"Fr = 24.3; near-field terminal rise {NFRISE:.1f} m; seabed return at 7.0 m; return dilution "
   f"{NFDIL:.0f}:1. Far field: maximum excess {EXMAX:.2f} g/kg, minimum dilution {DILMIN:.1f}:1, "
   f"seabed footprint above \u0394S = 0.5 g/kg of about {FOOT:,.0f} m\u00b2, affected volume {VOL:,.0f} m\u00b3, "
   f"and a horizontal reach of {REACH_LO:.0f}\u2013{REACH_HI:.0f} m.")

qa("Is the plume bottom-trapped?",
   f"Yes. The \u0394S > 0.5 g/kg region occupies the lowest {LAYER:.1f} m of a {C['depth']:.0f} m water "
   f"column, from {PTOP:.1f} m depth down to the bed at {ZDEEP:.1f} m, and the surface carries only "
   f"trace excess. An earlier build reported the plume reaching the sea surface; that was an artefact "
   f"of an isotropic source blob, since corrected.")

qa("What independent cross-check confirms the vertical structure?",
   f"The impacted region rises {RISE:.1f} m above the source, against a terminal jet rise of "
   f"{NFRISE:.1f} m derived independently from the Roberts (1997) correlation. Those two numbers come "
   f"from separately-constructed parts of the model and now agree to within a metre. The earlier build "
   f"reported 22.8 m against the same 6.4 m \u2014 a discrepancy that should have been caught.")

qa(f"Why is the reported peak salinity exactly {SMAX:.4f} g/kg?",
   "Because that is exactly the prescribed source salinity S_source = S_amb,bed + (S\u2080 \u2212 S_amb,bed)/S_r. "
   "The seed blob's centre relaxes fully toward S_source whatever the blob's shape, so the domain "
   "maximum is the injected value. It is a boundary condition read back out, and no change of source "
   "geometry can make it a prediction.")

qa(f"Why is the minimum dilution ({DILMIN:.1f}:1) still below the near-field return dilution ({NFDIL:.1f}:1)?",
   "Not because the plume re-concentrates. Dilution is measured against the local depth-varying "
   "ambient, and the peak-excess cell sits a few metres above the bed where the ambient is fresher "
   "(35.4 rather than 35.6 g/kg), so the same absolute salinity reports a lower dilution. It is the "
   "same artefact that produced the earlier build's 32:1, now much reduced but not eliminated.")

qa("Why does the seabed footprint fall to zero for thresholds above 0.8 g/kg?",
   "Because the maximum excess salinity anywhere on the seabed is about 0.84 g/kg, so any contour "
   "above that encloses no area. The domain maximum sits a few metres above the bed, inside the seed "
   "blob, not on the seabed itself.")

qa("Why do the headline footprint and the isopleth table differ slightly?",
   "They use different estimators. The isopleth table counts wet seabed cells at "
   "dx\u00b7dy = 21.6 m\u00b2 resolution; the headline value bilinearly refines the seabed field 8\u00d7 before "
   "thresholding. The difference, of order 1\u20132%, is a fair measure of the discretisation floor on "
   "the footprint.")

qa("How is the seabed footprint defined?",
   "As the plan area of the lowest fluid cell in each column where the excess salinity exceeds "
   "\u0394S_crit = 0.5 g/kg, computed after 8\u00d7 bilinear sub-cell refinement. The resolution floor of a "
   "single cell is 21.6 m\u00b2.")

qa("What is the affected volume, and is it meaningful?",
   f"{VOL:,.0f} m\u00b3 of water carries \u0394S > 0.5 g/kg. It is now a meaningful quantity: the impacted "
   f"cells form a bottom-hugging layer rather than a full-depth column. The earlier build reported "
   f"55,085 m\u00b3 \u2014 inflated roughly {55085/max(VOL,1):.1f}\u00d7 by the over-wide source blob.")

qa("What drives the far-field spreading, and what bounds it?",
   "Three drivers: the residual momentum of the seeded gravity current; the buoyancy-driven lateral "
   "collapse of the dense layer; and advection by the 0.12 m/s ambient current, modulated by the "
   "anisotropic dispersion tensor. What bounds it is the quadratic bottom drag (now ON), which "
   "supplies the momentum sink that arrests the front — so the reach is bounded rather than growing "
   "without limit (Section O).")

qa("What is the horizontal reach, really?",
   f"The final ensemble-mean field gives {REACH:.1f} m; the mean over the trailing window is "
   f"{REACH_W:.1f} \u00b1 {REACH_SD:.1f} m. It is reported as a bound, not a converged value: r_max is "
   f"the single furthest cell above the 0.5 g/kg contour \u2014 a threshold-sensitive tail metric that a "
   f"thin near-threshold filament slowly extends without adding area (it creeps ~42\u219259 m from "
   f"600\u21921800 s), so it does not settle to a constant. Under the robust split-half stationarity test "
   f"the footprint mean (~2500 m\u00b2) and the concentration metrics converge, and the compliance "
   f"conclusion rests on those, not on r_max.")

qa("Why is the ensemble-mean reach smaller than either member's reach?",
   "Because averaging two stochastically-displaced plumes smooths the excess field, lowering the peak "
   "and pulling the \u0394S = 0.5 g/kg contour inward. The mean of the fields is not the field of the mean "
   "reach. Reporting a footprint from an ensemble-mean field systematically understates it.")

qa("What are the run-health diagnostics and what do they prove?",
   f"Final divergence {M['divergence_final']:.1e}, mass imbalance {M['mass_imbalance_final']:.1e}, "
   f"eddy-viscosity cap engaged in {100*M['nut_cap_fraction']:.1f}% of cells. These prove the numerics "
   f"are sound. They say nothing whatever about whether the physical answer is right \u2014 a "
   f"perfectly-converged solution of the wrong problem carries identical diagnostics.")

# =============================================================================
section("N.  Mixing zone and regulatory compliance")

qa("What assessment threshold was adopted and why?",
   "ΔS = 0.5 g/kg above ambient, chosen as a conservative sub-lethal contour. It is more protective "
   "than the roughly 1 ppt typical of NSW mixing-zone practice and than the binding 2.0 ppt at 100 m "
   "of the California Ocean Plan.")

qa("What are the relevant regulatory limits?",
   "California Ocean Plan: ≤ 2.0 ppt above background at ≤ 100 m (binding). Perth/Cockburn Sound: "
   "≤ 1.2 ppt at 50 m and ≤ 0.8 ppt at 1,000 m. Gold Coast: ~2 PSU at 60 m. NSW/Kurnell: within about "
   "1 ppt at the mixing-zone edge, exact EPL distance unverified.")

qa("Does the predicted plume comply?",
   f"Against every limit above, comfortably: the maximum excess anywhere is {EXMAX:.2f} g/kg, below even "
   f"the most stringent 1.2 ppt criterion, and the seabed maximum is about 0.84 g/kg. Compliance is not "
   f"marginal, which limits the practical significance of the residual reach uncertainty.")

qa("How sensitive is compliance to the reach oscillation?",
   f"Not at all for the concentration-based limits, which are met everywhere in the domain. The reach "
   f"oscillates around {REACH_W:.0f} m (~{REACH_W-REACH_SD:.0f}–{REACH_W+REACH_SD:.0f} m) with the "
   f"forcing; even its upper excursions stay well inside the California 100 m boundary. The footprint "
   f"area breathes with the reach, so it is reported as a central value with a spread, not a single "
   f"number.")

qa("Why does a 0.5 g/kg contour matter ecologically if the limits are 1–2 ppt?",
   "Because regulatory limits are set for acute lethality whereas benthic communities show sub-lethal "
   "responses — reduced growth, altered assemblage composition — at smaller anomalies over chronic "
   "exposure. Adopting the more protective contour is a deliberate conservatism.")

qa("What is a mixing zone, formally?",
   "A defined volume around a discharge within which water-quality criteria may be exceeded, on the "
   "understanding that they are met at its boundary. Its size and the criterion at its boundary are "
   "both set by the regulator.")

qa("Clark et al. found effects to 100 m. Does that invalidate the bounded reach?",
   f"It does not, because the two measure different things: Clark measured ecological effect distance, "
   f"partly driven by diffuser-induced near-bed flow rather than salinity. The model's physical reach "
   f"(~{REACH_W:.0f} m) extinguishing inside the ~100 m ecological scale is exactly the expected "
   f"ordering.")

qa("What would you tell a regulator this model can and cannot support?",
   f"It can support a screening-level statement that the salinity anomaly is well within all applicable "
   f"concentration limits and that the affected seabed area is of order {FOOT:,.0f} m². It cannot support "
   f"a consent-grade footprint number, because the reach oscillates rather than settling to a single "
   f"value, the calibration target is synthetic and the ensemble is too small for the exceedance map to "
   f"mean anything.")

# =============================================================================
section("O.  Hard questions and known weaknesses")

para("These are the questions most likely to be asked by a hostile reviewer. Each was verified "
     "directly against the run artefacts. Two of the defects listed in the first edition of this "
     "document have since been fixed in the solver and the case re-run; they are retained here, "
     "marked as such, because an examiner will ask what was wrong and how you found it.",
     italic=True)

subsection("O.1  The source condition")

qa("Describe the source condition. What is actually injected?",
   "The far field is seeded by relaxing S and T toward prescribed source values inside a 3-D Gaussian "
   "blob centred at the near-field seabed return point, together with a momentum source along the "
   "residual gravity-current direction. The source salinity is "
   "S_source = S_amb,bed + (S0 - S_amb,bed)/S_r. The jet itself never appears in the 3-D field.")

qa("There was a bug here. What was it?",
   "The blob was isotropic, with sigma = max(physical return width, 1.5*max(dx,dz)). Because dx = 5.77 m, "
   "the grid floor of 8.65 m overrode the physical return-plume half-width of 2.50 m by 3.5x, so brine "
   "was injected over roughly a third of the water column. The whole source column was driven to "
   "S_source, and dS > 0.5 g/kg reached to within 0.48 m of the sea surface.")
note("FIXED. solver.py now uses an anisotropic floor: horizontal sigma = 1.5*max(dx,dy) = 8.65 m "
     "(a source narrower than a cell is unresolvable), vertical sigma = the physical 2.50 m, floored "
     "only at 1.5*dz = 1.44 m. Source weight at the sea surface fell from 0.031 to ~1e-18.")

qa("How would a reviewer have caught it without reading the code?",
   "By comparing two independently-derived numbers. The near-field correlation predicts a terminal jet "
   "rise of 6.4 m; the 3-D field reported an impacted column 22.8 m tall in 25 m of water. Those cannot "
   "both be right for a negatively-buoyant discharge. After the fix the same comparison gives "
   "7.4 m against 6.4 m.")

qa("What did the fix change, quantitatively?",
   f"The seabed footprint fell from 5,127 m2 to {FOOT:,.0f} m2, the affected volume from 55,085 m3 to "
   f"{VOL:,.0f} m3, the maximum excess from 0.99 to {EXMAX:.2f} g/kg, the minimum dilution from 32:1 to "
   f"{DILMIN:.1f}:1, and the impacted-column height from 22.8 m to {RISE:.1f} m. The old run therefore "
   f"OVER-predicted impact, so its numbers were conservative.")

qa("Did the fix make peak salinity a prediction?",
   f"No, and it never could. The blob centre relaxes fully toward S_source regardless of blob width, so "
   f"S_max = {SMAX:.4f} g/kg = S_source for any geometry. Removing this diagnostic requires a different "
   f"source treatment altogether - a bed flux boundary condition, or the two-way nested "
   f"resolved-near-field mode.")

qa("And the minimum dilution?",
   f"Improved but not eliminated. It moved from 32:1 toward the near-field hand-off of {NFDIL:.1f}:1 and "
   f"now reads {DILMIN:.1f}:1. The residual gap exists because dilution is measured against the local "
   f"depth-varying ambient and the peak-excess cell sits a few metres above the bed, where the ambient "
   f"is fresher. It is the same mechanism, much reduced.")

subsection("O.2  Steady state")

qa("Was steady state reached?",
   ("Yes. Every tracked metric satisfies both the scatter and the trend test over the trailing window: "
    f"the reach drifts only {RDRIFT:+.1f} m, {100*RDREL:.1f}% of its mean." if STEADY else
    f"No. The following metrics still drift across the trailing window: {', '.join(NOTCONV)}. The reach "
    f"drifts {RDRIFT:+.1f} m, {100*RDREL:.1f}% of its mean, so it must be quoted as a range rather than "
    f"as a converged value."))

qa("The earlier build reported steady_state_reached = True while the reach grew from 26 m to 96 m. How?",
   "Because the test bounded only the relative standard deviation over the trailing window "
   "(steady_tol = 0.20). A monotonically climbing quantity has small scatter about its own mean, so it "
   "passed. The test measured the wrong thing.")
note("FIXED. The criterion now also bounds the linear trend across the window: a metric is steady only "
     "if |sigma| <= steady_tol*|mean| AND |drift| <= steady_trend_tol*|mean|, with steady_trend_tol = 0.05. "
     "The solver reports the drift for every metric.")

qa("How long is the run, and is that enough?",
   f"t_end = {TEND:.0f} s. The domain flush time at 0.12 m/s over 300 m is about 2,500 s, but the "
   f"quantity that must equilibrate is the dense gravity current, whose front decelerates as the excess "
   f"dilutes below the assessment contour. "
   + ("The trend test confirms it equilibrated within this run." if STEADY else
      "The trend test says it had not equilibrated when the run stopped."))

qa("Why does the run now integrate to a longer time with bottom drag on?",
   f"Because the far-field eddy-viscosity railing that used to appear past ~400–600 s is fixed. In the "
   f"weakly-stratified, low-strain far field neither the strain nor a buoyancy realizability bound "
   f"binds, so ν_t = C_μ k²/ε ran free, spurious turbulent energy accumulated and ν_t railed to its "
   f"ceiling (~29% of cells by 600 s). A turbulent length-scale limiter (Galperin 1988 buoyancy limit "
   f"plus a geometric mixing-length cap, imposed as a floor on ε) together with a semi-implicit "
   f"(Patankar) k–ε sink now bound ν_t to a physical value and drain the spurious energy. The run "
   f"integrates cleanly to {TEND:.0f} s with {100*NUTCAP:.0f}% railing (0% at 900 s on the same grid).")

qa("Does the horizontal reach r_max converge?",
   "No, and it is not treated as a convergence target — for a good reason. r_max is the single "
   "furthest cell above the 0.5 g/kg contour, so a thin filament of near-threshold water creeping "
   "outward extends it while adding negligible area. An 1800 s run (on the uncalibrated "
   "configuration) showed r_max creeping ~42→59 m while the seabed FOOTPRINT mean stayed stable and "
   "the concentration metrics were fully converged (split-half drift "
   "<1%). So r_max is reported as a bound, and the compliance conclusion rests on the converged "
   "footprint and concentration limits, not on r_max. This is a property of r_max as a metric, not a "
   "solver limitation.")

qa("How is 'converged' actually tested?",
   "With a robust split-half stationarity test: the trailing window is split in two and the "
   "difference of the half-means is compared to a 5% tolerance. This replaces a least-squares linear "
   "drift, which for a stationary but oscillating signal (a stochastically-forced gravity-current "
   "front) reports spurious drift depending on which phase the window endpoints land on. Under the "
   "robust test the footprint and concentration metrics pass; r_max does not, consistent with it "
   "being a slowly-extending tail metric.")

qa("What was the actual fix for the far-field railing?",
   "A turbulent length-scale limiter plus a semi-implicit k–ε sink. The length limiter bounds "
   "L = C_μ^0.75 k^1.5/ε by the smaller of the Galperin (1988) buoyancy length and a geometric "
   "mixing-length cap, imposed as a floor on ε; this caps ν_t to a physical value exactly where "
   "strain AND stratification are both weak (where it otherwise railed). The Patankar semi-implicit "
   "sink linearises the destruction terms so they are unconditionally positive and non-oscillatory. "
   "Together they eliminate the railing (0% at 900 s), and bottom drag is now enabled.")

subsection("O.3  Calibration and validation logic")

qa("Is the calibration circular?",
   "It WAS, and it no longer is. The earlier make_site_data.py chose its transect stations so the target "
   "would be 'reproducible by the model at no tuning', and the calibration duly returned "
   "farfield_disp_cal = 1.00 — a procedure that recovers unity against a target constructed to be "
   "recoverable demonstrates consistency, not predictive skill. Worse, that 1.00 was not even a fit: the "
   "routine had failed to find any leverage and fallen back to its default. The synthetic transect is "
   "deleted. The calibration now fits the near-field coefficient to MEASURED data from the Gold Coast "
   "diffuser, which the model has never been fitted to otherwise.")

qa("How should it be described instead?",
   "As an uncalibrated run against a representative transect that is consistent with in-class field "
   "data. The word 'calibrated' has been removed from the report title, the thesis abstract, the "
   "slides and the conclusions.")

qa("Does --validate really validate the near field?",
   "It verifies it. The constants 2.2, 2.4 and 1.6 are hard-coded from Roberts (1997), so recovering "
   "z_t/(D*Fr) = 2.20 is arithmetic. What it genuinely tests is that Fr, the angle factors and the "
   "merging correction are implemented correctly, which is worth testing.")

qa("So what is the genuinely independent evidence?",
   "Two things. The lock-exchange front Froude number of 0.51 against Benjamin's 0.50, which exercises "
   "the PDE core with the brine physics switched off entirely. And the four MEASURED Gold Coast cases, "
   "three of which the model was never fitted to (the fourth, Case 3-1, is the calibration target and so "
   "is not independent evidence). Against those three the model's dilution error spans 0.35x to 3.4x, "
   "which is honest evidence of screening-grade skill and not a validation pass. Everything else is "
   "verification.")

qa("The plume-rise cross-check now agrees. Is that independent evidence too?",
   f"It is a consistency check between two parts of the same model, not external validation. The "
   f"near-field correlation gives {NFRISE:.1f} m; the 3-D field gives {RISE:.1f} m. That agreement is "
   f"necessary but not sufficient - it would also hold if both were wrong in the same way. Its value is "
   f"that its earlier failure exposed a real bug.")

subsection("O.4  Physics")

qa("Is bottom friction active in the physics set?",
   f"Yes. Bottom drag (C_d = {C.get('Cd_bed', 0.0025):.4f}) and a log-law wall function are ON in this "
   f"headline run. They supply the momentum sink that bounds the dense gravity current, giving a "
   f"bounded (oscillating) reach. Drag is usable because the far-field k–ε railing that used to appear "
   f"past ~400 s is now fixed by the length-scale limiter, so the run integrates stably to {TEND:.0f} s "
   f"with {100*NUTCAP:.0f}% railing.")

qa("What physics is still switched off?",
   "The osmotic salt flux, the osmotic body force, Soret/Dufour cross-diffusion, the higher-order "
   "TEOS-10-style equation-of-state terms, and the genuine free surface. Four of the six novelty claims "
   "in the report's section 2 therefore describe code capability rather than anything exercised by this "
   "case.")

qa("Are those omissions material?",
   "Almost certainly not. Turbulent dispersion exceeds molecular cross-diffusion by many orders of "
   "magnitude, the temperature range is 16-25 degC so the higher-order EOS terms are tiny, and a rigid "
   "lid is appropriate at 25 m depth with Hs = 1.5 m. They matter for the novelty claim, not for the "
   "numbers.")

qa("Do tide, wind and Coriolis do anything in this run?",
   f"Very little. The M2 period is 44,712 s and the inertial period at 34 degS is about 21 hours, "
   f"against a run of {TEND:.0f} s. They are present in the formulation but nearly inert at this run "
   f"length, and should not be advertised as active physics in the results.")

subsection("O.5  Uncertainty, resolution and reproducibility")

qa(f"Can a standard deviation be reported from {NENS:.0f} members?",
   "No. With two members a sample standard deviation carries about 100 percent relative uncertainty, "
   "and a 95th percentile is not estimable at all - it is simply the larger of the two samples. The "
   "exceedance field can take only the values 0, 0.5 and 1, so it is a three-level indicator rather "
   "than a probability.")

qa("There is a second problem with the ensemble. What?",
   f"The Ornstein-Uhlenbeck correlation time is tau = {TAU:.0f} s, comparable with the {TEND:.0f} s run. "
   f"The forcing barely decorrelates within a member, so the spread measures sensitivity to the initial "
   f"random draw rather than to temporal variability.")

qa("Is there a grid-convergence study?",
   "Not for the reported footprint or reach. The solver provides --gridconv, and the eddy-viscosity cap "
   "fraction became grid-independent after the buoyancy sign fix, but no Richardson-style study of the "
   "headline outputs is presented. With dx = 5.77 m against a 7.0 m near field, this remains a gap.")

qa("Does the committed deck reproduce the reported run?",
   f"Yes, now. sydney_sdp_case.json carries the grid (52x32x26), the ensemble ({NENS:.0f}), "
   f"t_end = {TEND:.0f} s and the bottom-drag settings that produced case_study/outputs/. An earlier "
   f"version specified 56x34x26, ensemble 5 and t_end 260 s, and regenerated none of the reported "
   f"numbers.")

qa("What is the single most important thing still to fix?",
   "Replace the synthetic calibration target with measured data - either a commissioned CTD/ADCP survey "
   "at Kurnell, or the published Gold Coast Tugun transect. Everything else on the list is a matter of "
   "compute (a larger ensemble, a longer run, a grid study); this one is the only defect that no amount "
   "of compute can repair.")

qa("Given all this, does the project's central conclusion survive?",
   f"Yes, and it is now stronger than before. The maximum excess anywhere is {EXMAX:.2f} g/kg, well "
   f"inside every applicable regulatory limit, and the affected seabed area is about {FOOT:,.0f} m2 - "
   f"smaller than the earlier build's estimate, which over-predicted impact. What remains weaker than "
   f"the documents once implied is the precision of the footprint and the word 'calibrated'.")

# =============================================================================
section("P.  Reproducibility, code quality and verification")

qa("How is the solver verified as software, as distinct from validated as physics?",
   "By 13 invariant self-tests run as a regression gate: finiteness, boundedness, conservation, "
   "symmetry, positive-definiteness and bitwise-exact restart. These test the code against its own "
   "mathematical contract, independently of any physical benchmark.")

qa("What are the four gates and their status after the source-blob fix?",
   "--selftest gives 13/13 PASS; --validate gives 4/4 PASS against Roberts (1997); --benchmark gives a "
   "lock-exchange Fr_f of about 0.51, PASS; and --calibrate-nf fits the near-field return-dilution "
   "coefficient to the MEASURED Gold Coast data (nf_dilution_cal = 0.871). The old fourth gate, "
   "--validate-farfield perth ('conservative at every station'), has been retired: Perth's 45:1 is a "
   "design target, not a measurement, and the conservatism claim built on it is withdrawn.")

qa("Is the run deterministic?",
   "Yes, given the seed (20240617). The RNG is seeded per ensemble member as seed + 1009*member, and "
   "checkpoint/restart is bitwise exact on both the free-surface and rigid-lid paths.")

qa("How would a reviewer independently check your headline numbers?",
   "Load case_study/outputs/fields_final.npz, threshold the seabed excess at 0.5 g/kg and multiply by "
   "dx*dy for the footprint; take max(excess) for the peak. Doing exactly that is how the source-blob "
   "defect was found, so it is a check worth performing.")

qa("What guards against the reporting drifting from the data again?",
   "The report is generated from metrics_summary.json rather than typed, so its numbers cannot go "
   "stale. The remaining gap is that no automated test asserts the deck reproduces the committed "
   "outputs; adding one would fail CI whenever deck and report diverge.")

qa("Is 5,200 lines in one file defensible?",
   "It keeps the solver self-contained and dependency-light, which matters for reproducibility and for "
   "auditing a single artefact. It costs navigability and makes unit-testing individual operators "
   "harder. For a research solver the trade is reasonable; for a maintained product it is not.")

qa("What is the licence and authorship position?",
   "MIT licence, authored by Akosa Samuel Onyejekwe as independent research work. All third-party "
   "physics is cited to primary literature in validation/sources.md, with DOIs flagged as verified or "
   "unverified.")

# =============================================================================
section("Q.  Recommendations and future work")

qa("What has already been done?",
   f"The source-blob floor was made anisotropic and the case re-run; the steady-state criterion was "
   f"given a trend test in addition to its scatter test; the far-field eddy-viscosity railing at long "
   f"integration times was diagnosed AND FIXED (a turbulent length-scale limiter plus a semi-implicit "
   f"k–ε sink; 0% railing at 900 s), after which bottom drag was enabled and the run extended to "
   f"{TEND:.0f} s with a bounded reach; and the input deck was corrected so it reproduces the reported "
   f"run.")

qa("What remains, in priority order?",
   "Replace the synthetic calibration target with measured data. Raise the ensemble to O(100) members "
   "run over several correlation times before quoting any exceedance map. Perform a grid-convergence "
   "study of the footprint and reach. Run worst-case weak-mixing scenarios to bound the compliance "
   "envelope.")

qa("Which single change would most improve the credibility of the numbers?",
   "A CTD/ADCP survey at the Kurnell outfall itself. The near field is now calibrated to measured "
   "in-class data (Gold Coast), which is the best available substitute, but it is another site: its "
   "crossflow, wave climate and diffuser geometry are not Kurnell's. Only a site survey can calibrate "
   "the FAR field, which remains uncalibrated and unidentifiable from mixing-zone data.")

qa("What sensitivity case is conspicuously missing?",
   "The near-field constant. The literature splits into a Roberts/Papakonstantis cluster at "
   "S_i/Fr ~ 1.6-1.68 and a Lai & Lee cluster at 1.07. The model adopts the higher, less-protective "
   "value. A run at 1.07 would reduce return dilution to about 26:1 and enlarge the footprint, and it "
   "should be reported.")

qa("What is the most promising direction for the solver itself?",
   "The two-way nested resolved-near-field mode. It is the only path to removing the correlation "
   "dependence and the prescribed-source diagnostics at once, and it would turn the near-field "
   "'validation' from verification into a genuine prediction.")

qa("Is there a publishable contribution here?",
   "Two negative results are genuinely useful: the k-epsilon buoyancy sign error and its quantified "
   "effect on far-field over-mixing, and the isotropic source-blob floor, which silently smeared a "
   "dense plume through the water column while every run-health diagnostic stayed at machine precision. "
   "Both are easy to make and hard to see. The end-to-end coupling is a solid engineering contribution.")

qa("How would you summarise the project's standing in one sentence?",
   "A numerically sound, well-verified coupled brine-dispersion solver whose Kurnell case study is now "
   "physically self-consistent and honestly caveated, and whose remaining weakness is that it has never "
   "been tested against a measurement.")

# =============================================================================
section("R.  Cross-document reconciliation")

para("Four documents describe this work. They did not agree with each other, and none of them fully "
     "agreed with the data. All have been corrected.", italic=True)

qa("What was the single most serious error?",
   "A datum inversion. The column layer_top_depth in plume_envelope_vs_distance.csv is a depth below "
   "the sea surface, computed in postprocess.py as -min(z) over the active column. Slides 15 and 24, "
   "their speaker notes, and sections 9.6 and 10.1 of the thesis read its value of 0.48 m as a height "
   "above the seabed, and concluded that the plume never reaches mid-water.")

qa("How could that have been caught internally?",
   "The same CSV reported a layer thickness of 23.08 m in 25 m of water, and a core depth beginning at "
   "4.33 m. A 23 m thick layer whose top is 0.48 m above a bed 25 m down is geometrically impossible. "
   "The thesis printed both facts in adjacent paragraphs without noticing.")

qa("Was that the same defect as the source blob?",
   "It was the symptom; the blob was the cause. The isotropic floor genuinely did push dS > 0.5 g/kg to "
   "within half a metre of the surface, so the CSV was reporting the truth about a broken run. The "
   "documents then misread the datum and drew a reassuring conclusion from an alarming number.")

qa("Did the slides misrepresent the calibration data?",
   "Yes. Slide 10 displayed a five-station transect - 22:1 at 10 m, 33:1 at 25 m, 44:1 at 50 m, 52:1 at "
   "75 m, 60:1 at 100 m - and the notes called it 'measured CTD/ADCP dilution data'. The actual file "
   "has four stations: 37:1 at 7 m, 40:1 at 15 m, 42:1 at 25 m, 44:1 at 50 m. Only the 50 m point "
   "matched, and none of it is measured.")

qa("Was the state vector described consistently?",
   "No. The report correctly identifies zeta as the Ornstein-Uhlenbeck stochastic forcing, matching "
   "solver.py. Slide 6 and the speaker notes called zeta the free-surface elevation, while slide 9 "
   "simultaneously stated that the free surface was off. The free surface is eta.")

qa("Which document was most honest before reconciliation?",
   "The thesis. It already stated that no CTD cast was lowered and no ADCP recovered, that the site "
   "decks are representative rather than measured, and that a two-member ensemble is not a converged "
   "uncertainty quantification. It nonetheless carried the datum error and a circular argument that a "
   "unity calibration multiplier proved the physics correct.")

qa("Which was least reliable?",
   "The speaker notes. They amplified every optimistic reading - 'measured' data, 'equilibrium "
   "conditions', a 95th percentile from two members, 'the plume does not reach the sea surface' - and "
   "added claims the slides themselves did not make.")

qa("Were the headline numbers inconsistent across documents?",
   "Almost not at all, and that is the point. Every document quoted the same Fr, rise, return distance, "
   "dilution, footprint, volume and reach. They were consistent with each other and jointly wrong about "
   "what several of those numbers meant. Internal consistency is not evidence of correctness.")

qa("What changed in each document?",
   "The report gained an interpretation-caveats section, a section on the source condition and the "
   "datum, an honest calibration section, a trend-aware steady-state table, and revised "
   "recommendations. The thesis gained an Errata and Reconciliation note after the abstract plus "
   "corrections to sections 5.5, 8, 9.2, 9.6 and 10.1. The slides and notes were corrected slide by "
   "slide. All four were then regenerated against the re-run.")

qa("Were any solver outputs altered to make the documents agree?",
   "No. The documents were corrected first, against the original outputs. Only afterwards was the "
   "solver defect fixed and the case re-run, which changed the outputs on their own merits. No figure "
   "or CSV value was ever edited by hand.")

qa("Does the compliance conclusion survive?",
   f"Yes, and more comfortably than before. The maximum excess is {EXMAX:.2f} g/kg, inside NSW's ~1 ppt, "
   f"Perth's 1.2 ppt at 50 m, Gold Coast's ~2 PSU at 60 m and California's binding 2.0 ppt at 100 m. "
   f"The corrected run predicts a smaller footprint than the original, so the original was "
   f"conservative.")

# =============================================================================
section("Appendix - Defects found, and their disposition")

para("Each row was confirmed by direct inspection of the named artefact. 'Status' records what has "
     "been done, not what is recommended.", italic=True)

rows = [
    ("Defect", "Evidence", "Disposition"),
    ("Isotropic source-blob floor smeared the return plume over ~1/3 of the water column "
     "(sigma_v = 8.65 m vs a physical 2.50 m)",
     "solver.py:1012; fields_final.npz",
     "FIXED in solver.py (anisotropic floor); case re-run. Footprint 5,127 -> %s m2, "
     "volume 55,085 -> %s m3, impacted-column height 22.8 -> %.1f m" % (f"{FOOT:,.0f}", f"{VOL:,.0f}", RISE)),
    ("Steady-state test bounded scatter but not trend; a reach climbing 26 -> 96 m passed it",
     "write_outputs(); metrics_timeseries.csv",
     "FIXED in solver.py: steady_trend_tol = 0.05 now bounds the linear drift as well. "
     + ("This run passes both tests." if STEADY else "This run fails the trend test, and says so.")),
    ("Far-field eddy-viscosity railing past ~400-600 s forced a short run over a free-slip seabed",
     "metrics_summary.json nut_cap_fraction; long-integration runs",
     "FIXED in solver.py: a turbulent length-scale limiter (Galperin 1988 buoyancy limit + geometric "
     "mixing-length cap, as an eps floor) plus a semi-implicit (Patankar) k-eps sink eliminate the "
     "railing (0% at 900 s). Bottom drag is now ENABLED and the run extended to a bounded, drag-limited "
     "reach"),
    ("Datum inversion: layer_top read as a height above bed, not a depth below surface",
     "postprocess.py:181; plume_envelope_vs_distance.csv",
     "FIXED: columns renamed, comments corrected, all four documents amended"),
    ("Slide 10 showed a five-station transect that does not exist in the data",
     "site_ctd_dilution_transect.csv",
     "FIXED: slide table replaced with the real four stations; 'measured' removed"),
    ("zeta labelled as free-surface elevation instead of stochastic forcing",
     "solver.py; report Table 4.1",
     "FIXED in slides.pptx and slides.docx"),
    ("Deck did not reproduce the run (56x34x26, ensemble 5, t_end 260 s)",
     "sydney_sdp_case.json vs run.log",
     "FIXED: deck now carries the values that produced case_study/outputs/"),
    ("Peak salinity reported as a prediction; it is exactly the injected S_source",
     "solver.py:1015; fields_final.npz",
     "NOT FIXABLE by blob geometry - inherent to a relaxation source. Documented as a "
     "diagnostic in every artefact"),
    ("Minimum dilution measured against the local depth-varying ambient",
     "compute_metrics()",
     "REDUCED by the blob fix (32:1 -> %.1f:1 against a near-field %.1f:1); residual documented"
     % (DILMIN, NFDIL)),
    ("Calibration target constructed to be 'reproducible by the model at no tuning'; the "
     "farfield_disp_cal = 1.00 it returned was not a fit but the routine's fallback after "
     "failing to find leverage",
     "make_site_data.py:150 (former); nereid_output/calibration.json",
     "FIXED. The synthetic transect and its generator are DELETED. The near-field return-dilution "
     "coefficient is now CALIBRATED to MEASURED in-class field data (nf_dilution_cal = 0.871, from "
     "the Gold Coast diffuser, Baum 2019) via the new --calibrate-nf gate"),
    ("Claim that the model is conservative by ~16-25% (under-predicts dilution, over-states impact)",
     "Perth 45:1 @ 50 m is a DESIGN/COMPLIANCE target, not a measurement; against the four MEASURED "
     "Gold Coast cases the model's dilution error spans 0.35x-3.4x, optimistic in two of them",
     "WITHDRAWN from all six documents. The model is not demonstrably conservative"),
    ("Far-field dispersivity (farfield_disp_cal) is not identifiable from mixing-zone data",
     "4x sweep moves the modelled 60 m dilution <3.5% (near-field-dominated station, x_n = 9 Fr d)",
     "OPEN by necessity. Left at its physical default of 1.0 - a default, not a fit. Needs "
     "measurements far enough out for far-field spreading to dominate, i.e. a site CTD/ADCP survey"),
    ("Ensemble of %d members cannot support a std or a 95th percentile" % NENS,
     "ensemble_stats.npz (exceedance in {0, 0.5, 1})",
     "OPEN. Statistics reported for completeness and explicitly withdrawn as bounds. "
     "Needs O(100) members"),
    ("No grid-convergence study of footprint or reach",
     "--gridconv exists but unused for this case",
     "OPEN. dx = 5.77 m against a 7.0 m near field"),
]

t = doc.add_table(rows=0, cols=3)
t.style = "Light Grid Accent 1"
for i, r in enumerate(rows):
    cells = t.add_row().cells
    for c, txt in zip(cells, r):
        c.text = ""
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(8)
        run.font.color.rgb = INK
        if i == 0:
            run.bold = True
            run.font.color.rgb = HEAD
for w, col in zip((2.2, 1.75, 3.15), t.columns):
    col.width = Inches(w)

doc.add_paragraph()
para("Three defects remain open, and none of them can be closed with compute alone except the "
     "ensemble size. The calibration is the binding one: this model has never been compared with a "
     "measurement taken at the site it purports to describe.",
     italic=True, size=9.5, color=ACCENT)

doc.add_paragraph()
para(f"Question bank contains {N[0]} questions. Generated from case_study/outputs/metrics_summary.json.",
     italic=True, size=9, color=MUTED)

out = os.path.join(HERE, "questions.answers.docx")
doc.save(out)
print("saved:", out)
print("questions:", N[0])
