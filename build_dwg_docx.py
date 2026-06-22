"""
build_dwg_docx.py  --  Marine-outfall engineering drawing set for the
Sydney Desalination Plant (Kurnell) offshore brine diffuser, derived from the
NEREID-B simulation report (6/report.docx, 6/metrics_summary.json,
6/sydney_case_input.json).

Produces five dimensioned, professionally-rendered A3 sheets:
  SDP-MO-001  Marine Outfall — General Layout (schematic long-section)
  SDP-MO-002  Outfall Pipeline Layout (plan + longitudinal profile)
  SDP-MO-003  Outfall Diffuser — Details
  SDP-MO-004  Brine Dispersion Diagram (environmental)
  SDP-MO-005  Marine Outfall General Arrangement (GA)

…and assembles them, with a cover sheet, drawing register and design basis,
into 6/dwg.docx.
"""
import os, json
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import (Rectangle, Polygon, Circle, Ellipse, FancyArrowPatch,
                                Arc, PathPatch)
from matplotlib.path import Path

import cad_lib as cad
from cad_lib import (new_sheet, add_area, view_label, hdim, vdim, leader, callout,
                     north_arrow, scale_bar, hatch_rect, sheet_note,
                     INK, INK2, ACCENT, WATER, WATER2, GROUND, ROCK, CONC,
                     PLUME, PLUME2, GRID)

HERE = os.path.dirname(os.path.abspath(__file__))
D6 = os.path.join(HERE, "6")

# ---------------------------------------------------------------- design data
CFG = json.load(open(os.path.join(D6, "sydney_case_input.json")))
MET = json.load(open(os.path.join(D6, "metrics_summary.json")))["metrics"]

S0      = CFG["S0"]            # 67.0  g/kg brine
T_B     = CFG["T_b"]           # 22 C
QP      = CFG["Q_d"]           # 0.0453 m3/s per port
DP      = CFG["d_p"]           # 0.120 m
THETA   = CFG["theta_deg"]     # 60 deg
NPORT   = CFG["n_ports"]       # 72
SPACE   = CFG["port_spacing"]  # 5.0 m
NZH     = CFG["nozzle_height"] # 1.0 m
DEPTH   = CFG["depth"]         # 25 m
SAMB_S  = CFG["S_amb_surf"]    # 35.4
SAMB_B  = CFG["S_amb_bot"]     # 35.6
SLOPE   = CFG["bathy_slope"]   # 0.006
UCUR    = CFG["U_current"]     # 0.12 m/s
UD      = 4.01                 # computed exit velocity (report Table 0)
QTOT    = QP*NPORT             # 3.26 m3/s
LDIFF   = (NPORT-1)*SPACE      # 355 m diffuser length
FRD     = MET["Fr_d"]          # 24.3
NF_RISE = MET["nf_rise_m"]     # 6.4
NF_RET  = MET["nf_return_dist_m"]   # 7.0
NF_DIL  = MET["nf_return_dilution"] # 38
EXC_MAX = MET["excess_max"]
SMAX    = MET["S_max"]
RMAX    = MET["r_max_m"]            # 64.8
FOOT    = MET["seabed_footprint_m2"]# 5127
VOL     = MET["affected_volume_m3"] # 55085
ZDEEP   = MET["z_deepest_m"]        # 24.5

PLOTS = []  # (path, caption) for docx


def _save(fig, name):
    # let dimension/leader lines extend slightly past the data window
    for a in fig.axes:
        for ch in a.get_children():
            try:
                ch.set_clip_on(False)
            except Exception:
                pass
    p = os.path.join(D6, name)
    fig.savefig(p, dpi=170, facecolor="white")
    plt.close(fig)
    return p


def set_view(ax, x0, x1, y0, y1, w_mm=None, h_mm=None, equal=True, pad=0.04,
             padx=None, pady=None):
    """Set limits to the true data extents (+pad). equal=True keeps true shape
    (plans/details, centred & letter-boxed in the inset); equal=False fills the
    inset with vertical exaggeration (long-sections / profiles)."""
    xs, ys = (x1-x0), (y1-y0)
    mx = xs*(padx if padx is not None else pad)
    my = ys*(pady if pady is not None else pad)
    ax.set_xlim(x0-mx, x1+mx); ax.set_ylim(y0-my, y1+my)
    ax.set_aspect("equal" if equal else "auto")


# ============================================================================
#  SHEET 1 — MARINE OUTFALL GENERAL LAYOUT (schematic long-section)
# ============================================================================
def sheet1():
    fig, sx = new_sheet("Marine Outfall — General Layout", "SDP-MO-001",
                        "NOT TO SCALE (SCHEMATIC)", "1 OF 5")
    ax = add_area(fig, 16, 78, 388, 196)
    # schematic horizontal axis (m, with break) ; vertical RL (m)
    X0, X1 = 0, 1180
    set_view(ax, X0, X1, -34, 26, 388, 196, equal=False)

    RL0 = 0.0           # sea level RL 0
    bed = -DEPTH        # -25 m at offshore diffuser
    # ----- land -----
    land = Polygon([(0, -8), (0, 18), (210, 18), (255, 2), (255, -8)],
                   closed=True, fc=GROUND, ec=INK, lw=0.9)
    ax.add_patch(land)
    hatch_rect(ax, 0, -8, 255, 0.01)  # baseline
    # SWRO plant block
    ax.add_patch(Rectangle((40, 6), 120, 9, fc="#EFF1F3", ec=INK, lw=1.0))
    ax.text(100, 12.2, "SDP KURNELL — SWRO PLANT", ha="center", va="center",
            fontsize=7.6, fontweight="bold", color=ACCENT)
    ax.text(100, 8.6, "250 ML/d product · recovery ≈ 47%", ha="center",
            va="center", fontsize=6.2, color=INK2)
    leader(ax, (160, 8), (205, -3),
           "Concentrate\n%.2f m³/s @ S₀=%.0f g/kg" % (QTOT, S0), fs=6.6,
           color=ACCENT, va="center")

    # ----- sea -----
    bx0 = 255
    # sloping seabed from shore to offshore (schematic): from -2 at shore to -25
    def bedfn(x):
        x = np.asarray(x, float)
        z = np.where(x < bx0, np.nan,
                     np.interp(x, [bx0, 360, 560, 1180],
                               [-2.0, -10.0, -22.0, -DEPTH-0.6]))
        return z
    xs = np.linspace(bx0, X1, 300)
    zb = bedfn(xs)
    ax.fill_between(xs, RL0, zb, color=WATER, zorder=0)          # water
    ax.fill_between(xs, zb, -34, color=ROCK, zorder=0)           # seabed soil
    ax.plot(xs, zb, color="#6b5d3e", lw=1.0)
    cad.wavy_surface(ax, bx0, X1, RL0, amp=0.5, n=26)
    ax.text(820, 1.7, "TASMAN SEA  (M.S.L. = RL 0.000)", fontsize=7,
            color="#2E6E8E", ha="center", style="italic")

    # ----- break line in tunnel -----
    bkx = 470
    ax.plot([bkx-8, bkx, bkx+4, bkx+8], [-30, -26, -30, -26], color=INK, lw=0.8)
    ax.plot([bkx-8, bkx, bkx+4, bkx+8], [-31, -27, -31, -27], color=INK, lw=0.8)
    ax.text(bkx, -32.6, "TUNNEL LENGTH NOT TO SCALE", fontsize=5.8,
            color=INK2, ha="center")

    # ----- bored tunnel (under seabed) -----
    tun_y = -29
    ax.plot([235, bkx-7], [-6, tun_y], color=INK, lw=2.4)          # land->down
    ax.plot([bkx+9, 980], [tun_y, tun_y], color=INK, lw=2.4)       # horizontal
    ax.plot([235, bkx-7], [-7.2, tun_y-1.2], color=INK2, lw=0.7)   # bottom wall
    ax.plot([bkx+9, 980], [tun_y-1.2, tun_y-1.2], color=INK2, lw=0.7)
    leader(ax, (700, tun_y), (640, -19),
           "Bored conveyance tunnel\n(brine outfall)", fs=6.6, color=INK)

    # ----- riser + diffuser at offshore end -----
    rx = 985
    ax.plot([rx, rx], [tun_y, bed+0.3], color=INK, lw=2.4)         # vertical riser
    # diffuser manifold along seabed
    dfx0, dfx1 = 1000, 1150
    ax.plot([dfx0, dfx1], [bed+0.3, bed+0.3], color=INK, lw=3.0)
    # ports / risers with nozzles
    for xx in np.linspace(dfx0+8, dfx1-6, 7):
        ax.plot([xx, xx], [bed+0.3, bed+NZH], color=INK, lw=1.4)
        ang = np.deg2rad(THETA)
        ax.annotate("", xy=(xx+6*np.cos(ang), bed+NZH+6*np.sin(ang)),
                    xytext=(xx, bed+NZH),
                    arrowprops=dict(arrowstyle="-|>", color=PLUME2, lw=1.4,
                                    mutation_scale=7))
    leader(ax, (rx, -16), (905, -9), "Tunnelled\nriser shaft", fs=6.6, color=INK)

    # ----- near-field brine plume schematic -----
    jx = np.linspace(0, 150, 60)
    apex = NF_RISE
    jz = bed + NZH + apex*np.sin(np.pi*np.clip(jx/120, 0, 1))   # rise then fall
    ax.plot(dfx0+jx*0.0+ jx*0.0, jz, alpha=0)   # noop to keep scale
    # plume envelope (schematic, anchored at diffuser)
    px = dfx0 + jx
    ax.fill_between(px, bed+0.2, bed+0.2+ (NZH+apex)*np.exp(-((jx-55)/55)**2),
                    color=PLUME, alpha=0.55, zorder=1)
    ax.plot(px, bed+0.2+ (NZH+apex)*np.exp(-((jx-55)/55)**2), color=PLUME2, lw=0.9)
    leader(ax, (dfx0+55, bed+NZH+apex), (1050, -12),
           "Negatively-buoyant\nbrine plume", fs=6.6, color=PLUME2)

    # ----- dimensions -----
    vdim(ax, bed+0.3, RL0, dfx1+24, "WATER DEPTH\n%.0f m" % DEPTH, off=0, fs=7,
         color=INK, rot=90)
    vdim(ax, bed+0.3, bed+NF_RISE+0.3, dfx0-26,
         "PLUME RISE\n%.1f m" % NF_RISE, fs=6.6, right=False, rot=90, color=PLUME2)

    # depth ticks (RL)
    for rl in [0, -5, -10, -15, -20, -25]:
        ax.plot([bx0-3, bx0], [rl, rl], color=INK2, lw=0.6)
        ax.text(bx0-5, rl, "RL %0.1f" % rl, fontsize=5.6, color=INK2,
                ha="right", va="center")

    ax.annotate("", xy=(1140, 4), xytext=(1090, 4),
                arrowprops=dict(arrowstyle="-|>", color="#2E6E8E", lw=1.4,
                                mutation_scale=9))
    ax.text(1115, 5.6, "PREVAILING FLOW", fontsize=6, color="#2E6E8E", ha="center")

    view_label(ax, "GENERAL LONGITUDINAL SECTION — PLANT TO OFFSHORE DIFFUSER",
               sub="Schematic; horizontal distances not to scale", loc="lower")

    sheet_note(sx, [
        "1. This sheet is a general schematic of the marine outfall system from the SWRO "
        "plant through a bored conveyance tunnel and tunnelled risers to a submerged "
        "multiport diffuser on the open Tasman shelf in ≈25 m water.",
        "2. Reject brine (concentrate) ≈%.2f m³/s at S₀=%.0f g/kg, T=%.0f °C; "
        "ambient ≈%.1f g/kg." % (QTOT, S0, T_B, SAMB_B),
        "3. Geometry is an indicative deep-diffuser configuration consistent with the "
        "public design basis; absolute figures from the NEREID-B simulation.",
        "4. All levels in metres relative to Mean Sea Level (RL 0.000). Refer to "
        "SDP-MO-002/003/005 for dimensioned layout and details.",
    ], x=16, y=12, w=120)
    PLOTS.append((_save(fig, "dwg_01_general.png"),
                  "Drawing SDP-MO-001 — Marine Outfall General Layout."))


# ============================================================================
#  SHEET 2 — OUTFALL PIPELINE LAYOUT (plan + longitudinal profile)
# ============================================================================
def sheet2():
    fig, sx = new_sheet("Outfall Pipeline Layout — Plan & Longitudinal Profile",
                        "SDP-MO-002", "AS SHOWN", "2 OF 5")

    # ---------- PLAN (upper) ----------
    axp = add_area(fig, 16, 168, 300, 110)
    set_view(axp, 0, 1500, -260, 260, 300, 110)
    # land (shore on left, sea to right). coastline diagonal
    coast_x = [120, 180, 230, 200, 160]
    coast_y = [260, 120, -20, -150, -260]
    land = Polygon(list(zip([0]+coast_x+[0], [260]+coast_y+[-260])),
                   closed=True, fc=GROUND, ec="#6b5d3e", lw=1.0)
    axp.add_patch(land)
    axp.fill_between([0, 1500], -260, 260, color=WATER, zorder=-2)
    axp.add_patch(land)
    axp.text(70, 200, "KURNELL\n(LAND)", fontsize=7.5, fontweight="bold",
             color="#6b5d3e", ha="center")
    axp.text(1150, 210, "TASMAN SEA", fontsize=8, color="#2E6E8E",
             style="italic", ha="center")
    # plant
    axp.add_patch(Rectangle((30, -40), 70, 60, fc="#EFF1F3", ec=INK, lw=1.0))
    axp.text(65, -10, "SDP", fontsize=7, fontweight="bold", ha="center", color=ACCENT)
    # tunnel route (straight, bearing)
    ox, oy = 95, -10
    dx, dy = 1180, 40
    axp.plot([ox, dx], [oy, dy], color=INK, lw=2.2)
    # chainage markers
    L = np.hypot(dx-ox, dy-oy)
    for ch in range(0, int(L)+1, 250):
        f = ch/L
        cxp, cyp = ox+(dx-ox)*f, oy+(dy-oy)*f
        axp.plot([cxp], [cyp], marker="|", ms=7, color=INK,
                 markeredgewidth=1.3)
        axp.text(cxp, cyp-18, "CH %d" % ch, fontsize=5.6, color=INK2, ha="center")
    leader(axp, (650, 12), (620, 150), "Outfall conveyance tunnel\n(bored, sub-seabed)",
           fs=6.8, color=INK)
    # diffuser line (perpendicular-ish, length 355 m) at offshore end
    ddx, ddy = dx, dy
    perp = np.array([-(dy-oy), (dx-ox)]); perp = perp/np.hypot(*perp)
    half = LDIFF/2 * 0.9   # scale fudge for plan visibility
    da = (ddx-perp[0]*LDIFF/2, ddy-perp[1]*LDIFF/2)
    db = (ddx+perp[0]*LDIFF/2, ddy+perp[1]*LDIFF/2)
    axp.plot([da[0], db[0]], [da[1], db[1]], color=ACCENT, lw=3.2,
             solid_capstyle="round")
    # port ticks along diffuser
    for f in np.linspace(0.06, 0.94, 14):
        pxp = da[0]+(db[0]-da[0])*f; pyp = da[1]+(db[1]-da[1])*f
        axp.plot([pxp, pxp+perp[0]*0], [pyp, pyp], marker=".", ms=3, color=ACCENT)
    callout(axp, (ddx, ddy), (1245, -150),
            "SUBMERGED DIFFUSER\n%d ports · L≈%.0f m" % (NPORT, LDIFF))
    # bearing annotation
    brg = (np.degrees(np.arctan2(dx-ox, dy-oy))) % 360
    axp.text(360, -70, "TUNNEL BRG ≈ %03.0f°" % brg, fontsize=6.2, color=INK2,
             rotation=np.degrees(np.arctan2(dy-oy, dx-ox)))
    # depth contours (schematic) parallel to coast
    for d, xx in [(10, 470), (20, 820), (25, 1120)]:
        axp.plot([xx-40, xx+30], [255, -255], color="#7fb4cf", lw=0.7, ls=(0, (5, 4)))
        axp.text(xx-5, -235, "−%dm" % d, fontsize=5.6, color="#2E6E8E")
    north_arrow(axp, 110, 215, r=42)
    scale_bar(axp, 760, -235, 600, 6, 100, "GRAPHIC SCALE — metres (PLAN, indicative)",
              h=14)
    view_label(axp, "PLAN — OUTFALL ROUTE", sub="Coast to offshore diffuser",
               loc="lower")

    # ---------- LONGITUDINAL PROFILE (lower) ----------
    axl = add_area(fig, 16, 86, 388, 74)
    CH0, CH1 = 0, 1500
    set_view(axl, CH0, CH1, -34, 22, 388, 74, equal=False)
    # grid
    for rl in range(-30, 21, 10):
        axl.plot([CH0, CH1], [rl, rl], color=GRID, lw=0.4, zorder=-2)
        axl.text(CH0-12, rl, "RL %d" % rl, fontsize=5.6, color=INK2, ha="right",
                 va="center")
    for ch in range(0, 1501, 250):
        axl.plot([ch, ch], [-32, 20], color=GRID, lw=0.4, zorder=-2)
        axl.text(ch, -33.5, "CH %d" % ch, fontsize=5.6, color=INK2, ha="center")
    # seabed profile (gentle slope to -25)
    chs = np.linspace(CH0, CH1, 200)
    seabed = np.interp(chs, [0, 120, 400, 900, 1500], [3, 0, -10, -21, -DEPTH-0.5])
    axl.fill_between(chs, 0, np.minimum(seabed, 0), color=WATER, zorder=-1)
    axl.fill_between(chs, seabed, -34, color=ROCK, zorder=-1)
    axl.plot(chs, seabed, color="#6b5d3e", lw=1.1)
    cad.wavy_surface(axl, 120, 1500, 0, amp=0.4, n=30)
    axl.text(800, 1.6, "M.S.L. RL 0.000", fontsize=6, color="#2E6E8E", ha="center")
    # tunnel invert
    tch = np.linspace(60, 1180, 100)
    tinv = np.interp(tch, [60, 1180], [2, -29])
    axl.plot(tch, tinv, color=INK, lw=2.4)
    axl.plot(tch, tinv-1.1, color=INK2, lw=0.6)
    leader(axl, (600, np.interp(600, tch, tinv)), (560, -7),
           "Tunnel invert (≈%.2f%% grade)" % (SLOPE*100), fs=6.4, color=INK)
    # riser + diffuser
    axl.plot([1180, 1180], [-29, -DEPTH+0.3], color=INK, lw=2.4)
    axl.plot([1180, 1320], [-DEPTH+0.3, -DEPTH+0.3], color=ACCENT, lw=3.0)
    for xx in np.linspace(1190, 1315, 8):
        axl.plot([xx, xx], [-DEPTH+0.3, -DEPTH+NZH], color=ACCENT, lw=1.2)
    leader(axl, (1180, -18), (1010, -22), "Riser shaft", fs=6.2, color=INK)
    callout(axl, (1250, -DEPTH+0.4), (1250, -10),
            "DIFFUSER\ninv. RL −%.1f" % (DEPTH-NZH))
    vdim(axl, -DEPTH+0.3, 0, 1360, "%.0f m" % DEPTH, fs=6.6, rot=90)
    axl.text(750, -36, "Vertical exaggeration ≈ 10×", fontsize=6, color=INK2,
             ha="center", style="italic")
    view_label(axl, "LONGITUDINAL PROFILE ALONG OUTFALL ALIGNMENT", loc="lower")

    sheet_note(sx, [
        "1. Plan and profile of the outfall conveyance tunnel from the SWRO plant to the "
        "submerged diffuser; alignment and chainages indicative.",
        "2. Seabed −25 m at the diffuser; regional slope ≈%.3f (≈%.2f%%). Levels to MSL "
        "(RL 0.000)." % (SLOPE, SLOPE*100),
        "3. Diffuser: %d ports, %.1f m c/c, total deployed length ≈%.0f m — see SDP-MO-003."
        % (NPORT, SPACE, LDIFF),
    ], x=16, y=12, w=118)
    PLOTS.append((_save(fig, "dwg_02_pipeline.png"),
                  "Drawing SDP-MO-002 — Outfall Pipeline Layout (plan and longitudinal profile)."))


# ============================================================================
#  SHEET 3 — OUTFALL DIFFUSER DETAILS
# ============================================================================
def sheet3():
    fig, sx = new_sheet("Outfall Diffuser — General Details", "SDP-MO-003",
                        "AS SHOWN", "3 OF 5")

    # ---------- PLAN of diffuser line (with break) ----------
    axp = add_area(fig, 16, 214, 388, 64)
    set_view(axp, -12, 90, -16, 16, 388, 64, equal=True)
    # show first 4 ports, break, last 4 ports (compressed)
    # manifold line
    axp.plot([-8, 38], [0, 0], color=INK, lw=3.0)
    axp.plot([52, 86], [0, 0], color=INK, lw=3.0)
    # break
    for bxk in (43, 47):
        axp.plot([bxk-1.4, bxk+1.4], [3, -3], color=INK, lw=0.8)
    axp.text(45, 5, "≈", fontsize=11, ha="center", color=INK)
    # ports (each a riser station with rosette of nozzles in plan)
    def rosette(x):
        axp.add_patch(Circle((x, 0), 1.3, fc="white", ec=ACCENT, lw=1.2, zorder=3))
        for a in (50, 130, 230, 310):
            ar = np.deg2rad(a)
            axp.annotate("", xy=(x+3*np.cos(ar), 3*np.sin(ar)), xytext=(x, 0),
                         arrowprops=dict(arrowstyle="-|>", color=PLUME2, lw=1.0,
                                         mutation_scale=6))
    xs_first = [0, 5, 10, 15]
    xs_last = [60, 65, 70, 75]
    for x in xs_first+xs_last:
        rosette(x)
    # spacing dim
    hdim(axp, 0, 5, 9.5, "%.1f m" % SPACE, fs=7)
    hdim(axp, 0, 15, 13.5, "3 spaces", fs=6.4)
    axp.text(37, 0, "…", fontsize=12, ha="center", va="center")
    leader(axp, (0, -1.3), (-6, -11), "Riser / rosette\nnozzle head (typ.)", fs=6.4,
           color=ACCENT)
    # total length dim spanning break
    hdim(axp, -8, 86, -13.5, "TOTAL DEPLOYED LENGTH  L ≈ %.0f m  (%d ports @ %.1f m c/c)"
         % (LDIFF, NPORT, SPACE), fs=7, above=False)
    north_arrow(axp, 80, 11, r=4.5)
    view_label(axp, "DIFFUSER PLAN  (ports shown part-only; centre omitted)", loc="lower")

    # ---------- SECTION / ELEVATION of a riser ----------
    axe = add_area(fig, 16, 96, 188, 108)
    set_view(axe, -7, 9, -1, 14, 188, 108)
    bed = 0.0
    # seabed
    axe.fill_between([-7, 9], bed, -1.5, color=ROCK, zorder=-1)
    axe.plot([-7, 9], [bed, bed], color="#6b5d3e", lw=1.1)
    hatch_rect(axe, -7, -1.5, 16, 1.5, pattern="xxxx", fc=ROCK, ec="#6b5d3e", lw=0)
    # buried tunnel crown
    axe.add_patch(Circle((0, -4.2), 0, fill=False))  # noop
    # riser pipe up from tunnel
    axe.add_patch(Rectangle((-0.35, bed-1.5), 0.7, 1.5+NZH, fc=CONC, ec=INK, lw=1.0))
    # rosette head
    axe.add_patch(Circle((0, NZH), 0.55, fc="white", ec=INK, lw=1.2))
    # nozzles at 60 deg both sides
    ang = np.deg2rad(THETA)
    for s in (-1, 1):
        x2 = s*4.2*np.cos(ang); z2 = NZH+4.2*np.sin(ang)
        axe.annotate("", xy=(x2, z2), xytext=(s*0.5, NZH),
                     arrowprops=dict(arrowstyle="-|>", color=PLUME2, lw=2.0,
                                     mutation_scale=10))
        # jet spreading cone
        axe.add_patch(Polygon([(s*0.5, NZH), (x2+ s*0.6, z2+0.9), (x2- s*0.0, z2-0.9)],
                              closed=True, fc=PLUME, alpha=0.4, ec="none"))
    # angle arc
    axe.add_patch(Arc((0.5, NZH), 3, 3, angle=0, theta1=0, theta2=THETA,
                      color=INK, lw=0.9))
    axe.text(2.1, NZH+0.8, "θ = %.0f°" % THETA, fontsize=7, color=INK)
    axe.plot([0.5, 3.5], [NZH, NZH], color=INK2, lw=0.6, ls=(0, (4, 3)))  # horiz ref
    # dims
    vdim(axe, bed, NZH, -3.2, "NOZZLE\nELEV.\n%.1f m" % NZH, fs=6.4, right=False, rot=0)
    leader(axe, (0, bed-1.5), (-5, -1.0), "Bored tunnel\n(below seabed)", fs=6.2,
           color=INK)
    leader(axe, (0, NZH), (-5.5, 6.5), "Rosette riser head", fs=6.4, color=INK)
    leader(axe, (3.0*np.cos(ang), NZH+3.0*np.sin(ang)), (5.5, 11.5),
           "Inclined dense-jet\nUd ≈ %.2f m/s" % UD, fs=6.4, color=PLUME2)
    view_label(axe, "SECTION A–A  —  RISER & NOZZLE", loc="lower")

    # ---------- ENLARGED NOZZLE DETAIL ----------
    axn = add_area(fig, 214, 96, 190, 108)
    set_view(axn, -0.05, 0.55, -0.18, 0.36, 190, 108)
    # nozzle body (tapered)
    axn.add_patch(Polygon([(0, -0.06), (0, 0.06), (0.28, 0.04), (0.28, -0.04)],
                          closed=True, fc=CONC, ec=INK, lw=1.2))
    # bore
    axn.add_patch(Rectangle((0.28, -DP/2), 0.06, DP, fc="white", ec=INK, lw=1.2))
    # exit flow
    axn.annotate("", xy=(0.50, 0), xytext=(0.34, 0),
                 arrowprops=dict(arrowstyle="-|>", color=PLUME2, lw=2.4,
                                 mutation_scale=12))
    # diameter dim
    vdim(axn, -DP/2, DP/2, 0.40, "d = %.0f mm" % (DP*1000), fs=7.2, rot=90)
    axn.plot([0.34, 0.40], [DP/2, DP/2], color=INK2, lw=0.5)
    axn.plot([0.34, 0.40], [-DP/2, -DP/2], color=INK2, lw=0.5)
    # centreline
    axn.plot([-0.02, 0.52], [0, 0], color=INK, lw=0.5, ls=(0, (8, 3, 1, 3)))
    leader(axn, (0.31, 0.0), (0.12, 0.30), "Nozzle bore\n(circular port)", fs=6.6,
           color=INK)
    leader(axn, (0.46, 0), (0.30, -0.13),
           "Exit jet  Q=%.4f m³/s\nUd≈%.2f m/s · Frd≈%.1f" % (QP, UD, FRD),
           fs=6.6, color=PLUME2)
    view_label(axn, "DETAIL 1 — NOZZLE/PORT  (enlarged)", loc="lower")

    sheet_note(sx, [
        "1. Submerged multiport diffuser: %d circular ports, %.0f mm dia., %.1f m c/c, "
        "deployed length ≈%.0f m, nozzle invert %.1f m above seabed, elevation θ=%.0f°."
        % (NPORT, DP*1000, SPACE, LDIFF, NZH, THETA),
        "2. Per-port discharge Q=%.4f m³/s, exit velocity Ud≈%.2f m/s, densimetric Froude "
        "number Frd≈%.1f (brine S₀=%.0f g/kg)." % (QP, UD, FRD, S0),
        "3. Rosette riser heads fed from a bored sub-seabed tunnel via tunnelled risers. "
        "Port geometry indicative of a 60° inclined dense-jet design.",
        "4. Dimensions in millimetres unless noted; levels in metres.",
    ], x=16, y=12, w=118)
    PLOTS.append((_save(fig, "dwg_03_diffuser.png"),
                  "Drawing SDP-MO-003 — Outfall Diffuser General Details."))


# ============================================================================
#  SHEET 4 — BRINE DISPERSION DIAGRAM (environmental)
# ============================================================================
def sheet4():
    fig, sx = new_sheet("Brine Dispersion Diagram — Environmental Assessment",
                        "SDP-MO-004", "AS SHOWN", "4 OF 5")

    # centreline data from report (far field)
    cl = np.array([
        [-2.962, 0.971, 32.49, 6.250], [14.346, 0.963, 32.77, 5.288],
        [31.653, 0.810, 38.80, 23.558], [48.961, 0.698, 45.01, 24.519],
        [66.269, 0.357, 87.93, 24.519], [83.576, 0.393, 79.91, 24.519],
        [100.884, 0.168, 187.3, 24.519], [118.192, 0.141, 223.8, 14.904],
    ])

    # ---------- SECTION (upper) ----------
    axs = add_area(fig, 16, 158, 388, 120)
    set_view(axs, -25, 175, -27, 6, 388, 120, equal=False)
    bed = -DEPTH
    axs.fill_between([-25, 175], 0, bed, color=WATER, zorder=-2)
    axs.fill_between([-25, 175], bed, -27, color=ROCK, zorder=-2)
    axs.plot([-25, 175], [bed, bed], color="#6b5d3e", lw=1.1)
    cad.wavy_surface(axs, -25, 175, 0, amp=0.25, n=34)
    axs.text(150, 0.9, "SEA SURFACE (RL 0.0)", fontsize=6, color="#2E6E8E", ha="center")
    # current arrow
    axs.annotate("", xy=(-12, -3), xytext=(-22, -3),
                 arrowprops=dict(arrowstyle="-|>", color="#2E6E8E", lw=1.6,
                                 mutation_scale=10))
    axs.text(-17, -1.8, "Uc=%.2f m/s" % UCUR, fontsize=6, color="#2E6E8E", ha="center")
    # diffuser
    axs.plot([-3, 3], [bed+0.2, bed+0.2], color=INK, lw=3)
    axs.plot([0, 0], [bed+0.2, bed+NZH], color=INK, lw=1.6)
    # NEAR FIELD: dense jet rising to NF_RISE then falling at NF_RET
    t = np.linspace(0, 1, 80)
    jx = NF_RET*1.6*t
    jz = bed+NZH + (NF_RISE)*np.sin(np.pi*np.clip(t/0.62, 0, 1))*np.where(t<0.62, 1,
            np.cos((t-0.62)/0.38*np.pi/2))
    jz = bed+NZH + NF_RISE*np.sin(np.pi*np.clip(t,0,1))   # simple ballistic arc
    axs.plot(jx, jz, color=PLUME2, lw=1.6)
    axs.fill_between(jx, bed, jz, color=PLUME, alpha=0.35)
    axs.add_patch(Arc((0,0),0,0))  # noop
    # FAR FIELD: bottom gravity current as shaded layer thinning downstream
    fx = cl[:,0]; thick = np.interp(fx, fx, np.clip(cl[:,1]*2.2, 0.4, 4.0))
    fxx = np.linspace(NF_RET, 175, 100)
    fth = np.interp(fxx, fx, np.clip(cl[:,1]*2.4, 0.3, 4.5), right=0.3)
    axs.fill_between(fxx, bed, bed+fth, color=PLUME, alpha=0.55, zorder=1)
    axs.plot(fxx, bed+fth, color=PLUME2, lw=1.0)
    # labels for NF rise & return
    vdim(axs, bed+NZH, bed+NZH+NF_RISE, -8, "TERMINAL\nRISE %.1f m" % NF_RISE,
         fs=6.4, right=False, rot=90, color=PLUME2)
    leader(axs, (NF_RET*1.6, bed+0.4), (24, -3),
           "Near-field return\nx≈%.0f m · dilution≈%.0f:1" % (NF_RET, NF_DIL),
           fs=6.4, color=PLUME2)
    leader(axs, (90, bed+0.7), (95, -6.5),
           "Far-field dense bottom gravity current", fs=6.6, color=PLUME2)
    # depth axis
    for rl in [0,-5,-10,-15,-20,-25]:
        axs.plot([-25, -23], [rl, rl], color=INK2, lw=0.6)
        axs.text(-25.5, rl, "RL %d" % rl, fontsize=5.6, color=INK2, ha="right", va="center")
    vdim(axs, bed, 0, 168, "DEPTH %.0f m" % DEPTH, fs=6.6, rot=90)
    # dilution callouts at 50 / 100
    for xx, dil, ex in [(50, 47, 0.67), (100, 175, 0.19)]:
        zz = bed + np.interp(xx, fxx, fth)
        leader(axs, (xx, zz), (xx, 4),
               "@%dm\nS≈%d:1\nΔS≈%.2f" % (xx, dil, ex), fs=6.0, color=INK, va="bottom")
        axs.plot([xx,xx],[bed, zz], color=INK2, lw=0.4, ls=(0,(3,3)))
    view_label(axs, "VERTICAL SECTION ALONG PLUME CENTRELINE", loc="lower")
    axs.text(0.5, 1.02, "BRINE PLUME BEHAVIOUR — NEAR-FIELD JET → FAR-FIELD GRAVITY CURRENT",
             transform=axs.transAxes, ha="center", fontsize=9, fontweight="bold",
             color=ACCENT)

    # ---------- PLAN footprint (lower-left) ----------
    axf = add_area(fig, 16, 80, 196, 74)
    set_view(axf, -40, 130, -70, 70, 196, 74)
    axf.fill_between([-40,130], -70, 70, color=WATER, zorder=-2)
    # nested excess-salinity contours (concentric, current-skewed downstream)
    contours = [(2.0, "#B0322A", 18), (1.0, "#E0705E", 34), (0.5, "#F2C9C2", 64)]
    for dS, col, reach in contours[::-1]:
        ell = Ellipse((reach*0.45, 0), width=reach*1.7, height=reach*1.25,
                      fc=col, ec="none", alpha=0.6, zorder=1)
        axf.add_patch(ell)
    for dS, col, reach in contours:
        axf.add_patch(Ellipse((reach*0.45, 0), width=reach*1.7, height=reach*1.25,
                      fill=False, ec=col, lw=1.0, zorder=3))
    # diffuser
    axf.plot([0,0], [-18,18], color=INK, lw=3, solid_capstyle="round")
    axf.text(0, 22, "DIFFUSER", fontsize=6, ha="center", color=INK, fontweight="bold")
    # current arrow
    axf.annotate("", xy=(-18, 50), xytext=(-34, 50),
                 arrowprops=dict(arrowstyle="-|>", color="#2E6E8E", lw=1.6, mutation_scale=9))
    axf.text(-26, 54, "current", fontsize=5.8, color="#2E6E8E", ha="center")
    # reach dim
    hdim(axf, 0, RMAX, -52, "HORIZONTAL REACH (ΔS>0.5) ≈ %.0f m" % RMAX, fs=6.4, above=False)
    leader(axf, (RMAX*0.45, 30), (95, 58),
           "Seabed footprint\n(ΔS>0.5 g/kg)\n≈ %.0f m²" % FOOT, fs=6.2, color=ACCENT)
    north_arrow(axf, 112, 52, r=10)
    scale_bar(axf, 30, -64, 100, 5, 20, "metres", h=4.5)
    view_label(axf, "SEABED FOOTPRINT — PLAN  (ΔS contours)", loc="lower")

    # ---------- legend / env panel (lower-right) ----------
    axg = add_area(fig, 222, 80, 182, 74)
    axg.set_xlim(0, 100); axg.set_ylim(0, 100); axg.set_aspect("auto")
    axg.text(2, 96, "EXCESS-SALINITY (ΔS) CONTOUR KEY", fontsize=8, fontweight="bold",
             color=ACCENT, va="top")
    keys = [("#B0322A", "ΔS ≥ 2.0 g/kg", "core / near-field"),
            ("#E0705E", "ΔS ≥ 1.0 g/kg", "mixing-zone order"),
            ("#F2C9C2", "ΔS ≥ 0.5 g/kg", "sub-lethal assessment contour")]
    for i,(c,a,b) in enumerate(keys):
        yy = 86-i*9
        axg.add_patch(Rectangle((3, yy-3), 8, 5, fc=c, ec=INK, lw=0.6))
        axg.text(14, yy-0.5, a, fontsize=7, fontweight="bold", va="center")
        axg.text(48, yy-0.5, b, fontsize=6.2, color=INK2, va="center")
    # key results box
    axg.add_patch(Rectangle((2, 4), 96, 52, fill=False, ec=INK, lw=0.8))
    axg.text(4, 53, "PREDICTED DISPERSION METRICS (NEREID-B)", fontsize=7,
             fontweight="bold", color=ACCENT, va="top")
    rows = [
        ("Near-field return dilution", "≈ %.0f : 1" % NF_DIL),
        ("Terminal plume rise", "%.1f m" % NF_RISE),
        ("Peak excess salinity ΔS_max", "%.2f g/kg" % EXC_MAX),
        ("Centreline dilution @ 50 m", "≈ 47 : 1   (ΔS≈0.67)"),
        ("Centreline dilution @ 100 m", "≈ 175 : 1  (ΔS≈0.19)"),
        ("Horizontal reach (ΔS>0.5)", "%.0f m" % RMAX),
        ("Seabed footprint (ΔS>0.5)", "%.0f m²" % FOOT),
        ("Affected water volume", "%.0f m³" % VOL),
    ]
    for i,(a,b) in enumerate(rows):
        yy = 48-i*5.4
        axg.text(5, yy, a, fontsize=6.4, va="center", color=INK)
        axg.text(96, yy, b, fontsize=6.4, va="center", ha="right", fontweight="bold",
                 color=INK)
    axg.axis("off")

    sheet_note(sx, [
        "1. Environmental dispersion diagram for the SDP brine plume from the NEREID-B "
        "coupled near-/far-field simulation. Contours are excess salinity ΔS above ambient.",
        "2. NSW mixing-zone practice limits the salinity anomaly to ≈1 ppt above ambient at "
        "the mixing-zone boundary; the 0.5 g/kg contour is a conservative sub-lethal screen "
        "(benthic / Posidonia protection).",
        "3. Values indicative and conservative (model under-predicts dilution). A site "
        "CTD/ADCP survey would calibrate absolute figures.",
    ], x=16, y=12, w=120)
    PLOTS.append((_save(fig, "dwg_04_dispersion.png"),
                  "Drawing SDP-MO-004 — Brine Dispersion Diagram (environmental assessment)."))


# ============================================================================
#  SHEET 5 — MARINE OUTFALL GENERAL ARRANGEMENT (GA)
# ============================================================================
def sheet5():
    fig, sx = new_sheet("Marine Outfall — General Arrangement (GA)", "SDP-MO-005",
                        "AS SHOWN", "5 OF 5")

    # ---------- GA PLAN (upper-left) ----------
    axp = add_area(fig, 16, 168, 250, 110)
    set_view(axp, -40, 400, -120, 120, 250, 110)
    axp.fill_between([-40,400], -120, 120, color=WATER, zorder=-2)
    # tunnel arriving
    axp.plot([-40, 0], [0, 0], color=INK, lw=2.4)
    leader(axp, (-25, 0), (-30, 55), "from tunnel /\nriser shaft", fs=6.2, color=INK)
    # riser shaft
    axp.add_patch(Circle((0,0), 6, fc=CONC, ec=INK, lw=1.2))
    # diffuser line of ports (plan), full length to scale
    axp.plot([0, LDIFF], [0,0], color=INK, lw=3.4, solid_capstyle="round")
    # show subset of ports (every ~6th) with alternating nozzles
    for i in range(0, NPORT, 4):
        x = i*SPACE
        axp.plot([x, x], [0, 7], color=ACCENT, lw=1.0)
        axp.plot([x, x], [0, -7], color=ACCENT, lw=1.0)
        axp.plot([x],[0], marker=".", ms=3, color=ACCENT)
    # dims
    hdim(axp, 0, SPACE, 18, "%.1f" % SPACE, fs=6.2)
    hdim(axp, 0, LDIFF, -30, "DIFFUSER LENGTH  L = (%d−1)×%.1f = %.0f m" % (NPORT, SPACE, LDIFF),
         fs=7, above=False)
    leader(axp, (4*4*SPACE, 7), (90, 80), "Ports @ %.1f m c/c (typ.)\n%d No. total"
           % (SPACE, NPORT), fs=6.4, color=ACCENT)
    north_arrow(axp, 360, 80, r=20)
    view_label(axp, "GA — PLAN OF DIFFUSER", loc="lower")

    # ---------- key data schedule (upper-right) ----------
    axt = add_area(fig, 274, 168, 130, 110)
    axt.set_xlim(0,100); axt.set_ylim(0,100); axt.set_aspect("auto"); axt.axis("off")
    axt.add_patch(Rectangle((1,1), 98, 98, fill=False, ec=INK, lw=1.0))
    axt.add_patch(Rectangle((1,90), 98, 9, fc=ACCENT, ec=INK, lw=1.0))
    axt.text(50, 94.5, "OUTFALL DESIGN SCHEDULE", color="white", fontsize=8.2,
             fontweight="bold", ha="center", va="center")
    sched = [
        ("Plant product capacity", "250 ML/d (→500)"),
        ("SWRO recovery", "≈ 47 %"),
        ("Concentrate flow Qtot", "%.2f m³/s" % QTOT),
        ("Brine salinity S₀", "%.0f g/kg" % S0),
        ("Brine temperature", "%.0f °C" % T_B),
        ("Ambient salinity", "%.1f–%.1f g/kg" % (SAMB_S, SAMB_B)),
        ("Water depth at diffuser", "%.0f m" % DEPTH),
        ("Number of ports", "%d" % NPORT),
        ("Port diameter d", "%.0f mm" % (DP*1000)),
        ("Port spacing", "%.1f m" % SPACE),
        ("Nozzle elevation θ", "%.0f°" % THETA),
        ("Nozzle height above bed", "%.1f m" % NZH),
        ("Per-port flow Q", "%.4f m³/s" % QP),
        ("Exit velocity Ud", "%.2f m/s" % UD),
        ("Diffuser length L", "%.0f m" % LDIFF),
        ("Densimetric Froude Frd", "%.1f" % FRD),
    ]
    for i,(a,b) in enumerate(sched):
        yy = 86 - i*5.25
        axt.text(4, yy, a, fontsize=6.3, va="center", color=INK)
        axt.text(96, yy, b, fontsize=6.3, va="center", ha="right", fontweight="bold",
                 color=INK)
        if i % 2 == 1:
            axt.add_patch(Rectangle((1.5, yy-2.6), 97, 5.25, fc="#F1F4F6",
                          ec="none", zorder=-1))

    # ---------- GA LONGITUDINAL ELEVATION (lower-left) ----------
    axe = add_area(fig, 16, 86, 250, 74)
    set_view(axe, -40, 400, -28, 6, 250, 74, equal=False)
    bed = -DEPTH
    axe.fill_between([-40,400], 0, bed, color=WATER, zorder=-2)
    axe.fill_between([-40,400], bed, -28, color=ROCK, zorder=-2)
    # slope the bed slightly
    bx = np.linspace(-40, 400, 50); bz = bed + (bx- (-40))*SLOPE*0  # ~flat at scale
    axe.plot([-40,400],[bed,bed], color="#6b5d3e", lw=1.1)
    cad.wavy_surface(axe, -40, 400, 0, amp=0.4, n=40)
    axe.text(180, 1.1, "M.S.L.  RL 0.000", fontsize=6, color="#2E6E8E", ha="center")
    # riser
    axe.plot([0,0],[bed+0.3, bed-2], color=INK, lw=2.4)
    # diffuser manifold
    axe.plot([0, LDIFF],[bed+0.3, bed+0.3], color=INK, lw=3.2)
    for i in range(0, NPORT, 3):
        x = i*SPACE
        axe.plot([x,x],[bed+0.3, bed+NZH], color=ACCENT, lw=0.8)
        ar = np.deg2rad(THETA)
        axe.annotate("", xy=(x+3*np.cos(ar), bed+NZH+3*np.sin(ar)), xytext=(x, bed+NZH),
                     arrowprops=dict(arrowstyle="-|>", color=PLUME2, lw=0.8, mutation_scale=5))
    vdim(axe, bed+0.3, 0, -28, "%.0f m" % DEPTH, fs=6.6, rot=90)
    vdim(axe, bed+0.3, bed+NZH, LDIFF+20, "%.1f m" % NZH, fs=6.2, rot=90)
    leader(axe, (0, bed-2), (-30, -10), "riser shaft", fs=6, color=INK)
    for rl in [0,-10,-20,-25]:
        axe.plot([-40,-37],[rl,rl], color=INK2, lw=0.6)
        axe.text(-41, rl, "RL %d" % rl, fontsize=5.6, color=INK2, ha="right", va="center")
    view_label(axe, "GA — LONGITUDINAL ELEVATION", loc="lower")

    # ---------- GA CROSS-SECTION (lower-right) ----------
    axc = add_area(fig, 274, 86, 130, 74)
    set_view(axc, -10, 10, -2, 12, 130, 74)
    bedc = 0
    axc.fill_between([-10,10], bedc, -2, color=ROCK, zorder=-2)
    axc.fill_between([-10,10], bedc, 12, color=WATER, zorder=-3)
    axc.plot([-10,10],[bedc,bedc], color="#6b5d3e", lw=1.1)
    # riser + rosette head
    axc.add_patch(Rectangle((-0.4, bedc-2), 0.8, 2+NZH, fc=CONC, ec=INK, lw=1.0))
    axc.add_patch(Circle((0, NZH), 0.5, fc="white", ec=INK, lw=1.2))
    ang = np.deg2rad(THETA)
    for s in (-1,1):
        x2=s*5*np.cos(ang); z2=NZH+5*np.sin(ang)
        axc.annotate("", xy=(x2,z2), xytext=(s*0.4,NZH),
                     arrowprops=dict(arrowstyle="-|>", color=PLUME2, lw=1.6, mutation_scale=8))
    axc.add_patch(Arc((0.4,NZH),3,3,theta1=0,theta2=THETA, color=INK, lw=0.8))
    axc.text(2.0, NZH+0.7, "θ=%.0f°" % THETA, fontsize=6.4)
    vdim(axc, bedc, NZH, -3.5, "%.1f m" % NZH, fs=6, right=False, rot=0)
    leader(axc, (0, NZH), (-6, 7), "rosette head", fs=6, color=INK)
    leader(axc, (0, bedc-2), (4.5, -1.4), "tunnel/riser", fs=6, color=INK)
    view_label(axc, "GA — TYPICAL CROSS-SECTION B–B", loc="lower")

    sheet_note(sx, [
        "1. General arrangement of the SDP offshore submerged multiport brine diffuser; "
        "refer to SDP-MO-003 for diffuser/nozzle details and SDP-MO-004 for dispersion.",
        "2. %d ports @ %.1f m c/c, L≈%.0f m, in %.0f m water; brine %.2f m³/s @ %.0f g/kg."
        % (NPORT, SPACE, LDIFF, DEPTH, QTOT, S0),
        "3. Levels in metres to MSL (RL 0.000); dimensions in metres unless noted. Indicative "
        "configuration on public design basis — not for construction.",
    ], x=16, y=12, w=120)
    PLOTS.append((_save(fig, "dwg_05_ga.png"),
                  "Drawing SDP-MO-005 — Marine Outfall General Arrangement."))


# ============================================================================
#  BUILD ALL + DOCX
# ============================================================================
def build_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

    BODY = "Calibri"; AC = RGBColor(0x0B, 0x3D, 0x5C)
    doc = Document()
    sec = doc.sections[0]
    # landscape A3-ish
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11.69), Inches(8.27)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.4))
    doc.styles["Normal"].font.name = BODY
    doc.styles["Normal"].font.size = Pt(11)

    def para(t="", bold=False, italic=False, size=11, color=None, align=None, sa=6):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa)
        if align: p.alignment = align
        r = p.add_run(t); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = BODY
        if color: r.font.color.rgb = color
        return p

    def heading(t, lvl=1):
        p = doc.add_heading(t, level=lvl)
        for r in p.runs:
            r.font.name = BODY; r.font.color.rgb = AC
        return p

    # ---- cover ----
    para("MARINE OUTFALL DRAWING SET", bold=True, size=22, color=AC,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    para("Sydney Desalination Plant (Kurnell, NSW) — Offshore Submerged Multiport "
         "Brine Diffuser", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para("Tasman-Shelf Outfall · Derived from the NEREID-B Coupled Stochastic-PDE "
         "Brine-Dispersion Simulation", italic=True, size=10.5,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para("A.S. Onyejekwe — Environmental Fluid Mechanics & Computational Modelling",
         size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    para("Issued for Design · Rev 0 · 2026-06-18", size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, sa=10, color=RGBColor(0x44,0x44,0x44))

    heading("Drawing Register", 2)
    reg = [
        ("SDP-MO-001", "Marine Outfall — General Layout (schematic long-section)", "NTS"),
        ("SDP-MO-002", "Outfall Pipeline Layout — plan & longitudinal profile", "As shown"),
        ("SDP-MO-003", "Outfall Diffuser — general details (plan, section, nozzle)", "As shown"),
        ("SDP-MO-004", "Brine Dispersion Diagram — environmental assessment", "As shown"),
        ("SDP-MO-005", "Marine Outfall — General Arrangement (GA)", "As shown"),
    ]
    tb = doc.add_table(rows=1, cols=3); tb.style = "Light Grid Accent 1"
    hdr = tb.rows[0].cells
    for c,txt in zip(hdr, ["Drawing No.", "Title", "Scale (A3)"]):
        r = c.paragraphs[0].add_run(txt); r.bold = True; r.font.size = Pt(10.5); r.font.name = BODY
    for no,ti,sc in reg:
        cells = tb.add_row().cells
        for c,txt in zip(cells, [no,ti,sc]):
            r = c.paragraphs[0].add_run(txt); r.font.size = Pt(10); r.font.name = BODY

    para("", sa=4)
    heading("Design Basis & Notes", 2)
    for t in [
        "Source: dimensions and predicted dispersion metrics are taken from the project "
        "report (6/report.docx) and the NEREID-B simulation outputs (6/metrics_summary.json, "
        "6/sydney_case_input.json) for the Sydney Desalination Plant offshore diffuser.",
        "Discharge: SWRO product 250 ML/d (expandable to 500 ML/d), recovery ≈47%%; "
        "concentrate ≈%.2f m³/s at S₀=%.0f g/kg and %.0f °C into a %.1f–%.1f g/kg ambient."
        % (QTOT, S0, T_B, SAMB_S, SAMB_B),
        "Diffuser: %d circular ports of %.0f mm diameter at %.1f m centres (deployed length "
        "≈%.0f m), nozzle invert %.1f m above seabed at %.0f° elevation, per-port flow "
        "%.4f m³/s, exit velocity ≈%.2f m/s, densimetric Froude number ≈%.1f, in %.0f m water."
        % (NPORT, DP*1000, SPACE, LDIFF, NZH, THETA, QP, UD, FRD, DEPTH),
        "Predicted performance: near-field return dilution ≈%.0f:1, terminal rise %.1f m, "
        "peak excess ΔS_max %.2f g/kg, horizontal reach (ΔS>0.5 g/kg) ≈%.0f m, seabed "
        "footprint ≈%.0f m², affected volume ≈%.0f m³."
        % (NF_DIL, NF_RISE, EXC_MAX, RMAX, FOOT, VOL),
        "Status: the per-port nozzle geometry is a representative deep-diffuser configuration "
        "consistent with the public design basis (60° inclined dense jet); figures are "
        "indicative and conservative. NOT FOR CONSTRUCTION — a site CTD/ADCP survey is "
        "required to calibrate absolute values.",
        "Convention: levels in metres relative to Mean Sea Level (RL 0.000); dimensions in "
        "metres unless noted otherwise; drawings sized for ISO A3.",
    ]:
        b = doc.add_paragraph(style="List Bullet"); b.paragraph_format.space_after = Pt(3)
        r = b.add_run(t); r.font.size = Pt(10); r.font.name = BODY

    # ---- drawing sheets, one per page ----
    # printable height = page_height - top - bottom margins; size image (A3
    # aspect h/w = 11.69/16.54) so the picture + caption fit one page and no
    # blank pages are generated by overflow.
    usable_h = (sec.page_height - sec.top_margin - sec.bottom_margin)  # EMU
    aspect = 11.69 / 16.54
    img_w = min(Inches(10.4), (usable_h - Inches(0.55)) / aspect)
    for path, cap in PLOTS:
        pic_par = doc.add_paragraph()
        pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_par.paragraph_format.page_break_before = True   # new page, no empty break para
        pic_par.paragraph_format.space_after = Pt(2)
        pic_par.add_run().add_picture(path, width=img_w)
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        r = cp.add_run(cap); r.italic = True; r.font.size = Pt(9.5); r.font.name = BODY
        r.font.color.rgb = RGBColor(0x44,0x44,0x44)

    out = os.path.join(D6, "dwg.docx")
    doc.save(out)
    print("WROTE", out)


if __name__ == "__main__":
    sheet1(); sheet2(); sheet3(); sheet4(); sheet5()
    build_docx()
