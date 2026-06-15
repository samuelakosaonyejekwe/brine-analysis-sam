# -*- coding: utf-8 -*-
"""
build_phd_proposal2.py — comprehensive PhD research proposal for the NEREID-B
brine-salinity modelling project -> phd_proposal2.docx (main folder).
Draws on combo_analysis.docx content; embeds preliminary-results figures.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
D2 = os.path.join(HERE, "2")
BODY = "Calibri"; EQF = "Cambria Math"
C = WD_ALIGN_PARAGRAPH.CENTER

DOC = Document()
DOC.styles["Normal"].font.name = BODY
DOC.styles["Normal"].font.size = Pt(11)


def _bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def h(t, lvl=1):
    return DOC.add_heading(t, level=lvl)


def para(t="", bold=False, italic=False, size=11, align=None, color=None, after=6):
    p = DOC.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if t:
        r = p.add_run(t); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = BODY
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def eq(t):
    p = DOC.add_paragraph(); p.alignment = C; p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = EQF; r.font.size = Pt(12)


def bullet(t, sub=False):
    p = DOC.add_paragraph(style="List Bullet 2" if sub else "List Bullet")
    r = p.add_run(t); r.font.name = BODY; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def numbered(t):
    p = DOC.add_paragraph(style="List Number")
    r = p.add_run(t); r.font.name = BODY; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def table(header, rows, widths=None, fs=9, mono0=False):
    t = DOC.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htx in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(htx); run.bold = True; run.font.size = Pt(fs)
        run.font.name = BODY; _bg(c, "1F4E79"); run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(fs)
            run.font.name = EQF if (mono0 and i == 0) else BODY
    if widths:
        for i, w in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)


def figure(fname, caption, width=5.6):
    path = os.path.join(D2, fname)
    if not os.path.exists(path):
        para(f"[figure {fname} not found]", italic=True); return
    DOC.add_picture(path, width=Inches(width))
    DOC.paragraphs[-1].alignment = C
    cp = DOC.add_paragraph(); cp.alignment = C
    r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9.5); r.font.name = BODY
    cp.paragraph_format.space_after = Pt(10)


# ======================================================================
#  TITLE PAGE
# ======================================================================
para("DOCTORAL RESEARCH PROPOSAL — PhD (Comprehensive)", bold=True, size=13,
     align=C, color=(0x55, 0x55, 0x55), after=16)
p = DOC.add_paragraph(); p.alignment = C
r = p.add_run("A Coupled Stochastic Partial-Differential-Equation Framework for "
              "Predicting Salinity Dispersion from Submarine Brine Outfalls")
r.bold = True; r.font.size = Pt(21); r.font.name = BODY; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
p = DOC.add_paragraph(); p.alignment = C
r = p.add_run("Development, Field Validation and Application of the NEREID-B "
              "Near-Field/Far-Field Model for Desalination Brine Plumes")
r.italic = True; r.font.size = Pt(14); r.font.name = BODY
para("", after=18)
for label, val in [("Discipline", "Environmental Fluid Mechanics / "
                    "Computational Hydro-Environmental Engineering"),
                   ("Candidate", "[Candidate name]"),
                   ("Supervisors", "[Principal supervisor]; [Co-supervisor]"),
                   ("Host", "[Department, University]"),
                   ("Mode / duration", "Full-time, 3.5–4 years"),
                   ("Submission date", "12 June 2026")]:
    pp = DOC.add_paragraph(); pp.alignment = C; pp.paragraph_format.space_after = Pt(3)
    rr = pp.add_run(f"{label}:  "); rr.bold = True; rr.font.size = Pt(11); rr.font.name = BODY
    rr2 = pp.add_run(val); rr2.font.size = Pt(11); rr2.font.name = BODY
DOC.add_page_break()

# ======================================================================
#  ABSTRACT
# ======================================================================
h("Abstract", 1)
para(
    "Seawater desalination now exceeds 95 Mm³/day globally and continues to "
    "grow; its hyper-saline reject (brine), discharged to the sea, is a "
    "leading marine-environmental concern because the dense effluent sinks, "
    "pools on the seabed and can exceed the salinity tolerance of benthic "
    "ecosystems. Reliable prediction of the three-dimensional, time-evolving "
    "salinity field around an outfall — together with a calibrated measure of "
    "uncertainty arising from the unpredictable sea — is essential for outfall "
    "design and regulatory consent, yet current tools are limited: integral "
    "jet models (CORMIX/VISJET) cannot resolve the unsteady 3-D far field, "
    "while Boussinesq CFD neglects the large brine/ambient density contrast and "
    "treats the sea deterministically.")
para(
    "This research develops, validates and applies NEREID-B, a novel coupled, "
    "unsteady, non-Boussinesq, stochastic partial-differential-equation (PDE) "
    "framework that solves the velocity, pressure, density, salinity, "
    "temperature, turbulence and free-surface fields simultaneously, closed by "
    "a nonlinear equation of state, irreversible-thermodynamics (Soret/Dufour) "
    "cross-diffusion and an explicit osmotic/reverse-osmosis coupling, and "
    "forced by intrinsic space-time stochastic noise so the model returns the "
    "probability distribution of salinity rather than a single value. A "
    "validated near-field/far-field coupling embeds the unresolvable diffuser "
    "jet through calibrated correlations. Extensive preliminary work — full "
    "model formulation, a robust verified solver, near-field validation against "
    "laboratory scaling, far-field grid-convergence, and two industrial case "
    "studies demonstrating that diffuser port spacing is a first-order control "
    "on impact — establishes proof of concept. The doctoral programme will "
    "extend the physics (non-Boussinesq continuity, multi-ion chemistry, "
    "higher-order numerics), deliver the first field validation of the coupled "
    "framework against CTD/ADCP surveys at an operating outfall, build a "
    "machine-learning surrogate for rapid probabilistic compliance screening, "
    "and release an open decision-support tool. The contribution is a new, "
    "validated, uncertainty-aware predictive capability for brine-outfall "
    "salinity dispersion, directly serving desalination engineering and marine "
    "environmental regulation.")
para("Keywords: brine dispersion; desalination outfall; dense jet; coupled PDE; "
     "stochastic modelling; uncertainty quantification; near-field/far-field "
     "coupling; salinity; environmental fluid mechanics.", italic=True, size=10)

# ======================================================================
#  CHAPTER 1 — INTRODUCTION
# ======================================================================
DOC.add_page_break()
h("1. Introduction", 1)

h("1.1 Background and context", 2)
para(
    "Water scarcity, intensified by climate change, population growth and "
    "over-abstraction, has made desalination a cornerstone of water security in "
    "arid and coastal regions. Seawater reverse osmosis (SWRO) dominates new "
    "capacity but recovers only ~40–50 % of the feed as product; the remainder "
    "is rejected as a brine typically 1.5–2× the ambient salinity, often warmer "
    "and containing residual antiscalants and cleaning chemicals. This reject "
    "is overwhelmingly returned to the sea through submarine outfalls, most "
    "modern installations using inclined multiport diffusers to promote rapid "
    "near-field dilution.")
para(
    "Because the reject is denser than seawater it is a negatively-buoyant "
    "effluent: discharged as an inclined jet it rises, bends over under gravity, "
    "falls back to the seabed and spreads as a bottom-hugging gravity current. "
    "Field and laboratory studies have linked the resulting salinity increment "
    "to stress and mortality in seagrasses (e.g. Posidonia), echinoderms and "
    "other osmoconforming benthic biota. Consequently, regulators worldwide "
    "impose salinity-increment limits at the edge of a defined mixing zone, and "
    "operators must demonstrate compliance by predictive modelling at the "
    "consenting stage and by monitoring thereafter.")

h("1.2 Statement of the problem", 2)
para(
    "Predicting outfall impact requires resolving the salinity field across two "
    "coupled regimes — the energetic near field (the buoyant jet) and the "
    "unsteady far field (the bottom gravity current under currents, tides, "
    "waves and stratification) — and quantifying the uncertainty introduced by "
    "the sea's variability. No existing tool does all three: integral models "
    "capture near-field dilution but not the 3-D unsteady far field; CFD "
    "resolves the far field but assumes small density differences and runs "
    "deterministically. The problem this research addresses is the absence of a "
    "single, validated, uncertainty-aware framework spanning both regimes for "
    "dense brine effluent.")

h("1.3 Rationale and motivation", 2)
para(
    "The gap is both scientifically significant and practically urgent. "
    "Scientifically, dense effluent dispersion couples non-Boussinesq "
    "buoyancy, stratified turbulence, anisotropic dispersion and stochastic "
    "forcing in ways not previously unified or validated. Practically, the "
    "rapid global expansion of desalination, alongside tightening marine "
    "protection regimes, creates strong demand for a credible, risk-based "
    "predictive capability that supports low-impact outfall design and "
    "evidence-based consenting.")

h("1.4 Aim, objectives, questions and hypotheses", 2)
para("Aim: to develop, validate and apply a coupled stochastic PDE framework "
     "(NEREID-B) for predicting near- and far-field salinity distribution and "
     "its uncertainty around submarine brine outfalls.")
para("Research questions:", bold=True, after=2)
numbered("Can a coupled stochastic PDE framework predict the unsteady 3-D "
         "salinity field of a brine outfall more completely than existing tools "
         "while remaining robust and affordable?")
numbered("How accurately can a near-field/far-field coupling reproduce measured "
         "dilution and footprint at an operating outfall?")
numbered("How do near-field jet merging and far-field gravity-current processes "
         "interact to determine impact and compliance?")
numbered("What is the predictive uncertainty given the unpredictability of the "
         "sea, and can it be quantified efficiently for design?")
para("Central hypothesis: a coupled non-Boussinesq stochastic PDE model with a "
     "validated near-field closure can predict brine-plume salinity and its "
     "uncertainty to engineering/regulatory accuracy, outperforming "
     "single-paradigm tools.", after=2)
para("Specific objectives are listed in Table 1.1.", after=4)
para("Table 1.1 — Objectives and status at proposal", italic=True, after=2)
table(["#", "Objective", "Status"],
      [["O1", "Formulate the coupled non-Boussinesq stochastic PDE system", "Done (prelim.)"],
       ["O2", "Develop a robust solver with near-field/far-field coupling", "Done (prelim.)"],
       ["O3", "Verify robustness; validate near field vs lab scaling", "Done (prelim.)"],
       ["O4", "Demonstrate stochastic uncertainty quantification", "Done (prelim.)"],
       ["O5", "Extend physics: non-Boussinesq, multi-ion, higher-order time", "Proposed (WP2)"],
       ["O6", "Field-validate the far field (CTD/ADCP at an operating outfall)", "Proposed (WP3)"],
       ["O7", "Build an ML surrogate for probabilistic screening", "Proposed (WP4)"],
       ["O8", "Apply to design/consenting cases; release open tool", "Proposed (WP5)"]],
      widths=[0.4, 4.3, 1.3])

h("1.5 Scope and delimitations", 2)
bullet("Focus: salinity (and coupled temperature/density) dispersion from "
       "submarine brine outfalls in coastal seas; ecological dose–response is "
       "out of scope (the model supplies the salinity exposure that such "
       "assessments require).")
bullet("Spatial scope: from the diffuser nozzle to the far-field mixing zone "
       "(hundreds of metres); regional-scale circulation is imposed as forcing.")
bullet("The near field is represented by validated correlations rather than "
       "nozzle-resolving CFD, which is infeasible at field scale.")

h("1.6 Operational definitions", 2)
table(["Term", "Definition"],
      [["Near field", "The momentum/buoyancy-dominated jet region (first tens "
        "of metres) where most dilution occurs"],
       ["Far field", "The ambient-controlled region where the diluted, dense "
        "plume spreads as a bottom gravity current"],
       ["Dilution", "Ratio (S₀−S_amb)/(S−S_amb) of source-to-local excess salinity"],
       ["Densimetric Froude, Fr", "U/√(g′D): jet momentum vs buoyancy, g′=g(ρ−ρ_a)/ρ_a"],
       ["Mixing zone", "Permitted area within which the salinity-increment limit "
        "may be exceeded"],
       ["Stochastic PDE (SPDE)", "A PDE driven by random forcing, whose solution "
        "is a probability distribution of fields"]],
      widths=[1.6, 4.6], fs=9)

# ======================================================================
#  CHAPTER 2 — LITERATURE REVIEW
# ======================================================================
DOC.add_page_break()
h("2. Literature Review and Theoretical Framework", 1)

h("2.1 Brine discharge and environmental impact", 2)
para(
    "Reviews (e.g. Roberts et al. 2010; Missimer & Maliva 2018) document that "
    "elevated salinity from desalination outfalls is the dominant stressor on "
    "benthic communities, with thresholds as low as +1–2 g kg⁻¹ for sensitive "
    "seagrasses. Regulatory frameworks (e.g. in the Mediterranean, Gulf and "
    "Australia) accordingly set salinity-increment limits, motivating accurate "
    "predictive modelling.")

h("2.2 Near-field jet and plume modelling", 2)
para(
    "The near field of inclined dense jets has been characterised extensively "
    "by laboratory experiment and integral modelling. Roberts et al. (1997), "
    "Cipollina et al. (2005), Roberts & Toms (1987), Lai & Lee (2012) and "
    "Abessi & Roberts (2015) established robust scaling for terminal rise, "
    "return distance and dilution as functions of the densimetric Froude "
    "number and nozzle angle (e.g. z_t/(D·Fr) ≈ 2.1–2.8 and return dilution "
    "S_r ≈ 1.6 Fr at 60°). Integral entrainment models (Jirka 2004; Lee & Chu "
    "2003) and the engineering tools CORMIX, VISJET and JETLAG encapsulate this "
    "knowledge but reduce the plume to a centreline and cannot represent the "
    "unsteady 3-D far field, complex bathymetry or the full sea state.")

h("2.3 Far-field and CFD modelling", 2)
para(
    "Far-field spreading has been modelled with shallow-water and 3-D "
    "hydrodynamic codes (e.g. Delft3D, ROMS, OpenFOAM-based RANS/LES). These "
    "resolve advection and dispersion but commonly invoke the Boussinesq "
    "approximation — questionable for hyper-saline reject — and require "
    "near-nozzle resolution that is infeasible at field scale, leading to the "
    "widely-used but loosely-coupled practice of hand-off from a near-field "
    "tool to a far-field model.")

h("2.4 Stochastic and uncertainty approaches", 2)
para(
    "Environmental flows are increasingly treated probabilistically (ensemble "
    "forecasting; polynomial chaos; stochastic PDEs). Such methods are "
    "established in oceanography and groundwater but have not been "
    "systematically applied to dense-effluent dispersion, where the "
    "unpredictability of currents, tides and wind strongly controls the "
    "footprint. Representing this variability intrinsically — rather than by ad "
    "hoc scenario runs — is a methodological opportunity.")

h("2.5 Thermodynamic and constitutive couplings", 2)
para(
    "The nonlinear seawater equation of state (TEOS-10; IOC et al. 2010) is "
    "essential for hyper-saline, stratified flows. Osmotic pressure, "
    "reverse-osmosis transport and Onsager (Soret/Dufour) cross-diffusion are "
    "well established in physical chemistry but are universally neglected in "
    "plume models; their role at the sharp brine front is unquantified.")

h("2.6 Synthesis and research gap", 2)
para(
    "The literature provides a validated near field, capable but Boussinesq and "
    "deterministic far-field codes, and mature but unexploited stochastic and "
    "thermodynamic methods. What is missing — and what this research will "
    "deliver — is their integration into one validated, non-Boussinesq, "
    "uncertainty-aware framework: the near field is already lab-validated and "
    "the far-field coupling has passed a preliminary multi-point in-class "
    "validation (conservative against the published Perth transect), while "
    "rigorous field validation via a dedicated CTD/ADCP campaign remains the "
    "core outstanding objective to tighten the absolute numbers.")

h("2.7 Theoretical and conceptual framework", 2)
para(
    "The work is grounded in conservation-law continuum mechanics (mass, "
    "momentum, scalar transport) closed by irreversible thermodynamics (Onsager "
    "reciprocity, the nonlinear EOS) and stochastic-process theory "
    "(Ornstein–Uhlenbeck forcing, ensemble/Fokker–Planck description of the "
    "salinity PDF). Conceptually, environmental impact is modelled as a chain: "
    "discharge → validated near-field dilution → stochastic far-field "
    "gravity-current transport → probabilistic salinity exposure → compliance, "
    "with uncertainty propagated end-to-end.")

# ======================================================================
#  CHAPTER 3 — METHODOLOGY
# ======================================================================
DOC.add_page_break()
h("3. Methodology", 1)

h("3.1 Research design", 2)
para(
    "The study follows a model-development-and-validation design within a "
    "post-positivist, quantitative paradigm: a mathematical model is "
    "formulated, implemented, verified (correct solution of the equations) and "
    "validated (correct representation of reality) against laboratory and field "
    "data, then applied to answer the research questions. Verification and "
    "validation (V&V) follow the established hierarchy (code verification → "
    "solution verification/grid-convergence → model validation → predictive "
    "use with quantified uncertainty).")

h("3.2 The coupled stochastic PDE model", 2)
para("NEREID-B solves the state vector q = (ρ, u, p, S, T, k, ε, η):")
eq("∂ρ/∂t + ∇·(ρu) = 0")
eq("∂(ρu)/∂t + ∇·(ρu⊗u) = −∇p + ∇·τ + ρg − 2ρΩ×u + F_wave + F_osm + F_stoch")
eq("∂S/∂t + u·∇S = −(1/ρ)∇·J_S + R_S + ξ_S")
eq("∂T/∂t + u·∇T = −(1/ρc_p)∇·J_Q + Q_rad/(ρc_p) + ξ_T")
eq("ρ = ρ_EOS(S,T,p)   (nonlinear TEOS-10)")
eq("∂ζ_m/∂t = −ζ_m/τ_m + σ_m 𝓛_m^{1/2} Ẇ_m(x,t)")
para("Constitutive closures:")
bullet("Salt flux J_S = −ρ D_S·∇S − ρ D_ST∇T − (ρ D_S L_p/R_gT)∇Π, fusing a "
       "full anisotropic dispersion tensor, Soret cross-diffusion and the "
       "osmotic/RO flux.")
bullet("Osmotic pressure Π(S,T) by a Pitzer-corrected van 't Hoff law.")
bullet("Turbulence by a buoyancy-modified k–ε closure with a Smagorinsky "
       "dissipation floor.")
bullet("Stochastic forcing channels ζ_m (currents, tide, wind, sub-grid "
       "turbulence) as Ornstein–Uhlenbeck colored noise with prescribed "
       "de-correlation time and spatial length scale.")
para("Solving the ensemble yields the predictive salinity PDF P[S(x,t)], its "
     "mean, variance and exceedance probabilities.")

h("3.3 Numerical solver", 2)
para(
    "A finite-volume, MAC-consistent fractional-step projection is used on a "
    "terrain-following grid with partial-cell (shaved-cell) bathymetry and an "
    "implicit free surface; scalars use a monotone TVD (van Leer) scheme; the "
    "constant-coefficient pressure-Poisson operator is LU-factorised once. "
    "Proposed extensions (WP2): a variable-density (non-Boussinesq) projection, "
    "second-order time integration, and multi-tracer (Maxwell–Stefan) "
    "transport.")

h("3.4 Near-field/far-field coupling", 2)
para(
    "The sub-grid diffuser is represented by validated dense-jet correlations "
    "(rise, return point, dilution; with a multiport merging correction), and "
    "the 3-D far field is seeded with the diluted return plume. WP3 will "
    "refine and field-calibrate this coupling, including jet-merging at real "
    "multiport diffusers.")

h("3.5 Verification, validation and uncertainty quantification", 2)
bullet("Verification: automated invariant tests (conservation, boundedness, "
       "divergence control, monotonicity, exact restart); method-of-"
       "manufactured-solutions for new physics (WP2).")
bullet("Solution verification: systematic grid-convergence (Richardson "
       "extrapolation / GCI).")
bullet("Validation: laboratory dense-jet scaling (done); a preliminary "
       "multi-point in-class far-field validation against the published Perth "
       "transect (done — the model matches near-field impact and is conservative "
       "at every far-field station); and the core field campaign (WP3) against "
       "CTD/ADCP/towed-sensor surveys at an operating outfall, testing predicted "
       "profiles, footprint, dilution and the probability-envelope coverage to "
       "tighten the absolute numbers.")
bullet("UQ: Monte-Carlo stochastic ensembles; Sobol sensitivity; and a "
       "polynomial-chaos / neural surrogate (WP4) for rapid screening.")

h("3.6 Field campaign and case-study design", 2)
para(
    "WP3 will instrument an operating SWRO outfall: vessel-based CTD casts and "
    "towed conductivity transects to map the salinity field, bottom-mounted "
    "ADCPs and a thermistor/conductivity mooring for the forcing and "
    "stratification, and met data for wind. Surveys will span tidal and "
    "seasonal states. Case studies (WP5) will cover single-port and multiport "
    "diffusers, stratified and energetic sites, to test generality and inform "
    "design guidance.")

h("3.7 Data management", 2)
para("All field and model data will be version-controlled, documented with "
     "metadata (CF conventions), archived in an open repository where "
     "agreements permit, and released alongside the software to support "
     "reproducibility.")

# ======================================================================
#  CHAPTER 4 — PRELIMINARY RESULTS
# ======================================================================
DOC.add_page_break()
h("4. Preliminary Results (Proof of Concept)", 1)
para("A complete preliminary implementation establishes feasibility and "
     "de-risks the programme:")
bullet("Robustness verified: the solver passes a six-check automated self-test "
       "(no NaN; salinity bounded; machine-level divergence control; "
       "EOS/transport monotonicity; bitwise-exact checkpoint/restart).")
bullet("Near field validated: modelled 60° dense-jet rise z_t/(D·Fr) = 2.20 "
       "lies within the published band 2.1–2.8; return dilution follows "
       "S_r ≈ 1.6 Fr (Figure 4.1).")
bullet("Far field grid-converged: peak excess salinity changes 0 % between the "
       "medium and fine meshes (4.44 → 4.45 g kg⁻¹ over a 4× cell-count range), "
       "with an independent lock-exchange benchmark giving a front Froude number "
       "Fr_f ≈ 0.44, close to the textbook ~0.5.")
bullet("Far field preliminarily validated and conservative: with a realizable "
       "k–ε closure (Durbin 1996) and the corrected buoyancy term the eddy "
       "viscosity no longer over-produces or rails on any grid, so the "
       "turbulence is physical and grid-independent. Against the published Perth "
       "in-class transect (WA EPA App D Table 3-3; Roberts & Abessi 2014) the "
       "model matches the near-field impact (~28.7:1 vs 27.7:1 at the ~5 m "
       "return/impact point) and is conservative at every far-field station "
       "(~28.7:1 vs 33.8:1 at 25.4 m; ~34.6:1 vs 45:1 at 50 m), i.e. it "
       "under-predicts dilution and over-predicts impact — a conservative (safe) "
       "error. This preliminary multi-point in-class validation is conservative "
       "rather than exact; a dedicated CTD/ADCP field campaign (WP3) remains to "
       "tighten the absolute numbers.")
bullet("Stochastic ensembles produce mean, spread and exceedance-probability "
       "maps of the salinity field.")
para("An industrial case — a 150,000 m³/day SWRO outfall — was simulated for "
     "two diffuser configurations. The results (Table 4.1, Figures 4.2–4.3) "
     "show that diffuser port spacing, through jet merging, is a first-order "
     "control on impact.")
para("Table 4.1 — Preliminary case results (single-port vs merged multiport)", italic=True, after=2)
table(["Quantity", "Single-port", "Merged multiport"],
      [["Near-field dilution", "19 ×", "8 ×"],
       ["Excess ΔS at seabed return", "1.5 g/kg", "3.7 g/kg"],
       ["Far-field peak excess", "2.13 g/kg", "4.45 g/kg"],
       ["Exceedance footprint (>2 g/kg)", "≈ 0 m²", "17,812 m²"],
       ["Maximum reach", "45 m", "167 m"],
       ["Compliance (far field)", "compliant", "defined mixing zone"]],
      widths=[2.6, 1.7, 1.9])
figure("fig_nearfield_trajectory.png",
       "Figure 4.1. Validated near-field inclined dense-jet trajectory; "
       "modelled rise and dilution match laboratory scaling.")
figure("fig_seabed_excess_map.png",
       "Figure 4.2. Predicted seabed excess-salinity footprint (merged "
       "diffuser) with the regulatory mixing-zone contour.")
figure("fig_vertical_section.png",
       "Figure 4.3. Vertical section: the dense plume is trapped near the bed "
       "and runs down-slope — the signature of a negatively-buoyant discharge.")
para("These results confirm that the framework is robust, near-field-validated "
     "and capable of resolving design-relevant differences — the foundation on "
     "which the proposed programme builds.")

# ======================================================================
#  CHAPTER 5 — RESEARCH PLAN
# ======================================================================
DOC.add_page_break()
h("5. Research Programme, Timeline and Management", 1)

h("5.1 Work packages", 2)
table(["WP", "Title", "Key activities", "Output"],
      [["WP1", "Consolidation & review",
        "Critical literature review; benchmark vs CORMIX/VISJET; harden the "
        "verified baseline", "Review paper"],
       ["WP2", "Physics extension",
        "Non-Boussinesq projection; multi-ion Maxwell–Stefan chemistry; "
        "2nd-order time; conservative free surface; method-of-manufactured-"
        "solutions verification", "Extended solver; methods paper"],
       ["WP3", "Field validation (core)",
        "Survey design; CTD/ADCP/towed campaign at an operating outfall; "
        "calibration; near- & far-field validation; envelope coverage test",
        "Validation dataset; flagship paper"],
       ["WP4", "UQ & ML surrogate",
        "Polynomial-chaos / neural surrogate; Sobol sensitivity; efficient "
        "probabilistic screening", "Surrogate; UQ paper"],
       ["WP5", "Application & tool",
        "Design/consenting case studies; open decision-support release; "
        "diffuser-design guidance", "Software; case papers"],
       ["WP6", "Synthesis & thesis",
        "Integration; dissemination; thesis", "Thesis; 4–6 papers"]],
      widths=[0.5, 1.4, 3.1, 1.5], fs=8.5)

h("5.2 Timeline (Gantt)", 2)
table(["Work package", "Y1H1", "Y1H2", "Y2H1", "Y2H2", "Y3H1", "Y3H2", "Y4H1"],
      [["WP1 Consolidation", "●", "●", "", "", "", "", ""],
       ["WP2 Physics extension", "", "●", "●", "●", "", "", ""],
       ["WP3 Field validation", "", "", "●", "●", "●", "", ""],
       ["WP4 UQ & surrogate", "", "", "", "●", "●", "●", ""],
       ["WP5 Application & tool", "", "", "", "", "●", "●", "●"],
       ["WP6 Synthesis & thesis", "", "", "", "", "", "●", "●"]],
      widths=[2.0, 0.6, 0.6, 0.6, 0.6, 0.6, 0.6, 0.6], fs=8.5)

h("5.3 Milestones and deliverables", 2)
table(["Milestone", "When", "Deliverable"],
      [["M1 Verified baseline & review", "Month 9", "Review paper submitted"],
       ["M2 Extended physics solver", "Month 21", "Methods paper; released code"],
       ["M3 Field campaign complete", "Month 30", "Validation dataset"],
       ["M4 Validation paper", "Month 36", "Flagship journal paper"],
       ["M5 Surrogate & UQ", "Month 39", "Surrogate tool; UQ paper"],
       ["M6 Tool release & thesis", "Month 45", "Open tool; PhD thesis"]],
      widths=[2.3, 1.0, 2.9], fs=9)

h("5.4 Training and development plan", 2)
bullet("Technical: advanced CFD/numerical methods, stochastic modelling and UQ, "
       "machine learning, oceanographic field methods and instrumentation.")
bullet("Transferable: scientific writing, open-source software engineering, "
       "data management, project and stakeholder management.")
bullet("Dissemination: international conferences (e.g. IAHR, IDA World Congress, "
       "EGU) and engagement with industry and regulators.")

h("5.5 Indicative resources and budget", 2)
table(["Item", "Indicative provision"],
      [["HPC / compute", "Ensemble simulations; cloud or institutional cluster allocation"],
       ["Field campaign", "Vessel time; CTD, ADCP, towed conductivity sensor; consumables"],
       ["Travel", "Conferences; field site visits"],
       ["Partnerships", "Plant operator (site access, plant data); environmental regulator"],
       ["Software", "Open-source toolchain (no licence cost)"]],
      widths=[1.8, 4.4], fs=9)

h("5.6 Risk management", 2)
table(["Risk", "L/I", "Mitigation"],
      [["Field access / weather (WP3)", "M/H",
        "Multiple candidate sites; early operator partnership; existing "
        "monitoring data as fallback"],
       ["Non-Boussinesq solver instability (WP2)", "M/M",
        "Incremental development behind the verified Boussinesq-buoyancy baseline"],
       ["Ensemble computational cost", "M/M", "ML surrogate; parallelisation"],
       ["Closure/merging-factor uncertainty", "M/M",
        "Calibrate to lab & field; report bounded estimates"]],
      widths=[2.3, 0.6, 3.3], fs=8.5)

h("5.7 Ethics", 2)
para("No human or animal subjects are involved. Marine field work will follow "
     "survey permitting, environmental best practice and any protected-area "
     "constraints; operator data will be governed by confidentiality/"
     "data-sharing agreements; outputs will be released openly where permitted.")

# ======================================================================
#  CHAPTER 6 — CONTRIBUTION & IMPACT
# ======================================================================
DOC.add_page_break()
h("6. Expected Contribution to Knowledge and Impact", 1)
para("Original contributions:")
bullet("A validated, non-Boussinesq, stochastic coupled-PDE framework for dense "
       "effluent dispersion returning the salinity probability distribution.")
bullet("The first field validation of a coupled near-field/far-field stochastic "
       "model for brine outfalls, with an openly-published dataset.")
bullet("Quantification of the osmotic, cross-diffusion and jet-merging effects "
       "neglected by current practice.")
bullet("A machine-learning surrogate enabling probabilistic compliance "
       "screening, and an open decision-support tool with design guidance.")
para("Impact: scientifically, the work advances environmental fluid mechanics "
     "by unifying and validating couplings previously treated in isolation; "
     "practically, it equips desalination engineers and regulators with a "
     "risk-based predictive capability that supports low-impact outfall design "
     "and defensible consenting — timely given the rapid global growth of "
     "desalination and tightening marine protection.")

h("7. Dissemination Plan", 1)
bullet("4–6 peer-reviewed papers (e.g. Water Research; Desalination; J. "
       "Hydraulic Engineering; Environmental Fluid Mechanics; Ocean Modelling).")
bullet("Conference presentations (IAHR, IDA, EGU) and an industry/regulator "
       "workshop.")
bullet("Open-source release of the validated solver and surrogate, with "
       "documentation and worked case studies.")

# ======================================================================
#  REFERENCES
# ======================================================================
h("8. Indicative References", 1)
for ref in [
    "Roberts, P.J.W., Ferrier, A. & Daviero, G. (1997). Mixing in inclined "
    "dense jets. J. Hydraulic Engineering, 123(8), 693–699.",
    "Roberts, P.J.W. & Toms, G. (1987). Inclined dense jets in flowing current. "
    "J. Hydraulic Engineering, 113(3), 323–341.",
    "Cipollina, A. et al. (2005). Bench-scale investigation of inclined dense "
    "jets. J. Hydraulic Engineering, 131(11), 1017–1022.",
    "Lai, C.C.K. & Lee, J.H.W. (2012). Mixing of inclined dense jets in "
    "stationary ambient. J. Hydro-environment Research, 6(1), 9–28.",
    "Abessi, O. & Roberts, P.J.W. (2015). Effect of nozzle orientation on dense "
    "jets in stagnant environments. J. Hydraulic Engineering, 141(8).",
    "Jirka, G.H. (2004). Integral model for turbulent buoyant jets, Part I. "
    "Environmental Fluid Mechanics, 4, 1–56.",
    "Lee, J.H.W. & Chu, V.H. (2003). Turbulent Jets and Plumes: A Lagrangian "
    "Approach. Kluwer.",
    "Bleninger, T. & Jirka, G.H. (2008). Modelling and environmentally sound "
    "management of brine discharges from desalination plants. Desalination, "
    "221, 585–597.",
    "Fischer, H.B. et al. (1979). Mixing in Inland and Coastal Waters. "
    "Academic Press.",
    "Missimer, T.M. & Maliva, R.G. (2018). Environmental issues in seawater "
    "reverse osmosis desalination: intakes and outfalls. Desalination, 434, "
    "198–215.",
    "Roberts, D.A. et al. (2010). Impacts of desalination plant discharges on "
    "the marine environment: A review. Water Research, 44(18), 5117–5128.",
    "IOC, SCOR & IAPSO (2010). The international thermodynamic equation of "
    "seawater – 2010 (TEOS-10). UNESCO, Manuals and Guides 56.",
    "Palomar, P. & Losada, I.J. (2011). Impacts of brine discharge on the "
    "marine environment: Modelling as a predictive tool. Desalination, 167.",
]:
    p = DOC.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(18); p.paragraph_format.first_line_indent = Pt(-18)
    r = p.add_run(ref); r.font.size = Pt(9.5); r.font.name = BODY

DOC.save(os.path.join(HERE, "phd_proposal2.docx"))
print("Saved phd_proposal2.docx")
print("paragraphs:", len(DOC.paragraphs), "tables:", len(DOC.tables))
