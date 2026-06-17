#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_model_docx.py  —  generate  model.docx  (salinity_prediction root)

The COUPLED PDE SYSTEM that NEREID-B (solver.py) discretises and integrates.
The equations below are transcribed DIRECTLY from the current solver source
(Rev 2.0), so they reflect the corrected physics and the optional extensions:
  * corrected k-eps buoyancy-production SIGN (damping in stable stratification);
  * realizable k-eps time-scale limiter + Smagorinsky LES floor;
  * conservative, concentration-weighted Soret/Dufour cross-fluxes (G8);
  * SSP-RK2 time integration option (H1);
  * full nonlinear / thermobaric equation of state option (H2);
  * extra passive tracers (H3);
  * Craik-Leibovich Stokes-drift advection + radiation-stress forcing (H4);
  * variable-density (non-Boussinesq) projection option (H5);
  * Sommerfeld/Orlanski radiative outflow option (H6).
Rev 1.6-2.0 added NUMERICAL/closure options only — WALE LES sub-grid viscosity,
one-/two-way grid nesting, the auto fine-mesh resolved-near-field driver, the
on-device (CuPy) pressure-Poisson solve, and the TEOS-10 GSW equation of state —
none of which changes the CONTINUOUS coupled PDE system stated below; they are
alternative closures / discretisations / hardware backends for the SAME equations.
Run:  python3 build_model_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOC = Document()
EQ_FONT = "Cambria Math"
BODY_FONT = "Calibri"
ACCENT = (0x1F, 0x4E, 0x79)

normal = DOC.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
_eqn = [0]


def h(text, level=1):
    p = DOC.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = BODY_FONT
    return p


def para(text="", bold=False, italic=False, size=11, align=None, color=None, space_after=6):
    p = DOC.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = BODY_FONT
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def eq(expr, label=True):
    p = DOC.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(expr)
    r.font.name = EQ_FONT; r.font.size = Pt(12)
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), EQ_FONT); rFonts.set(qn("w:hAnsi"), EQ_FONT); rFonts.set(qn("w:cs"), EQ_FONT)
    if label:
        _eqn[0] += 1
        tab = p.add_run("\t\t(" + str(_eqn[0]) + ")")
        tab.font.name = BODY_FONT; tab.font.size = Pt(11)
    return p


def bullet(text, sub=False):
    p = DOC.add_paragraph(style="List Bullet 2" if sub else "List Bullet")
    r = p.add_run(text); r.font.name = BODY_FONT; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def make_table(header, rows):
    t = DOC.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    for c, txt in zip(t.rows[0].cells, header):
        c.paragraphs[0].add_run(txt).bold = True
    for row in rows:
        cells = t.add_row().cells
        for c, txt in zip(cells, row):
            r = c.paragraphs[0].add_run(txt); r.font.size = Pt(9.5); r.font.name = BODY_FONT
    DOC.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ======================================================================
#  TITLE
# ======================================================================
title = DOC.add_heading("NEREID-B — Governing Coupled PDE System", level=0)
for r in title.runs:
    r.font.color.rgb = RGBColor(*ACCENT)
para("Nonlinear Eulerian Reactive-osmotic Effluent Integro-Dispersion model for a "
     "negatively-buoyant brine plume from an inclined submarine diffuser in a moving, "
     "stratified, wave-/tide-/wind-forced sea.", italic=True, color=ACCENT)
para("This document states the coupled system of partial differential equations that the "
     "reference solver (solver.py, Rev 2.0) discretises by a fractional-step finite-volume "
     "method. z is positive up with the still surface at z = 0; u = (u, v, w); bold symbols "
     "are vectors/tensors; repeated indices are summed.", size=10)

# ----------------------------------------------------------------------
h("1.  State vector and operators", 1)
para("The model advances the coupled state vector")
eq("q = ( u, v, w, p, S, T, ρ, k, ε, ζ, η )", label=False)
para("over a fluid domain Ω(x, y, z, t) bounded below by the bathymetry z = −H(x, y) and "
     "above by the free surface z = η(x, y, t). The material derivative and the resolved "
     "(eddy-viscous) RANS closure use")
eq("D/Dt = ∂/∂t + (u · ∇) ,    Sᵢⱼ = ½ ( ∂uᵢ/∂xⱼ + ∂uⱼ/∂xᵢ ) ,    |S|² = 2 Sᵢⱼ Sᵢⱼ")

# ----------------------------------------------------------------------
h("2.  Continuity (mass conservation)", 1)
para("The flow is treated as incompressible (Boussinesq); the divergence-free constraint is "
     "enforced each step by the pressure projection:")
eq("∇ · u = 0")
para("H5 (optional, non-Boussinesq / low-Mach): the projection instead solves the "
     "variable-density pressure-Poisson equation, so the local density enters the inertia "
     "and pressure coupling exactly:")
eq("∇ · ( (1/ρ) ∇p′ ) = (1/Δt) ∇ · u* ,    uⁿ⁺¹ = u* − Δt (1/ρ) ∇p′")

# ----------------------------------------------------------------------
h("3.  Momentum (buoyancy-modified RANS)", 1)
para("The Reynolds-averaged momentum balance with eddy-viscosity closure, full nonlinear "
     "buoyancy, rotation, and the resolved forcing terms:")
eq("∂uᵢ/∂t + ∂(uⱼ uᵢ)/∂xⱼ = − (1/ρ₀) ∂p/∂xᵢ "
   "+ ∂/∂xⱼ[ (ν + ν_t)(∂uᵢ/∂xⱼ) ] + bᵢ − εᵢⱼ₃ f uⱼ + Fᵢ + ζᵢ")
para("with the buoyancy (acting on the vertical component only)")
eq("bᵢ = − [ g (ρ − ρₐ) / ρ₀ ] δᵢ₃        (Boussinesq);     "
   "bᵢ = − [ g (ρ − ρₐ) / ρ ] δᵢ₃   (non-Boussinesq, H5)")
para("where ρₐ(z) is the ambient (background) density, f = 2Ω sin φ the Coriolis parameter "
     "(beta-plane option f = f₀ + β(y − y₀), G11), and Fᵢ collects the body forces:")
bullet("wind stress on the surface layer, τ = ρ_air C_d |U₁₀| U₁₀ ;")
bullet("quadratic bottom drag on the bed cell, τ_b = ρ₀ C_d^bed |u_b| u_b (optional B1);")
bullet("optional osmotic body force  F_os = − γ_os g ∇(S / S₀) ;")
bullet("radiation-stress (wave-setup) force F_rad (H4, eq. 18);")
para("and ζᵢ the stochastic acceleration (eq. 15). The pressure p is the Lagrange multiplier "
     "enforcing eq. (2); momentum advection is the divergence-consistent MAC flux, so no "
     "spurious φ(∇·u) is produced.")

# ----------------------------------------------------------------------
h("4.  Absolute-salinity transport", 1)
para("Salinity is advected by the divergence-free velocity and dispersed by the full "
     "anisotropic, state-dependent dispersion tensor D (eq. 9), with the thermodiffusive "
     "(Soret) and osmotic/reverse-osmosis cross-fluxes:")
eq("∂S/∂t + ∇ · (u S) = ∇ · ( D ∇S ) "
   "+ σ_So ∇ · ( w_S D ∇T ) + ∇ · ( D_os ∇S ) + Q_S")
para("The Soret salt flux is concentration-weighted (G8, conservative form) so it vanishes "
     "at zero and at saturation, and the osmotic effective diffusivity scales with the local "
     "salinity (because Π ∝ S ⇒ ∇Π ∝ ∇S):")
eq("w_S = 4 (S/S₀)(1 − S/S₀) ,        D_os = κ_os ( S / S₀ )")
para("Q_S is the diluted nozzle source (near-field coupling, §11). The discrete scheme is "
     "monotone (TVD MUSCL) and mass-conserving, with a hard, mass-redistributing bound "
     "0 ≤ S ≤ S₀ (G5).")

# ----------------------------------------------------------------------
h("5.  Temperature transport (with Dufour cross-flux)", 1)
para("Heat is transported analogously, with the reciprocal (Dufour) cross-flux driven by the "
     "salinity gradient:")
eq("∂T/∂t + ∇ · (u T) = ∇ · ( D_T ∇T ) + σ_So d_f ∇ · ( D_T ∇S ) + Q_T")
para("where D_T = κ_T + ν_t / Pr_t (+ geometric anisotropy) and d_f is the explicit Dufour "
     "coefficient (named parameter, G8).")

# ----------------------------------------------------------------------
h("6.  Nonlinear equation of state (density coupling)", 1)
para("Density closes the buoyancy coupling. The default is a linear-contraction law with a "
     "quadratic cabbeling term:")
eq("ρ = ρ₀ [ 1 − α_T (T − T₀) + β_S (S − S₀) ] − ½ ρ₀ c_cab (T − T₀)²")
para("H2 (optional, full_nonlinear, TEOS-10-style) adds higher-order T–S curvature and the "
     "pressure (thermobaric) dependence via the local hydrostatic pressure P:")
eq("ρ_full = ρ + ρ₀ [ ½ β_S2 (S−S₀)² − λ_TS (T−T₀)(S−S₀) + κ_p P − γ_tb P (T−T₀) ] ,   "
   "P = ρ₀ g |z| / 10⁴")

# ----------------------------------------------------------------------
h("7.  Anisotropic dispersion tensor", 1)
para("The symmetric, positive-definite dispersion tensor superposes an isotropic "
     "(molecular + turbulent) part, a flow-aligned Taylor-shear part, a wave-orbital part, "
     "and a bed-confined along-slope part:")
eq("D = ( D_m + ν_t / Sc_t + D_h ) I "
   "+ D_sh ( ê_u ⊗ ê_u ) + D_w ( ê_w ⊗ ê_w ) + D_b ( I − n̂ ⊗ n̂ )")
para("with ê_u the flow direction, ê_w the wave direction, n̂ the seabed normal, "
     "D_sh = a_sh |u_h| Δx (Taylor shear), D_w = g_w (π H_s / T_w) H_s (wave orbital), and "
     "D_b the bathymetric enhancement decaying with height above the bed. Each added term is "
     "positive semi-definite, so D stays SPD (well-posed anisotropic diffusion).")

# ----------------------------------------------------------------------
h("8.  Turbulence closure (buoyancy-modified, realizable k–ε)", 1)
para("Two transport equations close the eddy viscosity:")
eq("∂k/∂t + ∇·(u k) = ∇·[ (ν + ν_t/σ_k) ∇k ] + P_k + G_b − ε")
eq("∂ε/∂t + ∇·(u ε) = ∇·[ (ν + ν_t/σ_ε) ∇ε ] + (ε/k) [ C₁ ( P_k + C₃ max(G_b,0) ) − C₂ ε ]")
para("with shear production P_k and the buoyancy production with its CORRECTED sign "
     "(damping for stable stratification, ∂ρ/∂z < 0 ⇒ G_b < 0):")
eq("P_k = ν_t |S|² ,        G_b = ( ν_t / Pr_t )( g / ρ₀ ) ∂ρ/∂z")
para("The eddy viscosity uses the realizable (Durbin) time-scale limiter plus a Smagorinsky "
     "LES dissipation floor, so it cannot over-produce/rail on fine grids:")
eq("ν_t = C_μ k 𝒯 ,   𝒯 = min( k/ε , C_r / ( √6 C_μ √(|S|²) ) ) ,   "
   "ν_t ← max( ν_t , (C_s Δ)² √(|S|²) )")
para("Optional (les_mode=\"wale\"): a WALE sub-grid eddy viscosity may replace the Smagorinsky "
     "floor on fine/jet grids — correct near-wall cubic scaling and zero ν_t in pure shear:")
eq("ν_t^{WALE} = (C_w Δ)² ( S_{ij}^d S_{ij}^d )^{3/2} / "
   "[ (S̄_{ij} S̄_{ij})^{5/2} + (S_{ij}^d S_{ij}^d)^{5/4} ] ,   "
   "S_{ij}^d = ½(ḡ_{ij}² + ḡ_{ji}²) − ⅓ δ_{ij} ḡ_{kk}²")

# ----------------------------------------------------------------------
h("9.  Stochastic forcing (colored noise)", 1)
para("Each velocity channel is forced by a spatially-correlated Ornstein–Uhlenbeck process "
     "(stationary colored random stress), entering the momentum balance (eq. 3) as an "
     "acceleration:")
eq("∂ζᵢ/∂t = − (1/τ) ζᵢ + σ √(2/τ) · 𝒩ᵢ(x, t)")
para("where 𝒩ᵢ is unit-variance noise smoothed to a spatial correlation length ℓ; τ is the "
     "de-correlation time and σ the r.m.s. intensity. A Monte-Carlo ensemble over realisations "
     "yields the mean, variance and exceedance probability of the salinity field.")

# ----------------------------------------------------------------------
h("10.  Free surface (kinematic + dynamic)", 1)
para("The rigid lid is replaced by an evolving free surface η(x, y, t) with the kinematic "
     "and (linearised) dynamic conditions, folded implicitly into the pressure solve "
     "(unconditionally stable):")
eq("∂η/∂t = w |_{z=η} ,        p |_{z=η} = ρ g η")

# ----------------------------------------------------------------------
h("11.  Near-field coupling (sub-grid nozzle boundary condition)", 1)
para("The unresolvable nozzle is represented by the validated empirical inclined-dense-jet "
     "correlations (Roberts et al. 1997; CORMIX/VISJET-class), which set the diluted seabed "
     "return state that seeds the 3-D far field. For the densimetric Froude number "
     "Fr = u_d / √(g′₀ d) of a 60° jet:")
eq("S_return = S_amb + (S₀ − S_amb) / (1.6 Fr) ,   x_return = 2.4 d Fr ,   z_rise = 2.2 d Fr")
para("with reduced gravity g′₀ = g (ρ_amb − ρ_brine) / ρ_amb. This is the source term Q_S in "
     "eq. (4) and the corresponding momentum/temperature seeds.")
para("Alternatives (same boundary role): a VISJET/JETLAG-class Lagrangian integral jet adds "
     "ambient crossflow + stratification (nearfield_model=\"lagrangian\"); or the nozzle is left "
     "RESOLVED (near_field_coupling=False) and the jet is computed directly on an auto-sized, "
     "auto-refined two-way grid nest about the port (run_resolved_nearfield), so the resolved "
     "near field feeds back onto the far field instead of being prescribed by Q_S.")

# ----------------------------------------------------------------------
h("12.  Surface-wave momentum coupling (optional, H4)", 1)
para("Monochromatic surface waves contribute a depth-decaying Stokes drift that advects all "
     "fields (Craik–Leibovich), and a radiation-stress gradient body force from shoaling:")
eq("u_s(z) = u_s0 e^{2kz} ê_w ,   u_s0 = g_st ω k a² ,   a = H_s/2 ,   ω = 2π/T_w ,   k = ω²/g")
eq("Transport velocity:  u_L = u + u_s ;     "
   "F_rad = − (1 / (ρ₀ H)) ∇ · S_xx ,   S_xx = E (2n − ½) ,   E = ½ ρ₀ g a²")

# ----------------------------------------------------------------------
h("13.  Initial, boundary, and closure conditions", 1)
bullet("Inflow (−x): prescribed ambient current U_in(t) = U_c + U_tide sin(2πt/T_M2), carrying ambient S, T.")
bullet("Outflow (+x): open zero-gradient, or Sommerfeld/Orlanski radiation ∂φ/∂t + c ∂φ/∂x = 0 (H6).")
bullet("Lateral (±y): free-slip, no-flux walls (optional sponge).")
bullet("Seabed (z = −H): no-flux scalars; free-slip or log-law wall function + quadratic drag (optional).")
bullet("Surface (z = η): wind stress, free-surface kinematics (§10), no scalar flux.")
bullet("Ambient profiles: linear S_amb(z), T_amb(z) between bed and surface values; ρ_amb from eq. (6).")

# ----------------------------------------------------------------------
h("14.  Time integration", 1)
para("The default is a first-order fractional step (advection + cross-fluxes explicit; diagonal "
     "dispersion / turbulent stress backward-Euler implicit; MAC projection). The optional "
     "SSP-RK2 (Heun) scheme (H1) raises the temporal order to second:")
eq("q⁽¹⁾ = Φ(qⁿ) ,   q⁽²⁾ = Φ(q⁽¹⁾) ,   qⁿ⁺¹ = ½ ( qⁿ + q⁽²⁾ )")
para("where Φ is one fractional-step update; the average of two divergence-free states "
     "is divergence-free.")

# ----------------------------------------------------------------------
h("15.  Symbol table", 1)
make_table(
    ["Symbol", "Meaning", "Symbol", "Meaning"],
    [["u = (u,v,w)", "velocity (m/s)", "p", "dynamic pressure (Pa)"],
     ["S", "absolute salinity (g/kg)", "T", "temperature (°C)"],
     ["ρ, ρ₀, ρₐ", "density / reference / ambient", "η", "free-surface elevation (m)"],
     ["ν, ν_t", "molecular / eddy viscosity", "D, D_T", "salt / heat dispersion tensor"],
     ["k, ε", "TKE / dissipation rate", "P_k, G_b", "shear / buoyancy production"],
     ["C_μ,C₁,C₂,C₃", "k–ε constants", "Sc_t, Pr_t", "turbulent Schmidt / Prandtl no."],
     ["α_T, β_S", "thermal exp. / haline contr.", "c_cab", "cabbeling coefficient"],
     ["β_S2, λ_TS", "haline / thermohaline curvature", "κ_p, γ_tb", "compressibility / thermobaric"],
     ["σ_So, d_f", "Soret / Dufour coefficients", "κ_os", "osmotic diffusivity"],
     ["ζ, τ, σ, ℓ", "OU forcing, time, intensity, scale", "f, Ω, β", "Coriolis / rotation / beta"],
     ["u_s, k, a", "Stokes drift / wavenumber / amp.", "S_xx, E", "radiation stress / wave energy"],
     ["Fr, g′₀, d", "Froude no. / reduced g / port dia.", "Δ, C_s, C_r", "grid scale / Smagorinsky / realiz."]],
)

para("")
para("Reference implementation: solver.py (NEREID-B Rev 2.0). These equations are the "
     "continuous model; the solver is their numerical realisation (3-D structured "
     "finite-volume, partial-cell bathymetry, fractional-step Chorin projection, TVD-MUSCL "
     "advection, implicit anisotropic dispersion, variable-coefficient pressure-Poisson — "
     "LU-factorised on the host, or solved on-device by CuPy PCG when run on a GPU).",
     size=9, italic=True, color=ACCENT)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "6", "model.docx")
os.makedirs(os.path.dirname(out), exist_ok=True)
DOC.save(out)
print("wrote", out, "with", _eqn[0], "numbered equations")
