#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_source6_docx.py — generates  6/source.docx

A COMPLETE record of every source used to build, validate and run the NEREID-B
project: the scientific literature behind each model component (near-field jet
scaling, turbulence closure, numerics, equation of state, waves, boundary
conditions), the real-site / field / regulatory data sources, the software
libraries and tools the code depends on, and the project data files — with, for
each, what it was used FOR and what it DID (its role / the value it provided).
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DEST = os.path.join(HERE, "6", "source.docx")
ACCENT = (0x0B, 0x3D, 0x5C)
BODY = "Calibri"

DOC = Document()
DOC.styles["Normal"].font.name = BODY
DOC.styles["Normal"].font.size = Pt(11)


def _bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def h(t, level=1):
    p = DOC.add_heading(t, level=level)
    for r in p.runs:
        r.font.name = BODY
        if level <= 1:
            r.font.color.rgb = RGBColor(*ACCENT)
    return p


def para(t="", bold=False, italic=False, size=11, align=None, color=None, space_after=6):
    p = DOC.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if t:
        r = p.add_run(t); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = BODY
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def ref(num, citation, used, did):
    """One source entry: bold number+citation, then 'Used for' and 'What it did'."""
    p = DOC.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"[{num}]  "); r.bold = True; r.font.name = BODY; r.font.size = Pt(11)
    r = p.add_run(citation); r.font.name = BODY; r.font.size = Pt(11)
    p2 = DOC.add_paragraph(style="List Bullet 2"); p2.paragraph_format.space_after = Pt(1)
    r = p2.add_run("Used for: "); r.bold = True; r.font.size = Pt(10); r.font.name = BODY
    r = p2.add_run(used); r.font.size = Pt(10); r.font.name = BODY
    p3 = DOC.add_paragraph(style="List Bullet 2"); p3.paragraph_format.space_after = Pt(8)
    r = p3.add_run("What it did: "); r.bold = True; r.font.size = Pt(10); r.font.name = BODY
    r = p3.add_run(did); r.font.size = Pt(10); r.font.name = BODY


def table(header, rows, font=9, widths=None):
    t = DOC.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, ht in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(str(ht)); r.bold = True; r.font.size = Pt(font); r.font.name = BODY
        _bg(c, "0B3D5C"); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(font); r.font.name = BODY
    if widths:
        for i, w in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)
    return t


# ==============================================================================
#  TITLE
# ==============================================================================
ttl = DOC.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run("NEREID-B — Complete Source Register"); r.bold = True
r.font.size = Pt(24); r.font.name = BODY; r.font.color.rgb = RGBColor(*ACCENT)
sub = DOC.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Every scientific, data, regulatory and software source used to build, "
                "validate and run the coupled stochastic-PDE brine-salinity solver")
r.italic = True; r.font.size = Pt(12.5); r.font.name = BODY
para("", space_after=8)
para("This register lists every source the project depends on, grouped by role: (A) near-field "
     "laboratory jet scaling, (B) far-field real-site & field datasets, (C) numerical-method & "
     "turbulence theory, (D) equation-of-state standard, (E) wave & boundary-condition theory, "
     "(F) regulatory criteria, (G) software libraries & tools, and (H) project data files. For "
     "each source it states the full citation/identifier, what it was USED FOR, and what it DID "
     "(its role in the code or the value it provided). Sections G–H cover the computational "
     "sources (libraries, datasets, decks) that the literature-only validation record omits.")

# ==============================================================================
#  A. NEAR-FIELD LABORATORY DENSE-JET SCALING
# ==============================================================================
h("A.  Near-field model — laboratory inclined-dense-jet scaling", 1)
para("The unresolvable nozzle is represented by established inclined-dense-jet correlations "
     "(they ARE the laboratory data) implemented in nearfield_jet(): terminal rise z_t = 2.2·D·Fr, "
     "return distance x_r = 2.4·D·Fr, return dilution S_r = 1.6·Fr.")
ref("1", "Roberts, P.J.W., Ferrier, A. & Daviero, G. (1997). “Mixing in Inclined Dense Jets.” "
    "J. Hydraulic Engineering 123(8): 693–699.",
    "Primary near-field scaling for 60° negatively-buoyant jets, implemented in nearfield_jet().",
    "Sets the rise/return/dilution laws; gives the validated rise-ratio band 2.1–2.8 (NEREID-B "
    "reproduces 2.20). For the Sydney case it produced Fr=24.3, rise 6.4 m, return 7.0 m, 38:1.")
ref("2", "Cipollina, A., Brucato, A., Grisafi, F. & Nicosia, S. (2005). “Bench-Scale "
    "Investigation of Inclined Dense Jets.” J. Hydraulic Engineering 131(11): 1017–1022.",
    "Independent laboratory confirmation of the inclined dense-jet rise/dilution scaling.",
    "Corroborated the 2.1–2.8 rise band the solver uses; bounds confidence in the near-field law.")
ref("3", "Lai, C.C.K. & Lee, J.H.W. (2012). “Mixing of inclined dense jets in stationary "
    "ambient.” J. Hydro-environment Research 6(1): 9–28.",
    "Supporting trajectory/dilution scaling across Froude numbers and nozzle angles.",
    "Used to bound the sin/cos angle factors applied for non-60° nozzles in nearfield_jet().")
ref("4", "Abessi, O. & Roberts, P.J.W. (2014). “Multiport Diffusers for Dense Discharges.” "
    "J. Hydraulic Engineering 140(8): 04014032.",
    "Quantitative multiport near-field dilution benchmark (the scaling the WA EPA Perth model "
    "adopts) and the multiport plume-merging reduction.",
    "Calibrated the merge_factor for n_ports>1 with finite spacing; Perth check S_i 27.7 vs "
    "NEREID-B 28.7 (~3.5%).")
ref("5", "Abessi, O. & Roberts, P.J.W. (2017). “Effect of Nozzle Orientation on Dense Jets in "
    "Stagnant Environments.” J. Hydraulic Engineering 143(...).",
    "Crossflow rise-lowering and downstream-shift behaviour for the Lagrangian near-field option.",
    "Set the rise-lowering / downstream-shift trends in nearfield_jet_lagrangian().")
ref("6", "Porto Pereira, N. et al. (2024). Frontiers in Marine Science 11:1377252.",
    "Calibration of the near-field crossflow dilution enhancement (no longer 'experimental').",
    "Calibrated the crossflow factor: reproduces Perth-60° ~7× at 1 m/s and ~1.0× at the "
    "operational 0.08 m/s, preserving the validated stagnant baseline.")

# ==============================================================================
#  B. FAR-FIELD REAL-SITE & FIELD DATASETS
# ==============================================================================
h("B.  Far-field real-site & field datasets", 1)
ref("7", "BMT / Oceanica for the Water Corporation of WA. “Perth Desalination Plant Discharge "
    "Modelling: Model Validation.” Appendix D, PSDP2 referral, WA EPA.",
    "Authoritative diffuser specs (45 GL/yr, 45% recovery, ~163 m double-tee, 40×0.13 m ports "
    "at 60°, 2.51 m³/s, 61.4 into 36.5) and the 45:1@50 m target; the primary far-field benchmark.",
    "Defined perth_case_input.json and the far-field conservative validation: model ~34.6:1 vs "
    "45:1 at 50 m (ratio 0.77, conservative = over-predicts impact = safe).")
ref("8", "Water Corporation of Western Australia — Perth Seawater Desalination Plant (operational "
    "/ environmental information).",
    "Plant capacity, recovery and discharge context.",
    "Cross-checked the Perth diffuser configuration (45 GL/yr; ambient ≈36.5; reject ≈61–65).")
ref("9", "Marti, C.L., Antenucci, J.P., Luketina, D., Okely, P. & Imberger, J. (2011). "
    "“Near-Field Dilution Characteristics of a Negatively Buoyant Hypersaline Jet Generated by "
    "a Desalination Plant.” J. Hydraulic Engineering 137(1): 57–65.",
    "In-situ field validation of the Perth near-field dilution.",
    "Confirmed the near-field correlation against measured Perth data.")
ref("10", "Gacia, E., Invers, O., Manzanera, M., Ballesteros, E. & Romero, J. (2007). “Impact of "
    "the brine from a desalination plant on a shallow seagrass (Posidonia oceanica) meadow.” "
    "Estuarine, Coastal and Shelf Science 72(4): 579–590. doi:10.1016/j.ecss.2006.11.021.",
    "Measured shallow far-field ΔS transect (5.0/2.5/1.0 ppt at 10/20/30 m); the file "
    "gacia2007_ctd.csv used with --calibrate-ctd.",
    "Defined the OUT-OF-ENVELOPE boundary (shallow surface discharge): model decay ~23 m vs "
    "observed ~12 m — a documented class mismatch, reported transparently.")
ref("11", "Fernández-Torquemada, Y., Sánchez-Lizaso, J.L. & González-Correa, J.M. (2005). "
    "“Monitoring of the brine discharge of the SWRO plant of Alicante (SE Spain).” Desalination "
    "182: 395–402.",
    "Mediterranean SWRO reject salinity and ambient cross-check values.",
    "Provided contrast values (reject ≈68; ambient ≈37.5).")
ref("12", "Sydney Desalination Plant (Kurnell, NSW) — public design basis (250 ML/day capacity, "
    "~47% recovery, offshore Tasman-shelf diffuser ~25 m depth, tunnelled multiport risers).",
    "The industrial case study modelled in folder 6 (sydney_case_input.json).",
    "Set the public parameters; per-port geometry is a clearly-labelled representative "
    "engineering configuration (concentrate ~3.26 m³/s, ~67 g/kg from the recovery mass balance).")

# ==============================================================================
#  C. NUMERICAL METHOD & TURBULENCE THEORY
# ==============================================================================
h("C.  Numerical method & turbulence-closure theory", 1)
ref("13", "Chorin, A.J. (1968). “Numerical solution of the Navier–Stokes equations.” "
    "Mathematics of Computation 22(104): 745–762.",
    "The fractional-step / projection method for the incompressible pressure–velocity coupling.",
    "Underlies the per-step MAC projection that enforces ∇·u = 0 to machine precision "
    "(divergence ~1e-16 in every run).")
ref("14", "van Leer, B. (1979). “Towards the Ultimate Conservative Difference Scheme V. "
    "A Second-Order Sequel to Godunov's Method.” J. Computational Physics 32(1): 101–136.",
    "The TVD-MUSCL slope-limited advection scheme for the scalar transport.",
    "Made salinity/temperature advection monotone and non-amplifying (selftest TVD check PASS), "
    "preventing spurious over/undershoot of the brine field.")
ref("15", "Benjamin, T.B. (1968). “Gravity currents and related phenomena.” J. Fluid Mechanics "
    "31(2): 209–248.",
    "The classical dense lock-exchange gravity-current front-Froude benchmark.",
    "Independently validated the PDE CORE (no near-field correlation): --benchmark gives "
    "Fr_f ≈ 0.47 vs the textbook ~0.5.")
ref("16", "Durbin, P.A. (1996). “On the k–ε stagnation point anomaly.” Int. J. Heat and Fluid "
    "Flow 17(1): 89–90.",
    "The realizable time-scale limiter for the k–ε eddy viscosity.",
    "Stopped the eddy viscosity over-producing/railing on any grid (nut@cap 0% in the runs), "
    "making the turbulence physical and grid-independent.")
ref("17", "Smagorinsky, J. (1963). “General circulation experiments with the primitive "
    "equations.” Monthly Weather Review 91(3): 99–164.",
    "The Smagorinsky sub-grid dissipation floor on the eddy viscosity.",
    "Guarantees a grid-scale dissipation ν_t ≥ (C_s Δ)²|S|, preventing energy pile-up at the "
    "grid scale.")
ref("18", "Nicoud, F. & Ducros, F. (1999). “Subgrid-Scale Stress Modelling Based on the Square "
    "of the Velocity Gradient Tensor.” Flow, Turbulence and Combustion 62(3): 183–200.",
    "The WALE LES closure (les_mode=\"wale\").",
    "Optional higher-fidelity sub-grid viscosity with correct near-wall cubic scaling and zero "
    "ν_t in pure shear, for fine/jet/GPU grids.")
ref("19", "Menter, F.R. (1994). “Two-equation eddy-viscosity turbulence models for engineering "
    "applications.” AIAA Journal 32(8): 1598–1605.",
    "The k–ε production realizability (Pk ≤ limiter·ε) safeguard.",
    "Available production limiter (pk_limiter); tested and documented as not the nut-railing fix "
    "(the railing was buoyancy-driven, fixed by the corrected G_b sign).")

# ==============================================================================
#  D. EQUATION-OF-STATE STANDARD
# ==============================================================================
h("D.  Equation-of-state standard", 1)
ref("20", "IOC, SCOR & IAPSO (2010). “The international thermodynamic equation of seawater – "
    "2010 (TEOS-10).” UNESCO; McDougall, T.J. & Barker, P.M. (2011), GSW Oceanographic Toolbox.",
    "The official seawater density equation of state (eos_mode=\"teos10\"), via the GSW library.",
    "Provided the standards-grade nonlinear density ρ(S,T,p) used to close the buoyancy term "
    "when teos10 is selected (selftest 13/13); falls back to the built-in nonlinear EOS if GSW "
    "is absent.")

# ==============================================================================
#  E. WAVE & BOUNDARY-CONDITION THEORY
# ==============================================================================
h("E.  Wave & boundary-condition theory", 1)
ref("21", "Craik, A.D.D. & Leibovich, S. (1976). “A rational model for Langmuir circulations.” "
    "J. Fluid Mechanics 73(3): 401–426.",
    "The Stokes-drift (Craik–Leibovich) wave–current advection.",
    "Added the depth-decaying Stokes transport velocity u_L = u + u_s to all field advection "
    "(optional H4 surface-wave coupling).")
ref("22", "Orlanski, I. (1976). “A simple boundary condition for unbounded hyperbolic flows.” "
    "J. Computational Physics 21(3): 251–269; with the Sommerfeld radiation condition.",
    "The radiative (non-reflecting) outflow boundary at +x.",
    "Let the plume/current leave the +x boundary without spurious reflection (optional "
    "orlanski_bc); the η relaxation is the radiation surrogate when it is off.")

# ==============================================================================
#  F. REGULATORY CRITERIA
# ==============================================================================
h("F.  Regulatory mixing-zone criteria", 1)
ref("23", "Western Australian EPA — Perth / Cockburn Sound brine-discharge licence criteria.",
    "The Perth compliance benchmark.",
    "ΔS < 1.2 ppt within 50 m and < 0.8 ppt within 1000 m of the diffuser.")
ref("24", "NSW mixing-zone practice (POEO Act / EPL framework) — representative ~1 ppt boundary "
    "limit; the folder-6 case uses a more protective 0.5 g/kg sub-lethal assessment contour.",
    "The Sydney case assessment threshold (cfg.dS_crit).",
    "Set dS_crit = 0.5 g/kg, against which the seabed footprint (~5128 m²) and the 50 m ΔS "
    "(~0.67 g/kg) are assessed — deliberately conservative.")

# ==============================================================================
#  G. SOFTWARE LIBRARIES & TOOLS
# ==============================================================================
h("G.  Software libraries & tools (computational dependencies)", 1)
para("The solver and its document/figure pipeline depend on the following packages. Versions are "
     "PINNED where they matter for correctness (see the table note).")
table(["Package / tool", "Used for", "What it did in the project"],
      [["NumPy 1.26.4 (PINNED)", "core array backend",
        "all vectorised field kernels (advection, dispersion, k–ε, EOS). PIN CRITICAL: NumPy 2.x "
        "breaks SciPy 1.11.4's sparse import and kills the pressure solver."],
       ["SciPy 1.11.4", "sparse linear algebra",
        "assembled and LU-factorised (splu) the variable-coefficient pressure-Poisson operator — "
        "the projection's host linear solve; also colored-noise generation."],
       ["Matplotlib", "plotting",
        "rendered the solver fig_*.png suite and the simu-dossier line graphs / charts / maps "
        "(plot_*.png)."],
       ["GSW 3.6.22 (optional)", "TEOS-10 EOS",
        "provided the official seawater density when eos_mode=\"teos10\"; built-in nonlinear EOS "
        "used if absent."],
       ["CuPy (cupy-cuda12x, optional)", "GPU backend",
        "ran the per-step field kernels on the GPU and, in Rev 2.0, the on-device CG pressure "
        "solve (--gpu / --gpu-poisson); verified CPU-equivalent on a Colab T4."],
       ["Numba 0.65.1 (optional)", "JIT acceleration",
        "opt-in JIT+threaded Thomas solver (NEREID_NUMBA=1); measured NOT a net speedup (the "
        "NumPy path is already vectorised) — kept off by default."],
       ["python-docx 1.2.0", "report generation",
        "generated 6/model.docx, case_study.docx, simu.docx and this source.docx."],
       ["pypdf", "PDF mining",
        "extracted diffuser parameters from the WA EPA PDF report."]])
para("Python 3 standard library (argparse, csv, json, logging, math, multiprocessing) provided "
     "the CLI, IO, logging and the ensemble multiprocessing pool.", size=9.5, italic=True)

# ==============================================================================
#  H. PROJECT DATA FILES
# ==============================================================================
h("H.  Project data files & input decks", 1)
table(["File", "Used for", "What it did / provided"],
      [["6/sydney_case_input.json", "Sydney case input deck",
        "all discharge/ambient/met-ocean/numerical inputs for the folder-6 prediction + honest "
        "provenance fields."],
       ["5/perth_case_input.json", "Perth validation deck",
        "the field-validated Perth diffuser configuration (primary far-field benchmark)."],
       ["gacia2007_ctd.csv", "real CTD transect",
        "shallow Mediterranean ΔS transect for --calibrate-ctd; defined the out-of-envelope "
        "boundary."],
       ["perth_deep_ctd.csv", "real CTD (deep, field-validated)",
        "deep-diffuser field-validated dilution used as the default calibration site."],
       ["FIELD_SITES['pacific_ctd2023']", "in-code real survey",
        "an independent SE-Pacific SWRO CTD+ADCP far-field check embedded in the solver."],
       ["6/metrics_summary.json + CSVs + .npz", "solver outputs",
        "the prediction results consumed by the case-study and simulation-output documents."]])

# ==============================================================================
#  I. CONSOLIDATED REFERENCE LIST
# ==============================================================================
h("I.  Reference list", 1)
refs = [
    "Roberts, P.J.W., Ferrier, A. & Daviero, G. (1997). Mixing in Inclined Dense Jets. J. Hydraulic Eng. 123(8): 693–699.",
    "Cipollina, A., Brucato, A., Grisafi, F. & Nicosia, S. (2005). Bench-Scale Investigation of Inclined Dense Jets. J. Hydraulic Eng. 131(11): 1017–1022.",
    "Lai, C.C.K. & Lee, J.H.W. (2012). Mixing of inclined dense jets in stationary ambient. J. Hydro-environment Research 6(1): 9–28.",
    "Abessi, O. & Roberts, P.J.W. (2014). Multiport Diffusers for Dense Discharges. J. Hydraulic Eng. 140(8): 04014032.",
    "Abessi, O. & Roberts, P.J.W. (2017). Effect of Nozzle Orientation on Dense Jets in Stagnant Environments. J. Hydraulic Eng.",
    "Porto Pereira, N. et al. (2024). Frontiers in Marine Science 11:1377252.",
    "BMT/Oceanica for Water Corporation of WA. Perth Desalination Plant Discharge Modelling: Model Validation. Appendix D, PSDP2 referral, WA EPA.",
    "Water Corporation of WA. Perth Seawater Desalination Plant. watercorporation.com.au/our-water/desalination.",
    "Marti, C.L. et al. (2011). Near-Field Dilution Characteristics of a Negatively Buoyant Hypersaline Jet. J. Hydraulic Eng. 137(1): 57–65.",
    "Gacia, E. et al. (2007). Impact of brine from a desalination plant on a shallow Posidonia oceanica meadow. Est. Coastal Shelf Sci. 72(4): 579–590. doi:10.1016/j.ecss.2006.11.021.",
    "Fernández-Torquemada, Y. et al. (2005). Monitoring of the brine discharge of the SWRO plant of Alicante. Desalination 182: 395–402.",
    "Sydney Desalination Plant (Kurnell, NSW) — public design basis.",
    "Chorin, A.J. (1968). Numerical solution of the Navier–Stokes equations. Math. Comp. 22(104): 745–762.",
    "van Leer, B. (1979). Towards the Ultimate Conservative Difference Scheme V. J. Comp. Phys. 32(1): 101–136.",
    "Benjamin, T.B. (1968). Gravity currents and related phenomena. J. Fluid Mech. 31(2): 209–248.",
    "Durbin, P.A. (1996). On the k–ε stagnation point anomaly. Int. J. Heat Fluid Flow 17(1): 89–90.",
    "Smagorinsky, J. (1963). General circulation experiments with the primitive equations. Mon. Weather Rev. 91(3): 99–164.",
    "Nicoud, F. & Ducros, F. (1999). Subgrid-Scale Stress Modelling Based on the Square of the Velocity Gradient Tensor. Flow Turb. Combust. 62(3): 183–200.",
    "Menter, F.R. (1994). Two-equation eddy-viscosity turbulence models for engineering applications. AIAA J. 32(8): 1598–1605.",
    "IOC, SCOR & IAPSO (2010). TEOS-10. UNESCO; McDougall, T.J. & Barker, P.M. (2011). GSW Oceanographic Toolbox.",
    "Craik, A.D.D. & Leibovich, S. (1976). A rational model for Langmuir circulations. J. Fluid Mech. 73(3): 401–426.",
    "Orlanski, I. (1976). A simple boundary condition for unbounded hyperbolic flows. J. Comp. Phys. 21(3): 251–269.",
    "Western Australian EPA. Perth / Cockburn Sound brine-discharge licence criteria.",
    "NSW EPA — POEO Act / Environment Protection Licence mixing-zone framework.",
]
for i, c in enumerate(refs, 1):
    p = DOC.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"[{i}]  "); r.bold = True; r.font.name = BODY; r.font.size = Pt(9.5)
    r = p.add_run(c); r.font.name = BODY; r.font.size = Pt(9.5)

para("")
para("Software citation note: NumPy (Harris et al. 2020, Nature 585), SciPy (Virtanen et al. "
     "2020, Nature Methods 17), Matplotlib (Hunter 2007, CSE 9(3)). Companion documents: "
     "6/model.docx (governing PDEs), 6/case_study.docx (interpretation), 6/simu.docx (outputs).",
     size=9, italic=True, color=ACCENT)

os.makedirs(os.path.dirname(DEST), exist_ok=True)
DOC.save(DEST)
print("wrote", DEST)
