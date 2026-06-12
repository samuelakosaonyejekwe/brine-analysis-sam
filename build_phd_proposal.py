# -*- coding: utf-8 -*-
"""
build_phd_proposal.py — generate a standard PhD research proposal for the
NEREID-B brine-salinity modelling project -> phd_proposal.docx (main folder).
Draws on the consolidated combo_analysis.docx content and embeds two
preliminary-results figures from the case study (folder 2/).
"""
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
D2 = os.path.join(HERE, "2")
BODY = "Calibri"
EQF = "Cambria Math"

DOC = Document()
DOC.styles["Normal"].font.name = BODY
DOC.styles["Normal"].font.size = Pt(11)


def _bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def h(t, lvl=1):
    return DOC.add_heading(t, level=lvl)


def para(t="", bold=False, italic=False, size=11, align=None, color=None, after=6):
    p = DOC.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if t:
        r = p.add_run(t); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = BODY
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def eq(t):
    p = DOC.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = EQF; r.font.size = Pt(12)
    return p


def bullet(t, sub=False):
    p = DOC.add_paragraph(style="List Bullet 2" if sub else "List Bullet")
    r = p.add_run(t); r.font.name = BODY; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def numbered(t):
    p = DOC.add_paragraph(style="List Number")
    r = p.add_run(t); r.font.name = BODY; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def table(header, rows, widths=None, fs=9, mono0=False):
    t = DOC.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htx in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(htx); run.bold = True; run.font.size = Pt(fs)
        run.font.name = BODY; _bg(c, "1F4E79"); run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(fs)
            run.font.name = EQF if (mono0 and i == 0) else BODY
    if widths:
        for i, w in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)


def figure(path, caption, width=5.6):
    if not os.path.exists(path):
        para(f"[figure not found: {path}]", italic=True); return
    DOC.add_picture(path, width=Inches(width))
    DOC.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = DOC.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9.5); r.font.name = BODY
    cp.paragraph_format.space_after = Pt(10)


# ======================================================================
#  TITLE PAGE
# ======================================================================
para("DOCTORAL RESEARCH PROPOSAL (PhD)", bold=True, size=13,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55, 0x55, 0x55), after=18)
t = DOC.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Coupled Stochastic Partial-Differential-Equation Modelling of "
              "Salinity Dispersion from Submarine Brine Outfalls")
r.bold = True; r.font.size = Pt(22); r.font.name = BODY; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
s = DOC.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Development, Validation and Application of the NEREID-B Framework "
              "for Near-Field and Far-Field Prediction of Desalination Brine Plumes")
r.italic = True; r.font.size = Pt(14); r.font.name = BODY
para("", after=20)
for label, val in [("Discipline", "Environmental Fluid Mechanics / Computational "
                    "Hydro-Environmental Engineering"),
                   ("Candidate", "[Candidate name]"),
                   ("Supervisory team", "[Principal supervisor] · [Co-supervisor]"),
                   ("Department / Institution", "[Department, University]"),
                   ("Proposed duration", "3–4 years (full-time)"),
                   ("Date", "12 June 2026")]:
    p = DOC.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    rr = p.add_run(f"{label}:  "); rr.bold = True; rr.font.size = Pt(11); rr.font.name = BODY
    rr2 = p.add_run(val); rr2.font.size = Pt(11); rr2.font.name = BODY
DOC.add_page_break()

# ======================================================================
#  ABSTRACT
# ======================================================================
h("Abstract", 1)
para(
    "Seawater desalination is expanding rapidly worldwide, and the hyper-saline "
    "reject (brine) it discharges into the sea is a growing environmental "
    "concern: dense brine sinks, pools on the seabed and can raise salinity "
    "above the tolerance of benthic ecosystems. Predicting where, how far and "
    "how deep the elevated-salinity field extends — and with what certainty, "
    "given the unpredictable sea — is essential for outfall design and "
    "regulatory consent, yet existing tools are limited: integral jet models "
    "(CORMIX/VISJET) collapse the plume to a centreline and cannot resolve the "
    "unsteady far field, while conventional Boussinesq CFD neglects the large "
    "brine/ambient density contrast and treats the sea deterministically.")
para(
    "This research proposes the development, rigorous validation and "
    "application of NEREID-B — a novel coupled, unsteady, non-Boussinesq, "
    "stochastic partial-differential-equation (PDE) framework that solves the "
    "velocity, pressure, density, salinity, temperature, turbulence and "
    "free-surface fields simultaneously, closed by a nonlinear equation of "
    "state, irreversible-thermodynamics cross-diffusion and an explicit "
    "osmotic/reverse-osmosis coupling, with the sea's unpredictability "
    "represented as intrinsic stochastic forcing so the model returns the "
    "probability distribution of the salinity field rather than a single "
    "deterministic answer. A validated near-field/far-field coupling embeds the "
    "unresolvable diffuser jet through calibrated correlations and seeds the "
    "3-D far field with the diluted plume. Substantial preliminary work — a "
    "complete model formulation, a robust solver, automated verification, "
    "near-field validation against laboratory scaling, grid-convergence of the "
    "far field, and two industrial case studies — establishes proof of "
    "concept. The doctoral programme will (i) extend the physics "
    "(non-Boussinesq continuity, multi-ion chemistry, higher-order numerics), "
    "(ii) deliver the first field validation of the coupled framework against "
    "CTD/ADCP surveys at an operating outfall, (iii) develop a machine-learning "
    "surrogate for rapid probabilistic screening, and (iv) release an "
    "open, validated decision-support tool. The expected contribution is a new, "
    "validated, uncertainty-aware predictive capability for brine-outfall "
    "salinity dispersion, of direct value to desalination engineering and "
    "marine environmental regulation.")

# ======================================================================
#  1. INTRODUCTION & BACKGROUND
# ======================================================================
h("1. Introduction and Background", 1)
para(
    "Global desalination capacity now exceeds 95 million m³/day and is growing "
    "to meet water scarcity intensified by climate change and population "
    "growth. Seawater reverse osmosis (SWRO), the dominant technology, recovers "
    "only ~40–50% of the feed as fresh water; the remainder is rejected as a "
    "brine roughly twice as saline as the receiving sea, frequently warmer and "
    "carrying residual process chemicals. This brine is overwhelmingly "
    "discharged back to the sea through submarine outfalls. Because it is "
    "denser than seawater it behaves as a negatively-buoyant effluent: it "
    "sinks, hugs the seabed and spreads as a gravity current, with documented "
    "impacts on seagrass, benthic invertebrates and other salinity-sensitive "
    "biota.")
para(
    "Regulators increasingly impose a maximum salinity increment (commonly +2 "
    "to +5 % of ambient, or a fixed limit such as +2 g kg⁻¹) at the edge of a "
    "permitted mixing zone. Demonstrating compliance requires predicting the "
    "three-dimensional, time-evolving salinity field around the diffuser — "
    "across both the energetic NEAR FIELD (the buoyant jet, the first tens of "
    "metres, where most dilution occurs) and the FAR FIELD (the bottom gravity "
    "current spreading over hundreds of metres under currents, tides and "
    "waves). The sea is never static or predictable, so a defensible prediction "
    "must also quantify the uncertainty arising from that variability.")

# ======================================================================
#  2. PROBLEM STATEMENT & RESEARCH GAP
# ======================================================================
h("2. Problem Statement and Research Gap", 1)
para("Current predictive practice rests on tools with complementary but "
     "fundamental limitations:")
bullet("Integral / entrainment models (CORMIX, VISJET, JETLAG) reproduce "
       "near-field jet dilution well but reduce the plume to a 1-D centreline; "
       "they cannot resolve the unsteady 3-D far field, bathymetric steering, "
       "stratification, or the full sea state.")
bullet("Boussinesq Reynolds-Averaged Navier–Stokes / CFD models resolve the 3-D "
       "field but assume a small density difference — invalid for hyper-saline "
       "reject where (ρ−ρ₀)/ρ₀ can exceed 0.04 — and require infeasible "
       "near-nozzle resolution.")
bullet("Both families treat the sea deterministically, omit osmotic and "
       "irreversible-thermodynamic couplings, and provide no calibrated measure "
       "of predictive uncertainty.")
para(
    "The research gap is therefore the absence of a single, validated, "
    "uncertainty-aware framework that couples an accurate near field to a "
    "non-Boussinesq, stochastic far field and reports the probability "
    "distribution of salinity. Addressing this gap is the purpose of this PhD.")

# ======================================================================
#  3. RESEARCH QUESTIONS & HYPOTHESES
# ======================================================================
h("3. Research Questions and Hypotheses", 1)
para("The programme is organised around four questions:")
numbered("Can a coupled stochastic PDE framework predict the 3-D, unsteady "
         "salinity distribution of a brine outfall more completely than "
         "existing tools, while remaining numerically robust and affordable?")
numbered("How accurately can a near-field/far-field coupling reproduce measured "
         "near-field dilution and far-field footprint at an operating outfall?")
numbered("How do near-field processes (jet merging at multiport diffusers) and "
         "far-field processes (gravity-current spreading, stratification, "
         "currents) interact to determine environmental impact and compliance?")
numbered("What is the predictive uncertainty of the salinity field given the "
         "unpredictability of currents, tides, wind and turbulence, and can it "
         "be quantified efficiently for design and consenting?")
para("Central hypothesis: a coupled non-Boussinesq stochastic PDE model, with a "
     "validated near-field closure, can predict brine-plume salinity "
     "distribution and its uncertainty with accuracy sufficient for "
     "engineering and regulatory use, outperforming single-paradigm tools.")

# ======================================================================
#  4. AIM & OBJECTIVES
# ======================================================================
h("4. Aim and Objectives", 1)
para("Aim: to develop, validate and apply a novel coupled stochastic PDE "
     "framework (NEREID-B) for predicting the near- and far-field salinity "
     "distribution and its uncertainty around submarine brine outfalls.")
table(["#", "Objective", "Status at proposal"],
      [["O1", "Formulate the coupled non-Boussinesq stochastic PDE system with "
        "osmotic and cross-diffusion couplings", "Completed (preliminary)"],
       ["O2", "Develop a robust finite-volume solver with a validated "
        "near-field/far-field coupling", "Completed (preliminary)"],
       ["O3", "Verify robustness and validate the near field against laboratory "
        "dense-jet scaling", "Completed (preliminary)"],
       ["O4", "Quantify predictive uncertainty via stochastic ensembles", "Demonstrated (preliminary)"],
       ["O5", "Extend the physics: non-Boussinesq continuity, multi-ion "
        "chemistry, higher-order numerics", "Proposed (WP2)"],
       ["O6", "Field-validate the far field against CTD/ADCP surveys at an "
        "operating outfall", "Proposed (WP3)"],
       ["O7", "Develop a machine-learning surrogate for rapid probabilistic "
        "screening", "Proposed (WP4)"],
       ["O8", "Apply to design/consenting case studies and release an open "
        "decision-support tool", "Proposed (WP5)"]],
      widths=[0.5, 4.2, 1.8])

# ======================================================================
#  5. ORIGINALITY & CONTRIBUTION
# ======================================================================
h("5. Originality and Intended Contribution", 1)
para("The work is original in combining, in one self-consistent and validated "
     "framework, elements not previously brought together for brine dispersion:")
bullet("Osmotic / reverse-osmosis coupling: the osmotic-pressure gradient of "
       "the hyper-saline reject enters the salt flux as a thermodynamically "
       "consistent term, absent from conventional plume models.")
bullet("Intrinsic stochastic forcing: currents, tides, wind and turbulence are "
       "Ornstein–Uhlenbeck colored-noise fields, so the model is a genuine "
       "stochastic PDE returning the probability density of salinity — a mean "
       "plume plus a quantified confidence envelope.")
bullet("Non-Boussinesq buoyancy with a nonlinear equation of state, valid for "
       "the large brine/ambient density contrast.")
bullet("Onsager (Soret/Dufour) cross-diffusion and a full anisotropic, "
       "state-dependent dispersion tensor.")
bullet("A validated near-field/far-field coupling that combines the accuracy of "
       "calibrated jet correlations with a grid-converged stochastic 3-D far "
       "field — and the first field validation of such a coupled framework.")

# ======================================================================
#  6. METHODOLOGY
# ======================================================================
DOC.add_page_break()
h("6. Methodology", 1)

h("6.1 The coupled stochastic PDE model", 2)
para("NEREID-B solves the coupled state vector q = (ρ, u, p, S, T, k, ε, η). "
     "The governing balances are:")
eq("∂ρ/∂t + ∇·(ρu) = 0")
eq("∂(ρu)/∂t + ∇·(ρu⊗u) = −∇p + ∇·τ + ρg − 2ρΩ×u + F_wave + F_osm + F_stoch")
eq("∂S/∂t + u·∇S = −(1/ρ)∇·J_S + R_S + ξ_S")
eq("∂T/∂t + u·∇T = −(1/ρc_p)∇·J_Q + Q_rad/(ρc_p) + ξ_T")
eq("ρ = ρ_EOS(S,T,p)   (nonlinear, TEOS-10)")
eq("∂ζ_m/∂t = −ζ_m/τ_m + σ_m 𝓛_m^{1/2} Ẇ_m   (stochastic forcing)")
para("with the salt flux J_S = −ρ D_S·∇S − ρ D_ST ∇T − (ρ D_S L_p/R_gT) ∇Π "
     "fusing anisotropic dispersion, Soret cross-diffusion and the osmotic "
     "flux, and a buoyancy-modified k–ε closure for turbulence. Solving the "
     "ensemble of stochastic realisations yields P[S(x,t)].")

h("6.2 Numerical solver and near-field/far-field coupling", 2)
para("The system is discretised by a MAC-consistent fractional-step projection "
     "on a terrain-following finite-volume grid with partial-cell bathymetry "
     "and an implicit free surface; scalars use a monotone TVD scheme. The "
     "unresolvable diffuser nozzle is represented by validated empirical "
     "dense-jet correlations, and the 3-D far field is seeded with the diluted "
     "return plume — a CORMIX/VISJET-class coupling that is accurate where each "
     "paradigm is valid.")

h("6.3 Validation strategy", 2)
bullet("Verification: automated invariant tests (conservation, boundedness, "
       "divergence control, monotonicity, bitwise-exact restart).")
bullet("Near-field validation: reproduction of published laboratory scaling for "
       "inclined dense jets (terminal rise, return-point dilution).")
bullet("Grid-convergence: systematic far-field mesh-refinement study.")
bullet("FIELD VALIDATION (proposed core contribution): CTD/ADCP/towed-sensor "
       "surveys at an operating SWRO outfall; comparison of predicted vs "
       "measured salinity profiles, footprint and dilution, with the predicted "
       "probability envelope tested for coverage.")

h("6.4 Case studies and uncertainty quantification", 2)
para("Idealised and real-site cases will exercise the framework across "
     "discharge configurations (single-port, multiport with jet merging), sea "
     "states and bathymetries. Predictive uncertainty will be quantified by "
     "Monte-Carlo stochastic ensembles and, in WP4, by a polynomial-chaos / "
     "machine-learning surrogate for rapid probabilistic screening.")

# ======================================================================
#  7. PRELIMINARY RESULTS
# ======================================================================
h("7. Preliminary Results (Proof of Concept)", 1)
para("Substantial preliminary work already establishes feasibility:")
bullet("Model and solver implemented and made numerically robust (passes a "
       "6-check automated self-test, including machine-level divergence control "
       "and bitwise-exact checkpoint/restart).")
bullet("Near field validated: the modelled 60° dense-jet rise z_t/(D·Fr) = 2.20 "
       "lies squarely within the published laboratory band 2.1–2.8, with "
       "return-point dilution following the established S_r ≈ 1.6 Fr scaling.")
bullet("Far field shown grid-converged: peak excess salinity changes by 0 % "
       "between the medium and fine meshes (4.44 → 4.45 g kg⁻¹ across a 4× "
       "cell-count range).")
bullet("Stochastic ensembles demonstrated, returning mean, spread and "
       "exceedance-probability maps of the salinity field.")
para("A 150,000 m³/day SWRO outfall case was simulated in two diffuser "
     "configurations. Key preliminary findings (Table 2) show that diffuser "
     "port spacing — through jet merging — is a first-order control on "
     "environmental impact: closely-spaced ports cut near-field dilution from "
     "~19× to ~8×, more than doubling the far-field peak excess and turning an "
     "essentially-compliant discharge into one with a ~17,800 m² regulatory "
     "mixing zone.")
para("Table 2 — Preliminary case results (single-port vs merged multiport diffuser)", italic=True)
table(["Quantity", "Single-port", "Merged multiport"],
      [["Near-field dilution", "19 ×", "8 ×"],
       ["Excess ΔS at seabed return", "1.5 g/kg", "3.7 g/kg"],
       ["Far-field peak excess", "2.13 g/kg", "4.45 g/kg"],
       ["Exceedance footprint (>2 g/kg)", "≈ 0 m²", "17,812 m²"],
       ["Maximum reach", "45 m", "167 m"],
       ["Far-field compliance", "compliant", "defined mixing zone"]],
      widths=[2.6, 1.7, 1.9])
figure(os.path.join(D2, "fig_nearfield_trajectory.png"),
       "Figure 1 (preliminary). Validated near-field inclined dense-jet "
       "trajectory; modelled rise and dilution match laboratory scaling.")
figure(os.path.join(D2, "fig_seabed_excess_map.png"),
       "Figure 2 (preliminary). Predicted seabed excess-salinity footprint for "
       "the merged multiport diffuser, with the regulatory mixing-zone contour.")

# ======================================================================
#  8. RESEARCH PLAN / WORK PACKAGES
# ======================================================================
DOC.add_page_break()
h("8. Research Programme and Work Plan", 1)
para("The programme is organised into six work packages (WP) over 3.5 years.")
table(["WP", "Title", "Key activities", "Output"],
      [["WP1", "Consolidation & literature synthesis",
        "Critical review; benchmark against CORMIX/VISJET; refine model "
        "formulation and self-test suite", "Review paper; verified baseline"],
       ["WP2", "Physics extension",
        "Non-Boussinesq variable-density projection; multi-ion (Maxwell–Stefan) "
        "chemistry; 2nd-order time integration; conservative free surface",
        "Extended solver; methods paper"],
       ["WP3", "Field validation (core)",
        "Design & execute CTD/ADCP/towed-sensor survey at an operating outfall; "
        "calibrate; validate near & far field and the uncertainty envelope",
        "Validation dataset; key paper"],
       ["WP4", "Uncertainty & ML surrogate",
        "Polynomial-chaos / neural surrogate for rapid probabilistic screening; "
        "sensitivity (Sobol) analysis", "Surrogate tool; UQ paper"],
       ["WP5", "Application & tool release",
        "Design/consenting case studies (multiport, stratified, energetic "
        "sites); open decision-support release; guidance", "Software; case papers"],
       ["WP6", "Synthesis & thesis",
        "Integration, dissemination, thesis writing", "PhD thesis; 4–6 papers"]],
      widths=[0.5, 1.5, 3.0, 1.5], fs=8.5)

para("Indicative timeline (●= active):", after=4)
table(["Work package", "Y1 H1", "Y1 H2", "Y2 H1", "Y2 H2", "Y3 H1", "Y3 H2", "Y4 H1"],
      [["WP1 Consolidation", "●", "●", "", "", "", "", ""],
       ["WP2 Physics extension", "", "●", "●", "●", "", "", ""],
       ["WP3 Field validation", "", "", "●", "●", "●", "", ""],
       ["WP4 UQ & surrogate", "", "", "", "●", "●", "●", ""],
       ["WP5 Application & tool", "", "", "", "", "●", "●", "●"],
       ["WP6 Synthesis & thesis", "", "", "", "", "", "●", "●"]],
      widths=[2.2, 0.6, 0.6, 0.6, 0.6, 0.6, 0.6, 0.6], fs=8.5)

# ======================================================================
#  9. EXPECTED OUTCOMES
# ======================================================================
h("9. Expected Outcomes, Deliverables and Publications", 1)
bullet("A validated, uncertainty-aware coupled-PDE framework for brine-outfall "
       "salinity prediction, with an extended (non-Boussinesq, multi-ion) "
       "physics core.")
bullet("The first field-validated near-field/far-field stochastic coupling for "
       "dense effluent dispersion, with a published validation dataset.")
bullet("A fast machine-learning surrogate enabling probabilistic compliance "
       "screening in design workflows.")
bullet("An open, documented decision-support tool and design guidance on "
       "diffuser configuration (notably port spacing and jet merging).")
bullet("4–6 peer-reviewed journal papers and a PhD thesis.")

# ======================================================================
#  10. SIGNIFICANCE & IMPACT
# ======================================================================
h("10. Significance and Impact", 1)
para("Scientifically, the project advances environmental fluid mechanics by "
     "delivering a validated coupled stochastic-PDE treatment of dense effluent "
     "dispersion with quantified uncertainty. Practically, it equips "
     "desalination engineers and marine regulators with a more complete, "
     "risk-based predictive capability — supporting outfall design that "
     "minimises ecological impact and evidence-based consenting. Given the rapid "
     "global growth of desalination and tightening marine-protection regimes, "
     "the impact is timely and international.")

# ======================================================================
#  11. RISK MANAGEMENT
# ======================================================================
h("11. Risk Management and Contingency", 1)
table(["Risk", "Likelihood / impact", "Mitigation"],
      [["Field-survey access or weather constraints (WP3)", "Med / High",
        "Multiple candidate sites; partner with plant operators early; use "
        "existing monitoring datasets as fallback"],
       ["Non-Boussinesq solver instability (WP2)", "Med / Med",
        "Incremental development behind the verified baseline; retain stable "
        "Boussinesq-buoyancy default"],
       ["Computational cost of ensembles", "Med / Med",
        "ML surrogate (WP4); parallelisation; reduced-order screening"],
       ["Merging-factor / closure uncertainty", "Med / Med",
        "Calibrate against lab and field data; report as bounded estimate"]],
      widths=[2.2, 1.3, 2.7], fs=8.5)

# ======================================================================
#  12. RESOURCES & ETHICS
# ======================================================================
h("12. Resources, Facilities and Ethical Considerations", 1)
para("Resources required: HPC allocation for ensemble simulations; field "
     "instrumentation (CTD, ADCP, towed conductivity sensor) and vessel time "
     "for WP3; partnership with a desalination plant operator and the relevant "
     "environmental regulator. No human or animal subjects are involved; field "
     "work will follow marine-survey permitting and environmental best practice, "
     "and data-sharing agreements will govern operator-provided data. Software "
     "will be released open-source under an appropriate licence.")

# ======================================================================
#  13. REFERENCES
# ======================================================================
h("13. Indicative References", 1)
for ref in [
    "Roberts, P.J.W., Ferrier, A. & Daviero, G. (1997). Mixing in inclined "
    "dense jets. J. Hydraulic Engineering, 123(8), 693–699.",
    "Cipollina, A., Brucato, A., Grisafi, F. & Nicosia, S. (2005). Bench-scale "
    "investigation of inclined dense jets. J. Hydraulic Engineering, 131(11).",
    "Lai, C.C.K. & Lee, J.H.W. (2012). Mixing of inclined dense jets in "
    "stationary ambient. J. Hydro-environment Research, 6(1), 9–28.",
    "Roberts, P.J.W. & Toms, G. (1987). Inclined dense jets in flowing current. "
    "J. Hydraulic Engineering, 113(3).",
    "Abessi, O. & Roberts, P.J.W. (2015). Effect of nozzle orientation on dense "
    "jets in stagnant environments. J. Hydraulic Engineering, 141(8).",
    "Jirka, G.H. (2004). Integral model for turbulent buoyant jets in unbounded "
    "stratified flows. Environmental Fluid Mechanics, 4, 1–56.",
    "Bleninger, T. & Jirka, G.H. (2008). Modelling and environmentally sound "
    "management of brine discharges from desalination plants. Desalination, 221.",
    "Fischer, H.B., List, E.J., Koh, R.C.Y., Imberger, J. & Brooks, N.H. "
    "(1979). Mixing in Inland and Coastal Waters. Academic Press.",
    "Missimer, T.M. & Maliva, R.G. (2018). Environmental issues in seawater "
    "reverse osmosis desalination: intakes and outfalls. Desalination, 434.",
    "IOC, SCOR & IAPSO (2010). The international thermodynamic equation of "
    "seawater – 2010 (TEOS-10). UNESCO.",
    "Roberts, D.A. et al. (2010). Impacts of desalination plant discharges on "
    "the marine environment: A review. Water Research, 44(18).",
]:
    p = DOC.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(18); p.paragraph_format.first_line_indent = Pt(-18)
    r = p.add_run(ref); r.font.size = Pt(9.5); r.font.name = BODY

DOC.save(os.path.join(HERE, "phd_proposal.docx"))
print("Saved phd_proposal.docx")
print("paragraphs:", len(DOC.paragraphs), "tables:", len(DOC.tables))
