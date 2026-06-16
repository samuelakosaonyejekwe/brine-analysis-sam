# -*- coding: utf-8 -*-
"""
build_res_paper_slugs.py — assembles ~/res_paper_slugs.docx, a comprehensive
research paper on coupled slug–hydrate prediction in subsea pipelines, synthesised
from: phd_proposal.docx / work.docx (theory), explain3.docx (engineering case &
predictions), explain4.docx (interpreted outputs), further_output.docx (cross-
scenario CSV outputs). Embeds every generated output (charts, maps, curves, graphs)
with labelled captions and reconstructs the data tables. True .docx, saved to the
local home directory.
"""

import os, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HOME = os.path.expanduser("~")
FIG = "/home/akosa/salinity_prediction/slug_figs"
DEST = os.path.join(HOME, "res_paper_slugs.docx")
T = json.load(open("/tmp/slug_tables.json"))

DOC = Document()
BODY = "Calibri"
DOC.styles["Normal"].font.name = BODY
DOC.styles["Normal"].font.size = Pt(11)
DOC.styles["Normal"].paragraph_format.space_after = Pt(6)


def _bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def h(t, level=1): return DOC.add_heading(t, level=level)


def p(t="", bold=False, italic=False, size=11, align=None, color=None, after=8, just=True):
    par = DOC.add_paragraph(); par.paragraph_format.space_after = Pt(after)
    if just and align is None: par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if align is not None: par.alignment = align
    if t:
        r = par.add_run(t); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = BODY
        if color: r.font.color.rgb = RGBColor(*color)
    return par


def eq(t):
    par = DOC.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(8); par.paragraph_format.space_before = Pt(4)
    r = par.add_run(t); r.font.name = "Cambria Math"; r.font.size = Pt(11); r.italic = True


def bullet(t):
    par = DOC.add_paragraph(style="List Bullet"); par.paragraph_format.space_after = Pt(2)
    r = par.add_run(t); r.font.name = BODY; r.font.size = Pt(11)


TABN = [0]
def render_table(rows, caption, font=8.5, widths=None):
    TABN[0] += 1
    t = DOC.add_table(rows=0, cols=len(rows[0])); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            r = cells[ci].paragraphs[0].add_run(str(val))
            r.font.size = Pt(font); r.font.name = BODY
            if ri == 0:
                r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); _bg(cells[ci], "1F4E79")
    if widths:
        for ci, w in enumerate(widths):
            for c in t.columns[ci].cells: c.width = Inches(w)
    cap = DOC.add_paragraph(); cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f"Table {TABN[0]}.  "); r.bold = True; r.font.size = Pt(9); r.font.name = BODY
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.name = BODY


FIGN = [0]
def figure(fname, caption, width=6.0):
    FIGN[0] += 1
    path = os.path.join(FIG, fname)
    if os.path.exists(path):
        DOC.add_picture(path, width=Inches(width))
        DOC.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p(f"[missing {fname}]", italic=True, color=(0xB0, 0, 0))
    cap = DOC.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f"Figure {FIGN[0]}.  "); r.bold = True; r.font.size = Pt(9.5); r.font.name = BODY
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9.5); r.font.name = BODY
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ==============================================================================
#  TITLE
# ==============================================================================
tt = DOC.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tt.add_run("Coupled Prediction of Hydrodynamic Slugging and Gas-Hydrate Formation "
               "in Subsea Multiphase Pipelines: A Stochastic Coupled-PDE Field Theory and "
               "its Application to a North Sea Tie-Back")
r.bold = True; r.font.size = Pt(17); r.font.name = BODY; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
st = DOC.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("The Slug–Hydrate Coupled Transport (SHCT) System and the Slug–Hydrate "
               "Coupling Number Φ_SH for Predictive Flow Assurance")
r.italic = True; r.font.size = Pt(12); r.font.name = BODY
a = DOC.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = a.add_run("Independent Research  ·  2026"); r.font.size = Pt(10.5); r.font.name = BODY
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ==============================================================================
#  ABSTRACT
# ==============================================================================
h("Abstract", 1)
p("Subsea tie-backs carry unprocessed oil, gas and water across tens of kilometres of cold "
  "(≈4–5 °C), high-pressure (100–400 bar) seabed, conditions under which two flow-assurance "
  "hazards dominate: hydrodynamic slugging — the spontaneous growth of liquid slugs from "
  "interfacial instability — and gas-hydrate formation — the crystallisation of ice-like "
  "clathrates that can plug a line within hours. Industry treats these phenomena with "
  "separate, loosely-linked tools and defends against them by continuous, heavily over-dosed "
  "chemical inhibition and oversized slug catchers. This paper argues that the genuinely "
  "unsolved, high-value problem is the two-way, multiscale, stochastic coupling between the "
  "two, and presents the Slug–Hydrate Coupled Transport (SHCT) system: a single set of "
  "coupled partial differential equations that simultaneously evolves the hyperbolic two-"
  "fluid hydrodynamic fields, a stochastic hydrate phase-field, and a sheared population "
  "balance for agglomeration and wall deposition, closed by an interfacial-area–subcooling "
  "operator and a new dimensionless invariant — the Slug–Hydrate Coupling Number Φ_SH — that "
  "flags exactly where and when the coupling turns dangerous. A transient SHCT solver is "
  "applied to a representative 18 km, 254 mm North Sea crude-oil tie-back. The model predicts "
  "that the line is hydrate-safe but slug-dominated at full rate (arrival 38 °C, Φ_SH = 0.70, "
  "0 % plugging), yet tips into a fully hydrate-critical regime at 50 % turndown (arrival "
  "5 °C, Φ_SH = 50, 100 % plugging at a P50 of 8 h over a 2.5–15.5 km band) and after shut-in "
  "(no-touch window ≈ 5.5 h, plug at ≈ 16 h). A thermodynamic-inhibitor sweep identifies the "
  "minimum mono-ethylene-glycol (MEG) dose that clears the risk. Every generated output — "
  "along-line profiles, space–time holdup and Φ_SH maps, pressure–temperature envelopes, "
  "probabilistic time-to-plug curves, deposit growth, engineering deliverables and cross-"
  "scenario comparisons — is presented and interpreted. The result is a physically coherent, "
  "uncertainty-aware predictor that converts blanket flow-assurance defence into targeted, "
  "ahead-of-time action.")
p("Keywords:  subsea flow assurance; hydrodynamic slugging; gas hydrates; two-fluid model; "
  "stochastic PDE; phase-field; coupling number; MEG inhibition; North Sea tie-back.",
  italic=True, size=10)

# ==============================================================================
#  1. INTRODUCTION
# ==============================================================================
h("1.  Introduction", 1)
p("A modern deep-water field rarely processes fluids at the well; the raw multiphase stream "
  "is carried back to a host platform through long subsea flowlines and a riser. Along that "
  "path the fluid meets a near-isothermal seabed at ≈4–5 °C, pressures of 100–400 bar, and a "
  "continuously falling pressure toward the host — exactly the conditions under which both "
  "slugging and hydrate formation are most aggressive (Figure 1). The financial stakes are "
  "large and asymmetric: a single hydrate-plug remediation can cost tens of millions of "
  "dollars and weeks of deferred production, while a slug-overwhelmed separator or a pressure-"
  "surge event is a safety hazard. The defensive response — continuous injection of "
  "thermodynamic inhibitor (methanol or MEG) sized for the worst conceivable case, plus "
  "oversized slug catchers — is itself a multi-hundred-million-dollar burden over field life. "
  "The industry therefore pays twice: once for the risk, and once for the conservative "
  "defence against it.")
figure("work_01.png",
       "Schematic of a subsea tie-back showing where hydrodynamic slugging and gas-hydrate "
       "formation co-occur along the cold, high-pressure flow path from the subsea manifold, "
       "across the undulating seabed, to the riser and host platform.", width=5.6)
p("The two hazards are not independent layers on the same pipe; each is a boundary condition "
  "for the other (Figure 2). Slug → hydrate: the gas–liquid–water interfacial area, the "
  "turbulent mixing intensity and the local cooling rate that set hydrate nucleation and "
  "growth are produced by the slug hydrodynamics. Hydrate → slug: the hydrate solid fraction "
  "changes the mixture density, sharply raises the effective viscosity, roughens and narrows "
  "the wall, and removes water and gas from the flowing phases — all of which feed back into "
  "the momentum and continuity balances that govern where and when the next slug forms. "
  "Because the feedback is two-way, multiscale and stochastic, neither phenomenon can be "
  "predicted correctly while the other is frozen — yet freezing one is precisely what every "
  "deployed tool does.")
figure("work_02.png",
       "The unsolved two-way slug–hydrate coupling. The forward (slug→hydrate) and return "
       "(hydrate→slug) paths form a closed feedback loop that no current model represents "
       "mechanistically.", width=5.2)
p("Existing approaches fall short for architectural reasons: hydrodynamics and hydrate "
  "thermodynamics live in different modules with one-way information flow; hydrate risk is "
  "reduced to an ‘inside/outside the envelope’ check that discards the transient, spatially-"
  "varying subcooling the slugs create; onset is forced to a deterministic correlation, "
  "erasing the probability distribution of induction time that risk decisions require; and "
  "models are run offline as design studies, not assimilated against live subsea data. This "
  "paper closes these gaps with a single coupled field theory (Section 2), realises it in a "
  "transient solver, and applies it to a worked North Sea case (Sections 3–7), reporting and "
  "interpreting every generated output.")

# ==============================================================================
#  2. THEORY
# ==============================================================================
h("2.  The SHCT coupled field theory", 1)
p("The artificial separation is collapsed into one mechanistic, stochastic, space–time field "
  "model — the Slug–Hydrate Coupled Transport (SHCT) system — in which the same conservation "
  "laws that spontaneously grow slugs also drive, and are driven by, hydrate nucleation, "
  "growth, agglomeration and deposition. A sufficiently resolved two-fluid model already "
  "captures slug formation as an instability of its own equations (the ‘slug-capturing’ "
  "property); embedding a stochastic hydrate phase-field and a sheared population balance in "
  "the same conservative system, and closing the loop in both directions, lets one set of "
  "equations predict both hazards and their interaction.")

h("2.1  Governing equations", 2)
p("With subscripts g (gas), l (liquid = oil+water) and h (hydrate solid), volume fractions "
  "α_k summing to unity, the one-dimensional along-pipe system is:")
p("Phase mass conservation with hydrate mass exchange:", bold=True, after=2)
eq("∂(α_k ρ_k)/∂t + ∂(α_k ρ_k u_k)/∂x = Γ_k ,   Σ_k Γ_k = 0")
p("where Γ_h = +ṁ_hyd is the hydrate formation rate and Γ_g, Γ_l the matched gas/water "
  "consumption. Phase momentum (slug-capturing) with hydrate back-coupling:", bold=False, after=2)
eq("∂(α_k ρ_k u_k)/∂t + ∂(α_k ρ_k u_k² )/∂x = −α_k ∂p/∂x − τ_k^w(α_h, δ_h) ± τ^i − α_k ρ_k g sinθ")
p("Crucially the wall-shear closure τ_k^w and the flow area both depend on the hydrate "
  "fraction α_h and deposit thickness δ_h — the hydrate→slug return path absent from current "
  "models. Mixture energy carries the latent heat of crystallisation L_hyd and seabed "
  "exchange (overall coefficient U). Hydrate order φ = α_h is evolved by a stochastic Allen–"
  "Cahn/Cahn–Hilliard phase-field driven by subcooling and the slug-generated interfacial "
  "area a_i:")
eq("∂φ/∂t + u_m ∂φ/∂x = M ∇²(δF/δφ) + k_g · a_i · ⟨ΔT_sub⟩₊ + ξ(x,t)")
p("ΔT_sub = T_eq(p) − T is the local subcooling; a_i is the instantaneous interfacial area "
  "produced by the slug field (the slug→hydrate forward path); and the multiplicative noise ξ, "
  "scaled by the classical nucleation rate (barrier ∝ ΔT_sub⁻²), injects the irreducible "
  "stochasticity of induction time. A sheared Smoluchowski population balance, modulated by "
  "the slug-driven shear and wall shear, selects flowable slurry versus wall deposit; a "
  "deposit-growth equation closes the rheological feedback by narrowing the bore. The result "
  "is a mixed hyperbolic–parabolic–stochastic system in which the hyperbolic core generates "
  "slugs, the stochastic phase-field generates hydrates, and the population/deposit equations "
  "decide their fate and close the loop.")

h("2.2  The Slug–Hydrate Coupling Number Φ_SH", 2)
p("Non-dimensionalising the coupled equations yields a single dimensionless invariant that "
  "measures the rate of hydrate growth driven by slug-generated interface against the rate at "
  "which the slug structure renews that interface:")
eq("Φ_SH = k_g · a_i · ΔT_sub^n / f_slug")
p("where f_slug is the local slug-passage frequency. When Φ_SH ≪ 1 the slugs renew the "
  "interface faster than hydrate can exploit it (slurry-tolerant, low plug risk); when "
  "Φ_SH ≳ 1 hydrate growth outpaces slug renewal and locks the interface into solid — the "
  "critical regime for deposition and plugging (Figure 3). A computed field Φ_SH(x,t) is "
  "therefore a compact early-warning scalar that flags the exact pipe locations and times "
  "where the coupling turns dangerous.")
figure("work_03.png",
       "Operational regime map of the Slug–Hydrate Coupling Number Φ_SH. The Φ_SH = 1 contour "
       "separates the slurry-tolerant region (low plug risk) from the plugging-prone region.",
       width=5.2)

h("2.3  Solution and prediction architecture", 2)
p("The hyperbolic core is discretised by a finite-volume method with an approximate-Riemann "
  "flux that resolves slug fronts without smearing; an implicit–explicit (IMEX) scheme treats "
  "the stiff hydrate kinetics and phase-field diffusion stably alongside explicit advection; "
  "the stochastic source is integrated by an Euler–Maruyama scheme so ensembles return the "
  "induction-time probability distribution; and the population balance is closed by the method "
  "of moments for real-time use. On top of the solver, a physics-informed neural operator "
  "provides a millisecond surrogate, adjoint-based 4D-variational assimilation reconstructs "
  "the hidden internal state from sparse subsea sensors, and a digital twin issues ahead-of-"
  "time alarms and computes the minimum, spatially-targeted inhibitor dose needed to drive "
  "Φ_SH below threshold (Figure 4).")
figure("work_04.png",
       "The SHCT coupled field theory feeding a real-time prediction, data-assimilation and "
       "closed-loop inhibitor-control architecture (digital twin).", width=5.6)

# ==============================================================================
#  3. ENGINEERING CASE
# ==============================================================================
h("3.  The engineering case: a North Sea crude-oil tie-back", 1)
p("The solver is applied to a representative 18 km, 254 mm (10 in) internal-diameter North "
  "Sea crude-oil tie-back carrying ≈11,412 bbl/d of oil at 30 % water cut from a subsea "
  "manifold, across an undulating ≈5 °C seabed, to a host platform via a riser. The pipeline, "
  "fluid and operating inputs are given in Tables 1–3.")
render_table(T["explain3"][0], "Pipeline and wall input data.", widths=[2.6, 2.4, 1.0])
render_table(T["explain3"][1], "Fluid property input data.", widths=[2.6, 2.4, 1.0])
render_table(T["explain3"][2], "Operating-condition input data.", widths=[2.6, 2.4, 1.0])

# ==============================================================================
#  4. METHOD
# ==============================================================================
h("4.  Method", 1)
p("The SHCT transient solver integrates the coupled PDEs (two-phase momentum and pressure, "
  "energy, hydrate phase-field and deposition) with a stochastic ensemble. The workflow was: "
  "(i) verify the installation (21/21 invariant checks pass); (ii) encode the case inputs as "
  "JSON; (iii) run the predictions for the normal, turndown and shut-in scenarios on both the "
  "drift-flux (implicit) and full two-fluid engines; (iv) read the outputs (along-line fields, "
  "monitor time-series, probabilistic summaries, engineering deliverables and charts); and "
  "(v) design the MEG mitigation by an inhibitor-concentration sweep. The runs are mass-"
  "conservative for all flowing cases (≈0 %); in production, the kinetic and coupling "
  "constants would first be calibrated to the field's measured arrival temperature, pressure "
  "drop and any observed onset.")

# ==============================================================================
#  5. RESULTS
# ==============================================================================
h("5.  Results and generated outputs", 1)
p("The single most important finding is a regime transition: the same pipeline that is a "
  "well-behaved (if slugging) warm line at full rate becomes a cold, stratified, hydrate-"
  "plugging line the moment throughput is reduced. Every output below tells a consistent part "
  "of that story. The cross-scenario summary is given in Table 4.")
render_table(T["explain4"][0], "Cross-scenario summary of the predicted state and risk.",
             widths=[1.7, 1.7, 1.7, 1.7])

h("5.1  Full production — a warm, slugging, hydrate-safe line", 2)
p("At full rate the mixture flows at a peak 3.8 m/s (below the 5.2 m/s API-14E erosional "
  "limit) with a 29 bar total pressure drop, and arrives warm at 38 °C. The along-line "
  "profiles (Figure 5) show liquid holdup climbing from 0.34 at the inlet to ≈0.8 by 15 km "
  "and then oscillating violently (0.45–1.0) over the last 2 km — severe riser-induced "
  "slugging — while the temperature stays far above the ≈12–13 °C hydrate-equilibrium line. "
  "The pressure–temperature trajectory (Figure 6) sits firmly on the warm side of the hydrate "
  "envelope, so the coupling number stays sub-critical (Φ_SH = 0.70) and the plugging "
  "probability is 0 %. The dominant integrity issue at full rate is therefore mechanical — "
  "slugging and cyclic loading — not hydrates (Table 5).")
figure("explain4_01.png",
       "Full-production along-line profiles (P50), versus distance along the 18 km line. "
       "Top→bottom: terrain elevation (m); liquid holdup α_l; pressure (blue, bar) and "
       "temperature (red solid, °C) against the hydrate-equilibrium temperature (red dashed); "
       "and subcooling ΔT_sub (°C). ΔT_sub is negative throughout — the warm line runs outside "
       "the hydrate region.", width=5.0)
figure("explain4_02.png",
       "Full-production pressure–temperature trajectory (blue) against the hydrate-equilibrium "
       "curve (red). Axes: temperature (°C) vs pressure (bar). The trajectory lies to the warm "
       "(right) side, outside the hydrate-stable region.", width=4.6)
render_table(T["explain3"][3], "Full-production engineering read-out.", widths=[2.8, 1.6, 1.2])

h("5.2  Turndown — the line tips into the hydrate regime", 2)
p("At 50 % rate the picture inverts. The reduced throughput cannot carry its heat to the "
  "host, so the line arrives at the 5 °C seabed temperature, and the low velocity drops it out "
  "of slug flow into a stratified, liquid-filled state (slug frequency collapses to "
  "≈0.0001 Hz). The transient holdup field (Figure 7) shows the line filling with liquid "
  "(α_l ≈ 0.9) along almost its entire length after ~5–8 h. With the slugs gone, nothing "
  "renews the interface or scours the wall, so wherever the line is also subcooled the "
  "coupling number explodes: the Φ_SH(x,t) map (Figure 8) develops a broad supercritical band "
  "(Φ_SH at its cap of 50) after the line cools. The probabilistic outputs (Figure 9) resolve "
  "the timing and extent: 100 % of realisations plug, between ~5.8 and 8.5 h, and the entire "
  "2.5–15.5 km cold mid-section is supercritical. The model predicts a peak wall deposit of "
  "117 mm (full bore — a plug) and a non-transportable slurry; the deposit builds toward "
  "blockage as shown in Figure 10. The probabilistic summary is given in Table 6.")
figure("explain4_03.png",
       "Turndown transient liquid-holdup field α_l(x,t). Axes: distance along the line (km) "
       "vs time (h); colour = liquid holdup. After ~5–8 h the line fills with liquid "
       "(α_l ≈ 0.9, yellow) — stratified accumulation, not slugging.", width=5.4)
figure("explain4_04.png",
       "Turndown coupling-criticality map Φ_SH(x,t). Axes: distance (km) vs time (h); colour = "
       "Φ_SH. After the line cools (~6 h) a broad red band (Φ_SH at its cap of 50) develops; "
       "the black Φ_SH = 1 contour bounds the plugging-prone zone.", width=5.4)
figure("explain4_05.png",
       "Turndown probabilistic outputs. Left: time-to-plug cumulative distribution function — "
       "100 % of realisations plug, between ~5.8 and 8.5 h. Right: maximum Φ_SH along the line "
       "(distance, km), showing the supercritical 2.5–15.5 km mid-section bounded by the warm "
       "inlet and the re-accelerating riser.", width=5.8)
figure("explain3_06.png",
       "Predicted wall-hydrate deposit thickness at the monitor point versus time, building "
       "toward bore blockage. (Note: the fixed monitor at ≈16.5 km lies in the riser, outside "
       "the 2.5–15.5 km critical band; the 117 mm plug forms upstream and is captured by the "
       "spatial-field outputs.)", width=4.8)
render_table(T["explain3"][4], "Turndown probabilistic risk metrics (P10/P50/P90).",
             widths=[2.4, 1.1, 1.1, 1.1])

h("5.3  Shut-in — the cooldown / no-touch window", 2)
p("On an unplanned shut-in the isolated line cools from its warm operating state toward the "
  "5 °C seabed. The monitor time-series (Figure 11) captures the classic sequence: the inlet "
  "rate is cut at 12 h; subcooling crosses zero (enters the hydrate region) ~5 h later and "
  "climbs to +8.8 °C; and Φ_SH spikes to critical once the line is both cold and stagnant. "
  "The key engineering output is the cooldown / no-touch time — ≈5.5 h — the window within "
  "which the operator must restart, depressurise or inhibit before the line enters the "
  "hydrate-forming region. Left un-remediated, the model predicts a plug at a P50 of ≈16 h. "
  "(The shut-in run conserves liquid mass only approximately at full stagnation (~26 % error), "
  "so its precise plug time is indicative; it correctly identifies shut-in as hydrate-"
  "critical and gives a sound no-touch estimate.)")
figure("explain4_06.png",
       "Shut-in transient at the monitor point versus time (h). Top: inlet rate cut at 12 h. "
       "Middle: subcooling ΔT_sub crossing zero ~5 h after the cut and climbing to +8.8 °C. "
       "Bottom: Φ_SH spiking to critical once the line is cold and stagnant.", width=5.2)

h("5.4  Engineering deliverables", 2)
p("The solver returns the design deliverables directly (Table 7). The warm full-rate line "
  "needs only ≈23 wt% MEG over a 2.2 km cold tail, whereas the turndown line needs ≈32 wt% "
  "over a ~16 km stretch; the full cross-scenario deliverable set is given in Table 8.")
render_table(T["explain4"][1], "Interpreted engineering deliverables (full production vs "
             "turndown).", widths=[1.7, 1.3, 1.3, 2.2])
render_table(T["further"][0], "Full engineering_deliverables.csv across all five scenarios "
             "(Normal, Two-fluid, Shut-in, Turndown, Baseline).", font=8, widths=[1.9, 0.8, 0.9, 0.9, 0.9, 0.9, 0.9])

h("5.5  Probabilistic outputs and MEG mitigation", 2)
p("Because the solver carries a stochastic-nucleation ensemble, every risk metric comes with "
  "a P10/P50/P90 band (Table 6). The interpretation is that the thermal and coupling state is "
  "essentially deterministic (the line WILL be cold and supercritical at turndown) and the "
  "only genuine uncertainty is WHEN the first hydrate nucleates — a 6–8 h band — which is "
  "exactly the right framing for a risk-based decision. The MEG sweep (Table 9) identifies the "
  "minimum dose: ~15 wt% removes the plugging probability and ~35 wt% fully clears the under-"
  "inhibited length, at a required injection of ≈7,100 L/h on the turndown rate.")
render_table(T["explain3"][5], "MEG inhibitor concentration sweep (turndown case): the dose "
             "that clears the risk.", widths=[1.4, 1.5, 1.7, 1.3, 1.2])

h("5.6  Engine cross-check", 2)
p("The same normal-production case was run on the drift-flux (implicit) and full two-fluid "
  "engines (Table 10), both mass-conservative. They agree on the qualitative picture — "
  "intermittent slug flow, marginal at full rate and clearly at risk on turndown/shut-in — but "
  "differ on absolute pressure drop and arrival temperature: the two-fluid engine, with two "
  "independent phase momenta, predicts a lower mixture velocity and therefore stronger cooling "
  "(a more conservative thermal estimate). This engine spread signals where calibration to "
  "field data matters most, and the deliberately conservative two-fluid thermal result "
  "reinforces the central recommendation to protect turndown and shut-in operations.")
render_table(T["explain3"][6], "Engine cross-check: drift-flux vs full two-fluid (normal "
             "production).", widths=[2.4, 1.6, 1.4, 1.0])

h("5.7  Cross-scenario graphical synthesis (all CSV outputs)", 2)
p("To complete the output catalogue, the solver's four machine-readable products for every "
  "run — the along-line spatial profile (fields_profile.csv), the monitor-point time history "
  "(timeseries_monitor.csv), the probabilistic summary (probabilistic_summary.csv) and the "
  "engineering deliverables (engineering_deliverables.csv) — are plotted across all five "
  "scenarios (Normal, Two-fluid, Shut-in, Turndown, Baseline), first overlaid for direct "
  "comparison (Figures 12–15) and then individually (Figures 16–20). All axes, series and "
  "units are labelled. These confirm, from a single consistent dataset, the regime picture "
  "built up above: the shut-in transient drives the largest subcooling and hydrate risk, "
  "turndown lowers velocity and holdup into the critical regime, and the two-fluid variant "
  "gives a higher peak velocity and pressure drop than the homogeneous baseline.")
figure("further_01.png",
       "Cross-scenario spatial profiles along the pipeline (fields_profile.csv): pressure, "
       "temperature, liquid holdup and Φ_SH versus distance, overlaid for all scenarios.",
       width=6.3)
figure("further_02.png",
       "Cross-scenario monitor-point time histories (timeseries_monitor.csv): subcooling, "
       "holdup, mixture velocity and Φ_SH versus time, overlaid for all scenarios.", width=6.3)
figure("further_03.png",
       "Cross-scenario probabilistic risk metrics (probabilistic_summary.csv): P50 of each "
       "metric per scenario; whiskers span the P10–P90 range.", width=6.3)
figure("further_04.png",
       "Cross-scenario engineering deliverables (engineering_deliverables.csv): headline "
       "design quantities compared across scenarios.", width=6.3)
figure("further_05.png",
       "Normal (design-rate) scenario detail: monitor-point time histories — pressure, "
       "temperature, subcooling, holdup and velocity (top row) — and spatial profiles — "
       "elevation, temperature vs hydrate T_eq, Φ_SH and wall deposit (bottom row).", width=6.3)
figure("further_06.png",
       "Two-fluid momentum-model variant: monitor-point time histories (top) and spatial "
       "profiles (bottom), as in Figure 16.", width=6.3)
figure("further_07.png",
       "Shut-in / cooldown scenario detail: monitor-point time histories (top) and spatial "
       "profiles (bottom).", width=6.3)
figure("further_08.png",
       "Turndown (reduced-rate) scenario detail: monitor-point time histories (top) and "
       "spatial profiles (bottom).", width=6.3)
figure("further_09.png",
       "Baseline scenario detail: monitor-point time histories (top) and spatial profiles "
       "(bottom).", width=6.3)

# ==============================================================================
#  6. DISCUSSION
# ==============================================================================
h("6.  Discussion — key physical insights", 1)
p("Regime switch. ", bold=True, after=2)
p("The root cause of the hazard is a flow-regime transition, not a gradual change. Turndown "
  "moves the line out of slug flow into stratified flow; losing the slugs removes the natural "
  "wall-scouring that keeps hydrate from locking on. This is precisely why Φ_SH jumps from "
  "0.7 to its cap — it is the quantitative expression of ‘the slugs stopped cleaning the "
  "wall’.")
p("Coupling and trade-off. ", bold=True, after=2)
p("Slugging and hydrates are coupled and trade off. At full rate the problem is mechanical "
  "(severe riser slugging); at low rate it is chemical/thermal (hydrate plugging). The coupled "
  "Φ_SH number is what connects them and locates the hand-over — the central value of treating "
  "the two phenomena in one field theory rather than two separate tools.")
p("Locality and operations. ", bold=True, after=2)
p("The hazard is spatially concentrated in a well-defined 2.5–15.5 km band, so protection "
  "(insulation, inhibitor injection, monitoring) can be targeted there rather than along the "
  "whole route. And it is operations, not steady design, that set the risk: full-rate steady "
  "operation is safe; the danger lives in turndown and shut-in, which makes minimum-flow "
  "limits, cooldown management and restart inhibition the controlling mitigations.")

# ==============================================================================
#  7. RECOMMENDATIONS
# ==============================================================================
h("7.  Engineering recommendations", 1)
for tx in [
    "Define and enforce a minimum stable production rate that keeps the line in slug/"
    "intermittent flow and warm enough to stay sub-critical (Φ_SH < 1); the turndown run "
    "shows 50 % is already unsafe.",
    "Size the topside slug catcher to the predicted P90 surge volume and check riser-base / "
    "support fatigue against the severe riser slugging seen at full rate.",
    "Establish a shut-in no-touch time of ≈5–6 h; beyond it, inhibit (≈35 wt% MEG to fully "
    "clear the line) or depressurise before restart.",
    "Target hydrate protection at the 2.5–15.5 km coupled hot-spot band rather than uniformly "
    "along the route.",
    "Before committing capital, calibrate the kinetic constants to the field's measured "
    "arrival temperature, pressure drop and any observed onset, and run a grid-independence "
    "check on the chosen terrain.",
]:
    bullet(tx)

# ==============================================================================
#  8. LIMITATIONS
# ==============================================================================
h("8.  Limitations and confidence", 1)
p("The worked case uses literature-typical kinetic and coupling constants: the methodology "
  "and the relative predictions (which operation, where, when, what dose) are robust and the "
  "numerics are verified and mass-conservative for the flowing cases, but the absolute "
  "magnitudes (exact time-to-plug, exact arrival temperature) should be calibrated to field "
  "data before operational use. The shut-in run is only qualitatively mass-conservative "
  "(~26 % error at full stagnation), so its precise plug time is indicative. The drift-flux "
  "and two-fluid engines agree on the qualitative risk but differ on absolute ΔP and arrival "
  "temperature (the two-fluid being more conservative), and the temperature field carries some "
  "first-order-scheme grid sensitivity on steep terrain — a grid-independence check is "
  "recommended for the production geometry. None of these change the engineering conclusion. "
  "The solver complements, and does not replace, a certified transient multiphase simulator "
  "for final sign-off.")

# ==============================================================================
#  9. CONCLUSIONS
# ==============================================================================
h("9.  Conclusions", 1)
for tx in [
    "The SHCT system provides, in a single conservative stochastic field theory, a mechanistic "
    "two-way coupling of hydrodynamic slugging and gas-hydrate formation, closed by an "
    "interfacial-area–subcooling operator and a new dimensionless invariant, the Slug–Hydrate "
    "Coupling Number Φ_SH.",
    "Applied to an 18 km North Sea crude-oil tie-back, the model predicts a hydrate-safe but "
    "slug-dominated line at full rate (Φ_SH = 0.70, 0 % plugging), tipping into a fully "
    "hydrate-critical regime at 50 % turndown (Φ_SH = 50, 100 % plugging, P50 ≈ 8 h over a "
    "2.5–15.5 km band) and after shut-in (no-touch ≈ 5.5 h, plug ≈ 16 h).",
    "Φ_SH localises the hazard in space and time and converts the coupled physics into a "
    "single operational early-warning scalar; the probabilistic outputs frame the certain "
    "hazard with a quantified timing band; and the MEG sweep returns the minimum dose that "
    "clears the risk.",
    "The framework therefore replaces conservative blanket inhibition and oversized slug "
    "catchers with predictive, spatially-targeted flow assurance — serving design, operations "
    "and real-time control from one mathematical core.",
]:
    bullet(tx)

# ==============================================================================
#  REFERENCES / FURTHER READING
# ==============================================================================
h("References and source material", 1)
p("This paper synthesises the author's independent research documents: the invention "
  "disclosure / white paper defining the SHCT governing equations, the Φ_SH coupling invariant "
  "and the assimilation architecture; the worked North Sea engineering case and its method; "
  "the detailed interpretation of the prediction outputs; and the cross-scenario graphical "
  "output of all solver CSV products. Supporting flow-assurance background includes the "
  "classical treatments of hydrodynamic slugging (Kelvin–Helmholtz interfacial instability and "
  "the two-fluid slug-capturing model), gas-hydrate thermodynamics and kinetics (classical "
  "nucleation theory and the pressure–temperature stability envelope), and standard subsea "
  "design references (e.g. the API-14E erosional-velocity limit and thermodynamic-inhibitor "
  "(MEG/methanol) hydrate-suppression practice).", size=10)

DOC.save(DEST)
print("wrote", DEST, "|", FIGN[0], "figures,", TABN[0], "tables")
