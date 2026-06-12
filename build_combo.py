# -*- coding: utf-8 -*-
"""
build_combo.py — combine the eight project .docx files into a single coherent
document, combo_analysis.docx, preserving all text, tables and images.

Method (library-free): import the full style set from the first source, then
deep-copy each source document's body elements into the master, copying every
embedded image into the master package (de-duplicated by content) and remapping
the relationship IDs so the images render. Each source becomes a numbered Part
with a divider and a page break.
"""
from copy import deepcopy
from io import BytesIO
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
BODY = "Calibri"

# logical order: model -> inputs -> outputs -> single-port (report+discussion)
# -> merged (report+discussion) -> comparison
PARTS = [
    ("I", "The Coupled-PDE Model — NEREID-B", "salinity.docx"),
    ("II", "Input Data Requirements", "input.docx"),
    ("III", "Output Specification", "output.docx"),
    ("IV", "Single-Port Case — Simulation Report", "2/explain_simu_singleport.docx"),
    ("V", "Single-Port Case — Results Interpretation", "2/discuss_result2.docx"),
    ("VI", "Merged Multi-Port Case — Simulation Report", "2/explain_simu.docx"),
    ("VII", "Merged Multi-Port Case — Results Interpretation", "2/discuss_result.docx"),
    ("VIII", "Comparative Analysis — Single-Port vs Merged", "2/discuss_result3.docx"),
]

master = Document()

# --- import the full style set from the first source so every style resolves
src0 = Document(os.path.join(HERE, PARTS[0][2]))
ms = master.styles.element
for c in list(ms):
    ms.remove(c)
for c in src0.styles.element:
    ms.append(deepcopy(c))


def append_document(master, path):
    """Deep-copy a source document's body into master, carrying its images."""
    sub = Document(path)
    rid_map = {}
    for rid, rel in list(sub.part.rels.items()):
        if "image" in rel.reltype and not rel.is_external:
            blob = rel.target_part.blob
            new_part = master.part.package.image_parts.get_or_add_image_part(BytesIO(blob))
            rid_map[rid] = master.part.relate_to(new_part, rel.reltype)
    body = master.element.body
    for child in sub.element.body.iterchildren():
        if child.tag == qn("w:sectPr"):
            continue                       # keep the master's section properties
        new = deepcopy(child)
        for e in new.iter():
            for attr in (qn("r:embed"), qn("r:link"), qn("r:id")):
                v = e.get(attr)
                if v is not None and v in rid_map:
                    e.set(attr, rid_map[v])
        body.append(new)


# ---------------- combined cover page ----------------
def center(text, size, bold=False, italic=False, color=None, after=6):
    p = master.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = BODY
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


center("Prediction and Analysis of Salinity Distribution", 26, bold=True,
       color=(0x1F, 0x4E, 0x79), after=2)
center("Around a Submarine Brine-Effluent Outfall", 20, bold=True,
       color=(0x1F, 0x4E, 0x79), after=10)
center("A Consolidated Report: Novel Coupled-PDE Model, Solver, "
       "Industrial Case Simulations and Comparative Analysis", 13, italic=True)
center("NEREID-B — Nonlinear Eulerian Reactive-osmotic Effluent "
       "Integro-Dispersion model for Brine outfalls", 11)
master.add_paragraph()
center("Consolidated Edition  ·  combines 8 project documents  ·  "
       "Rev. 1.0  ·  11 June 2026", 11, color=(0x55, 0x55, 0x55))

master.add_paragraph()
ch = master.add_paragraph(); ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ch.add_run("Contents"); r.bold = True; r.font.size = Pt(15); r.font.name = BODY
for num, title, _ in PARTS:
    p = master.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Part {num} — {title}")
    r.font.size = Pt(11.5); r.font.name = BODY
center("(Each Part is the corresponding source document, reproduced in full.)",
       9.5, italic=True, color=(0x66, 0x66, 0x66))

# ---------------- append each part ----------------
for num, title, rel in PARTS:
    path = os.path.join(HERE, rel)
    master.add_page_break()
    dp = master.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp.paragraph_format.space_before = Pt(60); dp.paragraph_format.space_after = Pt(4)
    r = dp.add_run(f"PART {num}"); r.bold = True; r.font.size = Pt(28)
    r.font.name = BODY; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    tp = master.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(10)
    r = tp.add_run(title); r.bold = True; r.font.size = Pt(16); r.font.name = BODY
    sp = master.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run(f"(source: {rel})"); r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88); r.font.name = BODY
    master.add_page_break()
    append_document(master, path)
    print(f"  appended Part {num}: {rel}")

out = os.path.join(HERE, "combo_analysis.docx")
master.save(out)
print("Saved", out)
imgs = [r for r in master.part.rels.values() if "image" in r.reltype]
print("paragraphs:", len(master.paragraphs), "tables:", len(master.tables),
      "images:", len(imgs))
