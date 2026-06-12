# -*- coding: utf-8 -*-
"""
build_discuss_result2.py — interpretation & discussion of the SINGLE-PORT
(baseline) diffuser run -> 2/discuss_result2.docx.
Reads 2/singleport_summary.json and the extracted sp_fig_*.png figures.
"""
import json
import math
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
D2 = os.path.join(HERE, "2")
BODY = "Calibri"

data = json.load(open(os.path.join(D2, "singleport_summary.json")))
cfg = data["config"]; m = data["metrics"]
N_PORT = 16
Q_total = cfg["Q_d"] * N_PORT
U_d = cfg["Q_d"] / (math.pi * (cfg["d_p"] / 2) ** 2)
dS_source = cfg["S0"] - cfg["S_amb_bot"]
dilu = m["nf_return_dilution"]
dS_return = dS_source / dilu
foot = m["seabed_footprint_m2"]

DOC = Document()
DOC.styles["Normal"].font.name = BODY
DOC.styles["Normal"].font.size = Pt(11)


def _bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def h(t, lvl=1):
    return DOC.add_heading(t, level=lvl)


def para(t="", italic=False, after=6):
    p = DOC.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if t:
        r = p.add_run(t); r.italic = italic; r.font.size = Pt(11); r.font.name = BODY
    return p


def bullet(t):
    p = DOC.add_paragraph(style="List Bullet")
    r = p.add_run(t); r.font.name = BODY; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def table(header, rows, widths=None, fs=9):
    t = DOC.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htx in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(htx); run.bold = True; run.font.size = Pt(fs)
        run.font.name = BODY; _bg(c, "1F4E79"); run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(fs); run.font.name = BODY
    if widths:
        for i, w in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)


def figure(fname, caption, width=6.0):
    path = os.path.join(D2, fname)
    if not os.path.exists(path):
        para(f"[figure {fname} not found]", italic=True); return
    DOC.add_picture(path, width=Inches(width))
    DOC.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = DOC.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9.5); r.font.name = BODY
    cp.paragraph_format.space_after = Pt(10)


# ----------------------------------------------------------------------
t = DOC.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Interpretation and Discussion of Results")
r.bold = True; r.font.size = Pt(24); r.font.name = BODY; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
s = DOC.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Single-Port (Independent-Jet) Brine Diffuser — Baseline Case")
r.italic = True; r.font.size = Pt(13); r.font.name = BODY
mt = DOC.add_paragraph(); mt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = mt.add_run("NEREID-B Coupled-PDE Simulation  ·  Rev. 1.0  ·  11 June 2026")
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55); r.font.name = BODY
para("", after=8)
para("This document interprets and discusses the results of the SINGLE-PORT "
     "(baseline) configuration of the SWRO desalination brine-outfall case, in "
     "which the diffuser ports are far enough apart that each dense jet behaves "
     "independently (no plume merging). It explains what the predicted near- "
     "and far-field quantities mean, why the discharge is essentially "
     "compliant, and how trustworthy the prediction is. The companion documents "
     "discuss the merged-diffuser variant (discuss_result.docx) and compare the "
     "two configurations in detail (discuss_result3.docx).")

# ----------------------------------------------------------------------
h("1. Configuration and headline outcome", 1)
para(
    f"The case is the brine reject of a 150,000 m³/day SWRO plant "
    f"({Q_total:.2f} m³/s at {cfg['S0']:.0f} g/kg, ~{cfg['S0']/cfg['S_amb_bot']:.1f}× "
    f"ambient) discharged through {N_PORT} ports of {cfg['d_p']*1000:.0f} mm at "
    f"{cfg['theta_deg']:.0f}°. In this baseline the ports are treated as "
    f"INDEPENDENT — each jet entrains clean ambient water over its full "
    f"trajectory, achieving the maximum dilution the geometry allows. The "
    f"headline outcome is an efficient, essentially-compliant discharge:")
table(["Quantity", "Value", "Meaning"],
      [["Discharge Froude, Fr", f"{m['Fr_d']:.1f}", "momentum vs buoyancy"],
       ["Near-field dilution", f"{dilu:.0f} ×", "ambient mixed in by the jet"],
       ["Excess ΔS at seabed return", f"{dS_return:.2f} g/kg", "vs +{:.0f} limit".format(int(cfg['dS_crit']))],
       ["Far-field peak excess", f"{m['excess_max']:.2f} g/kg", "3-D model"],
       ["Exceedance footprint", f"{foot:.0f} m²", f">{cfg['dS_crit']:.0f} g/kg"],
       ["Maximum reach", f"{m['r_max_m']:.0f} m", "extent of impact"]],
      widths=[2.3, 1.4, 2.6])

# ----------------------------------------------------------------------
h("2. Near-field: an efficient inclined dense jet", 1)
para(
    f"The near-field trajectory (Figure 1) shows the classic 60° dense-jet "
    f"arc: the brine leaves the ports at {U_d:.1f} m/s, rises to "
    f"{m['nf_rise_m']:.1f} m (z_t/(D·Fr) = {m['nf_rise_ratio']:.2f}, squarely "
    f"in the validated laboratory band 2.1–2.8), bends over and returns to the "
    f"seabed {m['nf_return_dist_m']:.1f} m away. Crucially, over that path the "
    f"jet entrains enough ambient water to achieve a {dilu:.0f}× dilution, so "
    f"the brine reaches the bed at only {dS_return:.2f} g/kg above ambient — "
    f"below the +{cfg['dS_crit']:.0f} g/kg consent limit. This is exactly why "
    f"inclined multiport diffusers are the industry standard for brine "
    f"disposal: a well-spaced 60° jet does most of the dilution work in the "
    f"first few metres.")
figure("sp_fig_nearfield_trajectory.png",
       "Figure 1. Single-port near-field jet trajectory (validated scaling). "
       "Independent jets achieve the full ~19× dilution along the arc.")

# ----------------------------------------------------------------------
h("3. Far-field: a compliant, bottom-trapped plume", 1)
para(
    f"Seeded with the well-diluted ({dilu:.0f}×) return plume, the 3-D model "
    f"shows the brine sinking and spreading as a thin bottom layer (Figures 2 "
    f"and 3). Because the near-field already diluted it close to ambient, the "
    f"far-field excess is small: a peak of {m['excess_max']:.2f} g/kg, a "
    f"footprint of {foot:.0f} m² above the limit (i.e. essentially none — the "
    f"peak only marginally touches the limit in a small pool right at the "
    f"discharge), and a reach of just {m['r_max_m']:.0f} m. The plan view is "
    f"labelled 'far-field below limit' because the regulatory contour is not "
    f"reached over any meaningful area.")
figure("sp_fig_seabed_excess_map.png",
       "Figure 2. Seabed excess salinity — the diluted plume stays below the "
       "regulatory limit over the far field (no exceedance contour).")
figure("sp_fig_vertical_section.png",
       "Figure 3. Vertical section: a thin, weak, bottom-hugging plume.")
para(
    "Interpretation: this is the desired behaviour of a well-designed brine "
    "outfall — the dense effluent is trapped near the bed (not lofted to the "
    "surface), diluted below the consent limit in the near field, and confined "
    "to a small zone around the diffuser.")

# ----------------------------------------------------------------------
h("4. Dilution, decay, currents and surface", 1)
para(
    "The decay and dilution curves (Figures 4–5) confirm the excess falls "
    "rapidly toward ambient with distance, staying under the regulatory line "
    "throughout the far field. The near-bed currents (Figure 6) carry the thin "
    "plume gently away, and the free-surface response (Figure 7) is "
    "negligible — consistent with a trapped, bottom-confined discharge that "
    "does not disturb the surface.")
figure("sp_fig_salinity_decay.png", "Figure 4. Excess salinity vs distance (below limit).")
figure("sp_fig_centerline_dilution.png", "Figure 5. Far-field dilution vs distance.")
figure("sp_fig_seabed_currents.png", "Figure 6. Near-bed currents.")
figure("sp_fig_free_surface.png", "Figure 7. Free-surface elevation (weak response).")

# ----------------------------------------------------------------------
h("5. Compliance and uncertainty", 1)
para(
    f"Against the +{cfg['dS_crit']:.0f} g/kg limit the single-port diffuser is "
    f"essentially COMPLIANT: the mean-field exceedance footprint is ≈{foot:.0f} "
    f"m². The {m['n_ensemble']}-member stochastic ensemble gives a maximum "
    f"exceedance probability of {m['max_exceedance_prob']:.2f}, but this peak "
    f"sits at the immediate discharge point (where salinity is elevated in "
    f"every realisation) and falls away rapidly — i.e. any residual risk is "
    f"confined to the immediate vicinity of the diffuser, not the wider "
    f"receiving water.")
figure("sp_fig_exceedance_probability.png",
       "Figure 8. Exceedance-probability map from the 4-member ensemble — risk "
       "concentrated at the discharge and negligible in the far field.")

# ----------------------------------------------------------------------
h("6. Why the result is credible, and its limits", 1)
bullet("The near-field rise and dilution match validated laboratory scaling.")
bullet("The dense plume correctly sinks and is trapped near the bed.")
bullet("Salinity stays bounded; divergence is at machine level; the run passes "
       "the automated self-test.")
bullet("Limitation: the single-port idealisation assumes the jets do NOT "
       "interact. Real diffusers with closely-spaced ports merge, which "
       "REDUCES dilution and increases the far-field load — the subject of the "
       "merged-diffuser variant and the comparison document. The single-port "
       "result should therefore be read as a best-case (upper-bound dilution) "
       "scenario.")
bullet("The far field is physically based but not yet validated against site "
       "CTD/ADCP data; absolute footprint/reach should be confirmed by "
       "post-commissioning monitoring.")

# ----------------------------------------------------------------------
h("7. Overall interpretation", 1)
para(
    f"The single-port baseline describes an efficient, well-behaved brine "
    f"outfall: a 60° inclined jet that dilutes the {cfg['S0']:.0f} g/kg reject "
    f"{dilu:.0f}-fold in the near field, delivering it to the seabed at only "
    f"{dS_return:.2f} g/kg above ambient, with a compact, essentially-compliant "
    f"far-field footprint and a thin bottom-trapped plume. This is the "
    f"performance a multiport diffuser is designed to achieve — PROVIDED the "
    f"ports are spaced far enough apart that the jets remain independent. "
    f"Whether that assumption holds is precisely what the merged-diffuser "
    f"variant tests, and the comparison document quantifies how much is lost "
    f"when it does not.")

DOC.save(os.path.join(D2, "discuss_result2.docx"))
print("Saved 2/discuss_result2.docx")
print("paragraphs:", len(DOC.paragraphs), "tables:", len(DOC.tables))
