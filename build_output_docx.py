# -*- coding: utf-8 -*-
"""
Generator for output.docx — the complete catalogue of outputs produced by
the NEREID-B coupled stochastic PDE brine-salinity model: primary fields,
derived metrics, charts / graphs / curves, uncertainty products, compliance
diagnostics, validation diagnostics and export formats.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = Document()
BODY_FONT = "Calibri"
EQ_FONT = "Cambria Math"

normal = DOC.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)


def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def h(text, level=1):
    return DOC.add_heading(text, level=level)


def para(text="", bold=False, italic=False, size=11, align=None, color=None,
         space_after=6):
    p = DOC.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = BODY_FONT
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def bullet(text, sub=False):
    style = "List Bullet 2" if sub else "List Bullet"
    p = DOC.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p


def make_table(header, rows, col_widths=None, font_size=9, mono_first=False):
    t = DOC.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = BODY_FONT
        _set_cell_bg(hdr[i], "1F4E79")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = EQ_FONT if (mono_first and i == 0) else BODY_FONT
    if col_widths:
        for i, w in enumerate(col_widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)
    return t


# ======================================================================
#  TITLE
# ======================================================================
title = DOC.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("NEREID-B — Output Specification")
tr.bold = True
tr.font.size = Pt(24)
tr.font.name = BODY_FONT
tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = DOC.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Complete Catalogue of Predicted Outputs, Metrics, "
                 "Charts, Graphs, Curves and Data Products of the Coupled "
                 "Stochastic PDE Brine-Salinity Model")
sr.italic = True
sr.font.size = Pt(13)
sr.font.name = BODY_FONT

meta = DOC.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Companion to salinity.docx  ·  Rev. 1.0  ·  11 June 2026")
mr.font.size = Pt(11)
mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
mr.font.name = BODY_FONT

para("", space_after=8)
para(
    "This document enumerates every output the NEREID-B model generates to "
    "deliver its prediction: the raw solved fields, the derived engineering "
    "metrics, the full suite of visualizations (charts, graphs, curves, maps "
    "and animations), the statistical / uncertainty products that arise from "
    "the model's stochastic layer, the regulatory-compliance diagnostics, the "
    "model-health diagnostics, and the machine-readable export formats. "
    "Outputs are grouped by category; within each group a table gives the "
    "output name, what it is, its dimensionality / chart type, and its units "
    "or file format. A consolidated quick-reference index closes the document.",
)

# ======================================================================
#  0. OUTPUT HIERARCHY OVERVIEW
# ======================================================================
h("0. Output Hierarchy at a Glance", 1)
para("NEREID-B outputs fall into eight tiers, each built from the one above:")
bullet("Tier 1 — Primary solved fields (4-D: x, y, z, t).")
bullet("Tier 2 — Derived secondary fields computed from Tier 1.")
bullet("Tier 3 — Scalar engineering metrics (the headline numbers).")
bullet("Tier 4 — Spatial maps and cross-sections (2-D snapshots).")
bullet("Tier 5 — Curves and graphs (1-D relationships and profiles).")
bullet("Tier 6 — Time series and animations (evolution in time).")
bullet("Tier 7 — Statistical / uncertainty products (from the SPDE ensemble).")
bullet("Tier 8 — Compliance, validation and machine-readable data exports.")

# ======================================================================
#  1. PRIMARY SOLVED FIELDS
# ======================================================================
h("1. Tier 1 — Primary Solved Fields (4-D space–time)", 1)
para("These are the dependent variables the PDE system integrates directly. "
     "Each is a full four-dimensional array S(x, y, z, t):")
make_table(
    ["Output", "Description", "Dim.", "Units"],
    [
        ["S(x,y,z,t)", "Absolute salinity field — the primary prediction", "4-D", "g kg⁻¹ (psu)"],
        ["u, v, w", "Three velocity components of the moving sea + plume", "4-D", "m s⁻¹"],
        ["p(x,y,z,t)", "Pressure field", "4-D", "Pa"],
        ["ρ(x,y,z,t)", "Density field (from nonlinear EOS)", "4-D", "kg m⁻³"],
        ["T(x,y,z,t)", "Temperature field", "4-D", "°C"],
        ["k(x,y,z,t)", "Turbulent kinetic energy", "4-D", "m² s⁻²"],
        ["ε(x,y,z,t)", "Turbulent dissipation rate", "4-D", "m² s⁻³"],
        ["η(x,y,t)", "Free-surface elevation (incl. splash zone)", "3-D", "m"],
        ["α(x,y,z,t)", "Air–water volume fraction (interface)", "4-D", "–"],
        ["Π(x,y,z,t)", "Osmotic-pressure field", "4-D", "Pa"],
    ],
    col_widths=[1.1, 3.2, 0.6, 1.1], mono_first=True,
)

# ======================================================================
#  2. DERIVED SECONDARY FIELDS
# ======================================================================
h("2. Tier 2 — Derived Secondary Fields", 1)
para("Computed pointwise from Tier 1; each is also a full field that can be "
     "mapped, sliced or animated:")
make_table(
    ["Output", "Description", "Dim.", "Units"],
    [
        ["ΔS = S − S_amb", "Excess (above-ambient) salinity — the impact field", "4-D", "g kg⁻¹"],
        ["𝒟(x,y,z,t)", "Dilution field, (S₀−S_amb)/(S−S_amb)", "4-D", "– (×)"],
        ["μ_t(x,y,z,t)", "Turbulent eddy viscosity / diffusivity", "4-D", "m² s⁻¹"],
        ["|u|, speed", "Current speed magnitude", "4-D", "m s⁻¹"],
        ["ω = ∇×u", "Vorticity field (vortices, eddies)", "4-D", "s⁻¹"],
        ["Ri(x,y,z,t)", "Gradient Richardson number (mixing state)", "4-D", "–"],
        ["N²(x,y,z,t)", "Brunt–Väisälä stratification frequency", "4-D", "s⁻²"],
        ["Fr_d local", "Local densimetric Froude number", "4-D", "–"],
        ["b = g(ρ−ρ_a)/ρ_a", "Buoyancy field of the plume", "4-D", "m s⁻²"],
        ["∇S, |∇S|", "Salinity gradient & front sharpness", "4-D", "g kg⁻¹ m⁻¹"],
        ["E_rate", "Local entrainment / mixing rate", "4-D", "s⁻¹"],
        ["τ_res", "Residence / flushing time field", "3-D", "s"],
    ],
    col_widths=[1.2, 3.1, 0.6, 1.1], mono_first=True,
)

# ======================================================================
#  3. SCALAR ENGINEERING METRICS
# ======================================================================
h("3. Tier 3 — Scalar Engineering Metrics (the headline numbers)", 1)
para("Single-number diagnostics, each reported as an instantaneous value, a "
     "tidal-cycle envelope, and a probabilistic value with confidence bounds "
     "(see Tier 7). These directly answer 'how salty, how far, how deep'.")

h("3.1 Dilution & concentration metrics", 2)
make_table(
    ["Metric", "Description", "Units"],
    [
        ["S_max", "Peak salinity anywhere in the field", "g kg⁻¹"],
        ["𝒟_min", "Minimum (worst-case) dilution", "– (×)"],
        ["𝒟_centerline", "Centerline dilution along the jet", "– (×)"],
        ["𝒟_impact", "Dilution at seabed return / impact point", "– (×)"],
        ["𝒟_nf, 𝒟_ff", "Near-field and far-field dilution", "– (×)"],
        ["ΔS_max", "Maximum excess salinity above ambient", "g kg⁻¹"],
        ["ΔS_seabed", "Excess salinity at the seabed", "g kg⁻¹"],
    ],
    col_widths=[1.0, 3.6, 1.1], mono_first=True,
)

h("3.2 Geometric reach metrics (how far & how deep)", 2)
make_table(
    ["Metric", "Description", "Units"],
    [
        ["r_max", "Maximum horizontal reach of the ΔS_crit footprint", "m"],
        ["z_max", "Maximum depth of brine impact", "m"],
        ["z_rise", "Terminal rise height of the dense jet", "m"],
        ["x_return", "Horizontal distance to seabed return point", "m"],
        ["L_runout", "Run-out length of the seabed gravity current", "m"],
        ["b_plume", "Plume width / lateral half-width", "m"],
        ["h_layer", "Thickness of the spreading bottom layer", "m"],
        ["𝒜_footprint", "Seabed area with ΔS > ΔS_crit", "m²"],
        ["V_affected", "Water volume with ΔS > ΔS_crit", "m³"],
        ["θ_traj", "Jet trajectory / bend-over angle", "deg"],
    ],
    col_widths=[1.0, 3.6, 1.1], mono_first=True,
)

h("3.3 Mixing & dynamics metrics", 2)
make_table(
    ["Metric", "Description", "Units"],
    [
        ["Fr_d", "Discharge densimetric Froude number", "–"],
        ["Re", "Discharge Reynolds number", "–"],
        ["Ri_bulk", "Bulk Richardson number of the plume", "–"],
        ["E_total", "Total entrainment coefficient", "–"],
        ["τ_flush", "Mixing-zone flushing time", "s / h"],
        ["Ṁ_salt", "Net salt mass flux through control surfaces", "kg s⁻¹"],
        ["Π*", "Osmotic-to-inertial pressure ratio (novel)", "–"],
    ],
    col_widths=[1.0, 3.6, 1.1], mono_first=True,
)

# ======================================================================
#  4. SPATIAL MAPS & CROSS-SECTIONS
# ======================================================================
DOC.add_page_break()
h("4. Tier 4 — Spatial Maps and Cross-Sections (2-D charts)", 1)
para("Two-dimensional graphical products extracted from the 4-D fields, at "
     "user-chosen times and slice planes. Each is a colour/contour plot.")
make_table(
    ["Chart", "What it shows", "Plane / type"],
    [
        ["Plan-view salinity map", "Salinity over x–y at a chosen depth (e.g. seabed)", "horizontal contour"],
        ["Vertical cross-section", "Salinity over a vertical transect through the plume", "x–z / y–z contour"],
        ["Dilution map", "Dilution factor field", "contour / heatmap"],
        ["Excess-salinity (ΔS) map", "Above-ambient salinity footprint", "filled contour"],
        ["Plume iso-surface (3-D)", "3-D surface of S = S_crit (the plume body)", "3-D render"],
        ["Velocity vector / quiver map", "Current + plume flow direction & speed", "vector field"],
        ["Streamline / pathline map", "Flow and tracer paths around the outfall", "streamlines"],
        ["Vorticity map", "Eddies and vortices shed by the jet", "contour"],
        ["Density / stratification map", "ρ field and pycnocline position", "contour"],
        ["Temperature map", "Thermal field & thermal plume", "contour"],
        ["Turbulence (k, ε, μ_t) map", "Mixing-intensity distribution", "contour"],
        ["Richardson-number map", "Where mixing is active vs suppressed", "contour"],
        ["Bathymetry overlay", "Plume draped on seabed/continental-shelf terrain", "3-D terrain"],
        ["Bed-shear / scour map", "Bed stress under the gravity current", "contour"],
    ],
    col_widths=[1.7, 3.2, 1.3],
)

# ======================================================================
#  5. CURVES & GRAPHS
# ======================================================================
h("5. Tier 5 — Curves, Graphs and Profiles (1-D relationships)", 1)
para("Line plots that quantify how key variables change with distance, "
     "depth, angle or driving parameter — the classic engineering design "
     "curves.")
make_table(
    ["Curve / graph", "x-axis vs y-axis", "Purpose"],
    [
        ["Centerline dilution curve", "distance along plume vs dilution 𝒟", "core design curve"],
        ["Salinity decay curve", "distance from outfall vs S (or ΔS)", "how fast brine dilutes"],
        ["Plume trajectory curve", "horizontal x vs vertical z of centerline", "rise & fall path"],
        ["Vertical salinity profile", "S vs depth z at a station", "stratification of impact"],
        ["Vertical density / T profiles", "ρ, T vs depth", "ambient & plume structure"],
        ["Velocity profile", "speed vs depth (boundary layer)", "current shear"],
        ["Concentration–distance decay", "log S vs log r (power-law fit)", "scaling-law extraction"],
        ["Footprint-vs-time curve", "time vs affected area 𝒜", "growth of impact"],
        ["r_max & z_max vs time", "time vs reach / depth", "envelope of impact"],
        ["Dilution vs Froude number", "Fr_d vs 𝒟", "design sensitivity"],
        ["Dilution vs nozzle angle", "θ vs 𝒟_impact", "diffuser-angle optimisation"],
        ["Dilution vs current speed", "|U| vs 𝒟", "effect of ambient currents"],
        ["Salinity vs tidal phase", "tidal phase vs S at receptor", "tidal modulation"],
        ["Cumulative exposure curve", "ΔS threshold vs exposed area/volume", "dose–area relationship"],
        ["Power spectrum of S", "frequency vs spectral energy", "variability / periodicity"],
    ],
    col_widths=[1.8, 2.3, 2.1],
)

# ======================================================================
#  6. TIME SERIES & ANIMATIONS
# ======================================================================
h("6. Tier 6 — Time Series and Animations (unsteady evolution)", 1)
para("Because the model is unsteady and the sea is moving, every metric and "
     "field is also produced as a function of time.")
make_table(
    ["Output", "Description", "Format"],
    [
        ["S(t) at monitoring stations", "Salinity time series at user-placed receptors / CTD points", "line series"],
        ["Tidal-cycle evolution", "Field snapshots through flood–ebb cycle", "frame set / animation"],
        ["Storm / event response", "Plume response to a wind or current event", "animation"],
        ["Plume-spread animation", "Time-lapse of the salinity footprint growing", "MP4 / GIF"],
        ["3-D plume animation", "Rotating iso-surface evolving in time", "MP4"],
        ["Metric time histories", "r_max(t), z_max(t), 𝒜(t), 𝒟_min(t)", "multi-line plot"],
        ["Startup / shutdown transient", "Response to discharge turning on/off", "line series"],
        ["Spin-up convergence", "Field statistics approaching quasi-steady state", "line series"],
    ],
    col_widths=[1.8, 3.2, 1.2],
)

# ======================================================================
#  7. STATISTICAL / UNCERTAINTY PRODUCTS
# ======================================================================
DOC.add_page_break()
h("7. Tier 7 — Statistical & Uncertainty Products (from the SPDE ensemble)", 1)
para("Because NEREID-B is a stochastic PDE forced by the unpredictable sea, "
     "wind and turbulence, it runs as a Monte-Carlo ensemble and returns "
     "probability distributions, not just single values. These products are a "
     "distinguishing output of the model.")
make_table(
    ["Output", "Description", "Type"],
    [
        ["⟨S⟩(x,t)", "Ensemble-mean salinity field", "4-D field"],
        ["σ_S(x,t)", "Salinity standard-deviation (spread) field", "4-D field"],
        ["P[S(x,t)]", "Full probability density of salinity at each point", "PDF field"],
        ["Percentile fields", "5th / 50th / 95th percentile salinity fields", "4-D fields"],
        ["Exceedance-probability map", "ℙ(ΔS > ΔS_crit) — risk map", "2-D / 3-D map"],
        ["Confidence-envelope bands", "Shaded ±90 % band on every curve & metric", "band plot"],
        ["Salinity histogram / PDF", "Distribution of S at a chosen receptor", "histogram"],
        ["Exceedance (CDF) curve", "ℙ(S > threshold) vs threshold", "curve"],
        ["Return-period estimates", "Expected recurrence of a salinity level", "table / curve"],
        ["Confidence on r_max, z_max", "Probabilistic reach & depth of impact", "value ± CI"],
        ["Ensemble spaghetti plot", "All realizations of the plume edge overlaid", "multi-line"],
        ["Sensitivity / tornado chart", "Ranking of inputs driving output variance", "bar chart"],
        ["Sobol / variance indices", "Quantified contribution of each parameter", "table"],
        ["Risk / probability rose", "Directional exceedance vs current/wind direction", "rose plot"],
    ],
    col_widths=[1.6, 3.3, 1.3], mono_first=True,
)

# ======================================================================
#  8. COMPLIANCE & ENVIRONMENTAL DIAGNOSTICS
# ======================================================================
h("8. Tier 8a — Regulatory & Environmental Compliance Diagnostics", 1)
para("Outputs framed against permit and ecological limits, supporting "
     "Environmental Impact Assessment and sustainability reporting.")
make_table(
    ["Output", "Description", "Units"],
    [
        ["Mixing-zone boundary", "Locus where ΔS falls below the permitted limit", "polygon / m"],
        ["Regulatory mixing-zone area", "Area enclosed by the compliance boundary", "m²"],
        ["Compliance pass/fail flags", "Status at each defined receptor / station", "boolean"],
        ["Threshold-exceedance duration", "Time ΔS > limit at a location", "h"],
        ["Seabed exposure area", "Benthic area above ecological salinity limit", "m²"],
        ["ΔS at sensitive habitats", "Salinity rise at seagrass / reef receptors", "g kg⁻¹"],
        ["Cumulative dose maps", "Time-integrated ΔS exposure (∫ΔS dt)", "g kg⁻¹·h"],
        ["Far-field background rise", "Long-term ambient salinity increment", "g kg⁻¹"],
        ["Compliance summary table", "Limit vs predicted vs margin per receptor", "table"],
    ],
    col_widths=[1.7, 3.2, 1.3],
)

# ======================================================================
#  9. MODEL-HEALTH / VALIDATION DIAGNOSTICS
# ======================================================================
h("8. Tier 8b — Model-Health and Validation Diagnostics", 1)
para("Outputs that demonstrate the solution is trustworthy.")
make_table(
    ["Output", "Description", "Units"],
    [
        ["Salt & mass balance closure", "Global conservation error vs time", "% / kg"],
        ["Residual convergence", "Per-equation residual history", "–"],
        ["CFL / time-step record", "Stability metric over the run", "–"],
        ["Grid-convergence (GCI) report", "Solution sensitivity to mesh resolution", "%"],
        ["Entropy-production check", "Positivity of Onsager fluxes (consistency)", "≥0"],
        ["Model-vs-measurement scatter", "Predicted vs observed S at survey points", "scatter + R², RMSE"],
        ["Skill scores", "RMSE, bias, Nash–Sutcliffe, Willmott index", "–"],
        ["Calibration parameter log", "Fitted coefficients & their uncertainty", "table"],
    ],
    col_widths=[1.8, 3.1, 1.3],
)

# ======================================================================
#  10. EXPORT FORMATS
# ======================================================================
h("9. Tier 8c — Machine-Readable Data Exports", 1)
para("Every field, metric and figure is also emitted in standard formats for "
     "downstream GIS, reporting and coupling to other models.")
make_table(
    ["Product", "Contents", "Format"],
    [
        ["4-D field archive", "All Tier 1–2 fields with coordinates & time", "NetCDF / HDF5"],
        ["3-D scene files", "Iso-surfaces, slices, vectors for 3-D viewers", "VTK / VTU"],
        ["Geo-referenced maps", "Footprint, exceedance & mixing-zone maps", "GeoTIFF / shapefile / KML"],
        ["Metric time series", "All scalar metrics vs time", "CSV / Parquet"],
        ["Statistical summary", "Means, percentiles, exceedance, Sobol indices", "CSV / JSON"],
        ["Figures & charts", "All curves, maps, profiles", "PNG / SVG / PDF"],
        ["Animations", "Time-lapse and 3-D fly-throughs", "MP4 / GIF"],
        ["Compliance report", "Auto-generated EIA-style summary", "PDF / DOCX"],
    ],
    col_widths=[1.6, 3.3, 1.3],
)

# ======================================================================
#  11. CONSOLIDATED INDEX
# ======================================================================
DOC.add_page_break()
h("10. Consolidated Quick-Reference Index of Outputs", 1)
para("The headline deliverables, in priority order, that together constitute "
     "the model's prediction:")
make_table(
    ["#", "Output", "Answers the question"],
    [
        ["1", "4-D salinity field S(x,y,z,t)", "What is the salinity everywhere, at all times?"],
        ["2", "Excess-salinity ΔS & dilution 𝒟 fields", "How much above ambient / how diluted?"],
        ["3", "r_max & z_max (with confidence)", "How FAR and how DEEP does the brine reach?"],
        ["4", "Seabed footprint area 𝒜 & affected volume", "How large is the impacted zone?"],
        ["5", "Plume trajectory & rise/return geometry", "Where does the jet go and land?"],
        ["6", "Centerline-dilution & decay curves", "How quickly does it dilute with distance?"],
        ["7", "Exceedance-probability / risk maps", "With what probability is a limit exceeded?"],
        ["8", "Confidence envelopes on all metrics", "How certain is the prediction?"],
        ["9", "Mixing-zone boundary & compliance flags", "Does it meet environmental limits?"],
        ["10", "Time series & animations", "How does it evolve in the unsteady sea?"],
        ["11", "Vertical profiles of S, T, ρ, u", "What is the vertical structure?"],
        ["12", "Vorticity, turbulence & Richardson maps", "How and where does mixing occur?"],
        ["13", "Sensitivity / Sobol charts", "Which inputs control the outcome?"],
        ["14", "Validation scatter & skill scores", "Can the prediction be trusted?"],
        ["15", "NetCDF / GeoTIFF / CSV / report exports", "How is it delivered and reused?"],
    ],
    col_widths=[0.4, 2.5, 3.3],
)

para("", space_after=8)
para(
    "Note on possibly-overlooked outputs now included above: vorticity/eddy "
    "fields, residence (flushing) time, bed-shear / scour potential, "
    "cumulative exposure (dose) maps, return-period statistics, Sobol "
    "variance indices, power spectra of salinity variability, entropy-"
    "production consistency checks, and directional risk roses — these are "
    "frequently omitted from conventional brine-plume reporting but are "
    "produced natively by NEREID-B because of its stochastic, fully-coupled "
    "formulation.",
    italic=True,
)

DOC.save("/home/akosa/salinity_prediction/output.docx")
print("Saved output.docx")
